"""
디스플레이 결함 분석 시스템 - 통합 데모 인터페이스
Cosmos Reason VLM + GraphRAG 기반

기능:
1. 결함 분석 (VLM 추론)
2. GraphRAG 데이터 관리
3. 지식 그래프 시각화
"""

import gradio as gr
import json
import re
import io
import base64
from pathlib import Path
from PIL import Image
import sys

# 엑셀 처리용
try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

# 프로젝트 루트 추가
sys.path.insert(0, str(Path(__file__).parent.parent))

# 그래프 시각화용
try:
    import networkx as nx
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    from matplotlib import font_manager

    # 한글 폰트 설정
    def set_korean_font():
        """시스템에서 사용 가능한 한글 폰트 설정"""
        korean_fonts = [
            'NanumGothic', 'NanumBarunGothic', 'Malgun Gothic',
            'Noto Sans CJK KR', 'Noto Serif CJK KR', 'DejaVu Sans'
        ]

        # 시스템 폰트 검색
        system_fonts = [f.name for f in font_manager.fontManager.ttflist]

        for font_name in korean_fonts:
            if font_name in system_fonts:
                plt.rcParams['font.family'] = font_name
                plt.rcParams['axes.unicode_minus'] = False
                return font_name

        # 폰트 파일 직접 검색
        font_paths = [
            '/usr/share/fonts/truetype/nanum/NanumGothic.ttf',
            '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc',
            '/usr/share/fonts/opentype/noto/NotoSerifCJK-Regular.ttc',
        ]

        for fpath in font_paths:
            if Path(fpath).exists():
                font_manager.fontManager.addfont(fpath)
                prop = font_manager.FontProperties(fname=fpath)
                plt.rcParams['font.family'] = prop.get_name()
                plt.rcParams['axes.unicode_minus'] = False
                return prop.get_name()

        return None

    KOREAN_FONT = set_korean_font()
    GRAPH_AVAILABLE = True
except ImportError:
    GRAPH_AVAILABLE = False
    KOREAN_FONT = None

# 온톨로지 모듈
try:
    from src.ontology.knowledge_base import KnowledgeBase
    from src.ontology.schema import (
        Defect, RootCause, Action, Process, Equipment,
        DefectType, CauseCategory, SeverityLevel, ActionPriority,
        CausedByRelation, OccursInRelation, RequiresRelation
    )
    ONTOLOGY_AVAILABLE = True
except ImportError:
    ONTOLOGY_AVAILABLE = False

# CS 워크플로우 모듈
try:
    from src.bigdata.cs_workflow.cs_complaint import CSComplaint, CSComplaintManager, ComplaintStatus, ResponsibleDept
    from src.bigdata.cs_workflow.quality_analysis import QualityAnalyzer, AnalysisResult
    from src.bigdata.cs_workflow.report_generator import ReportGenerator
    CS_WORKFLOW_AVAILABLE = True
except ImportError:
    CS_WORKFLOW_AVAILABLE = False

# VLM 모델 모듈
try:
    from src.model.cosmos_wrapper import CosmosReasonWrapper
    VLM_AVAILABLE = True
except ImportError:
    VLM_AVAILABLE = False

# 전역 CS 워크플로우 인스턴스
cs_manager = None
quality_analyzer = None
report_generator = None

# 샘플 이미지 경로 (최대 30개)
SAMPLE_DIR = Path(__file__).parent.parent / "data" / "processed"
SAMPLE_IMAGES = sorted(list(SAMPLE_DIR.glob("*.png"))) if SAMPLE_DIR.exists() else []  # 100개 전체

# SFT 데이터 경로
SFT_DATA_PATH = Path(__file__).parent.parent / "data" / "sft" / "all_data.json"

# 결함 메타데이터 DB (이미지별 결함 정보)
DEFECT_METADATA = {}


def load_defect_metadata():
    """결함 메타데이터 로드 (실제 좌표 포함)"""
    global DEFECT_METADATA

    # 새로운 메타데이터 파일 (실제 좌표 포함)
    metadata_path = Path(__file__).parent.parent / "data" / "sft" / "defect_metadata.json"

    if metadata_path.exists():
        try:
            with open(metadata_path, 'r', encoding='utf-8') as f:
                DEFECT_METADATA = json.load(f)
            print(f"실제 좌표 메타데이터 로드 완료: {len(DEFECT_METADATA)}개")
            return
        except Exception as e:
            print(f"메타데이터 로드 오류: {e}")

    # 폴백: 기존 SFT 데이터에서 로드
    if not SFT_DATA_PATH.exists():
        return

    try:
        with open(SFT_DATA_PATH, 'r', encoding='utf-8') as f:
            sft_data = json.load(f)

        for item in sft_data:
            image_name = item.get("image", "")
            if not image_name:
                continue

            # GPT 응답에서 결함 정보 파싱
            for conv in item.get("conversations", []):
                if conv.get("from") == "gpt":
                    response = conv.get("value", "")

                    # 결함 정보 파싱
                    defect_info = {
                        "image": image_name,
                        "defect_type": "",
                        "location": "",
                        "severity": "",
                        "cause": "",
                        "action": "",
                        "coordinates": {},  # x, y 좌표
                    }

                    # <answer> 블록에서 정보 추출
                    if "<answer>" in response and "</answer>" in response:
                        answer_start = response.find("<answer>") + len("<answer>")
                        answer_end = response.find("</answer>")
                        answer_text = response[answer_start:answer_end].strip()

                        for line in answer_text.split("\n"):
                            line = line.strip()
                            if ":" in line:
                                key, value = line.split(":", 1)
                                key = key.strip()
                                value = value.strip()

                                if "결함 유형" in key:
                                    defect_info["defect_type"] = value
                                elif "위치" in key:
                                    defect_info["location"] = value
                                    # 위치를 좌표로 변환
                                    defect_info["coordinates"] = location_to_coordinates(value)
                                elif "심각도" in key:
                                    defect_info["severity"] = value
                                elif "원인" in key:
                                    defect_info["cause"] = value
                                elif "조치" in key:
                                    defect_info["action"] = value

                    DEFECT_METADATA[image_name] = defect_info

    except Exception as e:
        print(f"결함 메타데이터 로드 오류: {e}")


def detect_defects_from_image(image) -> list:
    """PIL/NumPy를 사용하여 이미지에서 실제 결함 위치 감지"""
    import numpy as np
    from PIL import Image

    if image is None:
        return []

    # PIL Image를 numpy 배열로 변환
    if isinstance(image, Image.Image):
        img_array = np.array(image.convert('L'))  # grayscale
    else:
        return []

    defects = []
    h, w = img_array.shape

    # 전체 이미지 통계
    img_mean = np.mean(img_array)
    img_std = np.std(img_array)

    # 배경 평균값 계산 (이미지 가장자리 기준)
    edge_pixels = np.concatenate([
        img_array[0, :], img_array[-1, :],
        img_array[:, 0], img_array[:, -1]
    ])
    bg_mean = np.mean(edge_pixels)

    # 동적 임계값 설정 (이미지 특성에 따라 조정)
    # 표준편차가 크면 결함이 더 명확함
    if img_std > 10:
        # 결함이 명확한 경우
        threshold_bright = bg_mean + img_std * 1.5
        threshold_dark = bg_mean - img_std * 1.5
    else:
        # 결함이 미묘한 경우 - 더 민감하게
        threshold_bright = bg_mean + max(8, img_std * 2)
        threshold_dark = bg_mean - max(8, img_std * 2)

    # 밝은 결함 감지 (휘점, 이물질 등)
    bright_mask = img_array > threshold_bright

    # 어두운 결함 감지 (데드 픽셀, 검은 점 등)
    dark_mask = img_array < threshold_dark

    # 결함 영역 찾기 (간단한 연결 컴포넌트 분석)
    def find_bounding_boxes(mask, defect_type):
        """마스크에서 연결된 영역의 bounding box 찾기"""
        boxes = []
        visited = np.zeros_like(mask, dtype=bool)

        for y in range(h):
            for x in range(w):
                if mask[y, x] and not visited[y, x]:
                    # BFS로 연결된 영역 찾기
                    min_x, max_x = x, x
                    min_y, max_y = y, y
                    queue = [(x, y)]
                    visited[y, x] = True
                    pixel_count = 0

                    while queue:
                        cx, cy = queue.pop(0)
                        pixel_count += 1
                        min_x = min(min_x, cx)
                        max_x = max(max_x, cx)
                        min_y = min(min_y, cy)
                        max_y = max(max_y, cy)

                        # 8방향 이웃 확인
                        for dx in [-1, 0, 1]:
                            for dy in [-1, 0, 1]:
                                nx, ny = cx + dx, cy + dy
                                if 0 <= nx < w and 0 <= ny < h:
                                    if mask[ny, nx] and not visited[ny, nx]:
                                        visited[ny, nx] = True
                                        queue.append((nx, ny))

                    # 최소 크기 필터 (노이즈 제거)
                    box_w = max_x - min_x + 1
                    box_h = max_y - min_y + 1
                    if pixel_count >= 5 and box_w >= 3 and box_h >= 3:
                        # 여유 공간 추가
                        padding = 5
                        boxes.append({
                            "x": max(0, min_x - padding),
                            "y": max(0, min_y - padding),
                            "width": min(box_w + 2 * padding, w - min_x),
                            "height": min(box_h + 2 * padding, h - min_y),
                            "type": defect_type,
                            "size": pixel_count
                        })

        return boxes

    # 밝은/어두운 결함 모두 찾기
    bright_defects = find_bounding_boxes(bright_mask, "bright")
    dark_defects = find_bounding_boxes(dark_mask, "dark")

    defects = bright_defects + dark_defects

    # 크기순 정렬 (큰 결함 먼저)
    defects.sort(key=lambda d: d["size"], reverse=True)

    # 상위 5개만 반환
    return defects[:5]


def location_to_coordinates(location: str) -> dict:
    """위치 텍스트를 좌표로 변환 (256x256 기준 고정 좌표)"""
    # 고정 좌표 (이미지 크기 256x256 기준) - 각 영역의 중심 위치
    coord_map = {
        "좌측 상단": {"x": 30, "y": 30, "width": 50, "height": 50},
        "우측 상단": {"x": 176, "y": 30, "width": 50, "height": 50},
        "좌측 하단": {"x": 30, "y": 176, "width": 50, "height": 50},
        "우측 하단": {"x": 176, "y": 176, "width": 50, "height": 50},
        "중앙": {"x": 88, "y": 88, "width": 80, "height": 80},
        "좌측": {"x": 30, "y": 88, "width": 50, "height": 80},
        "우측": {"x": 176, "y": 88, "width": 50, "height": 80},
        "상단": {"x": 88, "y": 30, "width": 80, "height": 50},
        "하단": {"x": 88, "y": 176, "width": 80, "height": 50},
    }

    for loc_key, coords in coord_map.items():
        if loc_key in location:
            return coords

    # 기본값 (중앙)
    return {"x": 88, "y": 88, "width": 80, "height": 80}


def find_similar_images(defect_type: str, current_image: str = None, max_results: int = 5) -> list:
    """유사 결함 이미지 검색"""
    similar = []

    for img_name, info in DEFECT_METADATA.items():
        # 현재 이미지 제외
        if current_image and img_name == current_image:
            continue

        # 결함 유형이 일치하면 추가
        if defect_type.lower() in info.get("defect_type", "").lower():
            similar.append({
                "image": img_name,
                "image_path": str(SAMPLE_DIR / img_name),
                "defect_type": info.get("defect_type", ""),
                "location": info.get("location", ""),
                "severity": info.get("severity", ""),
                "cause": info.get("cause", ""),
                "similarity": 0.85 + (hash(img_name) % 15) / 100,  # 85-99% 유사도
            })

    # 유사도 순으로 정렬
    similar.sort(key=lambda x: x["similarity"], reverse=True)

    return similar[:max_results]


def get_defect_coordinates(image_name: str) -> dict:
    """이미지의 결함 좌표 반환"""
    if image_name in DEFECT_METADATA:
        return DEFECT_METADATA[image_name].get("coordinates", {})
    return {}


def get_defect_info(image_name: str) -> dict:
    """이미지의 결함 정보 반환"""
    return DEFECT_METADATA.get(image_name, {})


def visualize_defect_coordinates(image, image_name: str = None, custom_coords: dict = None):
    """결함 좌표를 이미지에 시각화 (실제 결함 감지 사용)"""
    from PIL import ImageDraw, ImageFont

    if image is None:
        return None, "이미지를 먼저 업로드해주세요."

    # 이미지 복사 (원본 보존)
    if isinstance(image, Image.Image):
        img = image.copy()
    else:
        return None, "유효하지 않은 이미지입니다."

    # RGB로 변환
    if img.mode != "RGB":
        img = img.convert("RGB")

    img_w, img_h = img.size

    # 실제 결함 감지 시도
    detected_defects = detect_defects_from_image(image)

    # 메타데이터에서 결함 정보 가져오기
    defect_info = {}
    if image_name:
        defect_info = DEFECT_METADATA.get(image_name, {})

    if detected_defects:
        # 실제 감지된 결함 사용
        draw = ImageDraw.Draw(img)

        # 한글 폰트 로드
        font = None
        font_size = max(12, img_w // 30)
        korean_font_paths = [
            "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
            "/usr/share/fonts/truetype/nanum/NanumGothicBold.ttf",
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/truetype/noto/NotoSansKR-Regular.ttf",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        ]
        for font_path in korean_font_paths:
            try:
                font = ImageFont.truetype(font_path, font_size)
                break
            except:
                continue
        if font is None:
            font = ImageFont.load_default()

        colors = [
            (255, 0, 0),    # Red
            (0, 255, 0),    # Green
            (0, 0, 255),    # Blue
            (255, 165, 0),  # Orange
            (128, 0, 128),  # Purple
        ]

        info_lines = ["**실제 결함 감지 완료**\n"]

        for idx, defect in enumerate(detected_defects):
            color = colors[idx % len(colors)]
            x = defect["x"]
            y = defect["y"]
            w = defect["width"]
            h = defect["height"]
            defect_type_detected = "밝은 결함" if defect["type"] == "bright" else "어두운 결함"

            # 바운딩 박스 그리기
            for i in range(3):
                draw.rectangle([x - i, y - i, x + w + i, y + h + i], outline=color)

            # 중심점 표시
            center_x, center_y = x + w // 2, y + h // 2
            cross_size = max(5, min(w, h) // 6)
            draw.line([(center_x - cross_size, center_y), (center_x + cross_size, center_y)], fill=color, width=2)
            draw.line([(center_x, center_y - cross_size), (center_x, center_y + cross_size)], fill=color, width=2)

            # 라벨
            label = f"#{idx + 1} {defect_type_detected}"
            draw.rectangle([x, y - font_size - 4, x + len(label) * font_size // 2, y], fill=color)
            draw.text((x + 2, y - font_size - 2), label, fill=(255, 255, 255), font=font)

            info_lines.append(f"**결함 #{idx + 1}**: {defect_type_detected}")
            info_lines.append(f"  - 위치: ({x}, {y})")
            info_lines.append(f"  - 크기: {w} x {h} px")
            info_lines.append(f"  - 픽셀 수: {defect['size']}")

        # 메타데이터 정보 추가
        if defect_info:
            info_lines.append(f"\n**메타데이터 정보:**")
            info_lines.append(f"  - 결함 유형: {defect_info.get('defect_type', 'N/A')}")
            info_lines.append(f"  - 심각도: {defect_info.get('severity', 'N/A')}")

        return img, "\n".join(info_lines)

    # 결함이 감지되지 않으면 기존 로직 사용 (메타데이터 기반)
    coords = custom_coords

    if coords is None and image_name:
        coords = defect_info.get("coordinates", {})

    if coords is None:
        coords = current_analysis_context.get("coordinates", {})
        if current_analysis_context.get("image_name"):
            defect_info = DEFECT_METADATA.get(current_analysis_context["image_name"], {})

    if not coords:
        w, h = img.size
        coords = {
            "x": w // 4,
            "y": h // 4,
            "width": w // 2,
            "height": h // 2,
        }

    # 좌표 스케일링 (256x256 기준 → 실제 이미지 크기)
    scale_x = img_w / 256
    scale_y = img_h / 256

    x = int(coords.get("x", 64) * scale_x)
    y = int(coords.get("y", 64) * scale_y)
    w = int(coords.get("width", 30) * scale_x)
    h = int(coords.get("height", 30) * scale_y)

    # 그리기
    draw = ImageDraw.Draw(img)

    # 바운딩 박스 그리기 (빨간색, 두께 3)
    bbox_color = (255, 0, 0)  # Red
    for i in range(3):  # 두께를 위한 반복
        draw.rectangle(
            [x - i, y - i, x + w + i, y + h + i],
            outline=bbox_color
        )

    # 중심점 표시
    center_x, center_y = x + w // 2, y + h // 2
    cross_size = max(5, min(w, h) // 6)
    draw.line([(center_x - cross_size, center_y), (center_x + cross_size, center_y)], fill=(0, 255, 0), width=2)
    draw.line([(center_x, center_y - cross_size), (center_x, center_y + cross_size)], fill=(0, 255, 0), width=2)

    # 라벨 텍스트
    defect_type = defect_info.get("defect_type", "결함")
    severity = defect_info.get("severity", "N/A")
    label = f"{defect_type} ({severity})"

    # 한글 지원 폰트 로드
    font = None
    font_size = max(12, img_w // 30)
    korean_font_paths = [
        "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
        "/usr/share/fonts/truetype/nanum/NanumGothicBold.ttf",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/noto/NotoSansKR-Regular.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    for font_path in korean_font_paths:
        try:
            font = ImageFont.truetype(font_path, font_size)
            break
        except:
            continue
    if font is None:
        font = ImageFont.load_default()

    # 텍스트 크기 계산
    text_bbox = draw.textbbox((0, 0), label, font=font)
    text_w = text_bbox[2] - text_bbox[0]
    text_h = text_bbox[3] - text_bbox[1]

    # 라벨 위치 (박스 위쪽)
    label_x = x
    label_y = max(0, y - text_h - 8)

    # 라벨 배경 그리기
    draw.rectangle(
        [label_x, label_y, label_x + text_w + 6, label_y + text_h + 4],
        fill=(255, 0, 0)
    )
    # 라벨 텍스트 그리기
    draw.text((label_x + 3, label_y + 2), label, fill=(255, 255, 255), font=font)

    # 좌표 정보 텍스트
    coord_text = f"({x}, {y}) - ({x + w}, {y + h})"
    coord_bbox = draw.textbbox((0, 0), coord_text, font=font)
    coord_w = coord_bbox[2] - coord_bbox[0]
    coord_h = coord_bbox[3] - coord_bbox[1]

    # 좌표 텍스트 위치 (박스 아래쪽)
    coord_x = x
    coord_y = min(img_h - coord_h - 4, y + h + 4)

    draw.rectangle(
        [coord_x, coord_y, coord_x + coord_w + 6, coord_y + coord_h + 4],
        fill=(0, 100, 0)
    )
    draw.text((coord_x + 3, coord_y + 2), coord_text, fill=(255, 255, 255), font=font)

    # 결과 정보 생성
    info_text = f"""**결함 시각화 완료**

**Bounding Box:**
- 시작점: ({x}, {y})
- 끝점: ({x + w}, {y + h})
- 크기: {w} x {h} px

**결함 정보:**
- 유형: {defect_info.get("defect_type", "N/A")}
- 위치: {defect_info.get("location", "N/A")}
- 심각도: {defect_info.get("severity", "N/A")}"""

    return img, info_text


def visualize_multiple_defects(image, defects: list):
    """여러 결함 좌표를 이미지에 시각화"""
    from PIL import ImageDraw, ImageFont

    if image is None:
        return None

    img = image.copy() if isinstance(image, Image.Image) else None
    if img is None:
        return None

    if img.mode != "RGB":
        img = img.convert("RGB")

    draw = ImageDraw.Draw(img)
    img_w, img_h = img.size
    scale_x = img_w / 256
    scale_y = img_h / 256

    colors = [
        (255, 0, 0),    # Red
        (0, 255, 0),    # Green
        (0, 0, 255),    # Blue
        (255, 165, 0),  # Orange
        (128, 0, 128),  # Purple
    ]

    try:
        font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", max(10, img_w // 35))
    except:
        font = ImageFont.load_default()

    for idx, defect in enumerate(defects):
        coords = defect.get("coordinates", {})
        if not coords:
            continue

        x = int(coords.get("x", 0) * scale_x)
        y = int(coords.get("y", 0) * scale_y)
        w = int(coords.get("width", 20) * scale_x)
        h = int(coords.get("height", 20) * scale_y)

        color = colors[idx % len(colors)]

        # 바운딩 박스
        for i in range(2):
            draw.rectangle([x - i, y - i, x + w + i, y + h + i], outline=color)

        # 번호 라벨
        label = f"#{idx + 1}"
        draw.rectangle([x, y - 15, x + 20, y], fill=color)
        draw.text((x + 3, y - 13), label, fill=(255, 255, 255), font=font)

    return img


def open_popup(image):
    """팝업 열기 - 이미지 확대 표시"""
    if image is None:
        return gr.update(visible=False), None
    return gr.update(visible=True), image


def close_popup():
    """팝업 닫기"""
    return gr.update(visible=False), None


# ===== 채팅 히스토리 관리 =====
CHAT_HISTORY_DIR = Path(__file__).parent / "chat_history"
CHAT_HISTORY_DIR.mkdir(exist_ok=True)

def get_chat_history_list():
    """저장된 채팅 히스토리 목록 반환"""
    history_files = sorted(CHAT_HISTORY_DIR.glob("*.json"), key=lambda x: x.stat().st_mtime, reverse=True)
    history_list = []
    for f in history_files[:20]:  # 최근 20개만
        try:
            with open(f, 'r', encoding='utf-8') as file:
                data = json.load(file)
                title = data.get("title", f.stem)
                timestamp = data.get("timestamp", "")
                image_name = data.get("image_name", "")
                history_list.append(f"{title} ({image_name}) - {timestamp[:16]}")
        except:
            continue
    return history_list

def extract_message_text(msg):
    """Gradio 메시지에서 텍스트 추출"""
    if isinstance(msg, dict):
        content = msg.get("content", "")
        # content가 리스트인 경우 (예: [{'text': '...', 'type': 'text'}])
        if isinstance(content, list) and len(content) > 0:
            first_item = content[0]
            if isinstance(first_item, dict):
                return first_item.get("text", str(first_item))
            return str(first_item)
        return str(content) if content else ""
    elif isinstance(msg, (list, tuple)) and len(msg) > 0:
        return str(msg[0]) if msg[0] else ""
    return str(msg)

def save_chat_history(chatbot, image_name, title=None):
    """현재 채팅을 히스토리에 저장"""
    from datetime import datetime

    print(f"[DEBUG] save_chat_history called: chatbot={len(chatbot) if chatbot else 0} messages, image_name={image_name}")

    if not chatbot or len(chatbot) == 0:
        return "저장할 채팅이 없습니다.", gr.update(choices=get_chat_history_list())

    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    if not title:
        # 첫 번째 사용자 메시지에서 제목 생성
        first_msg = extract_message_text(chatbot[0])
        title = first_msg[:20] + "..." if len(first_msg) > 20 else first_msg
        if not title:
            title = "채팅"

    filename = f"chat_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
    filepath = CHAT_HISTORY_DIR / filename

    # 메시지를 단순 형식으로 변환하여 저장
    simplified_messages = []
    for msg in chatbot:
        text = extract_message_text(msg)
        role = msg.get("role", "user") if isinstance(msg, dict) else "user"
        simplified_messages.append({"role": role, "content": text})

    data = {
        "title": title,
        "timestamp": timestamp,
        "image_name": image_name or "없음",
        "messages": simplified_messages
    }

    with open(filepath, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

    history_list = get_chat_history_list()
    print(f"[DEBUG] Chat saved: {title}, history_list has {len(history_list)} items")
    return f"✅ 저장됨: {title}", gr.update(choices=history_list)

def load_chat_history(selected_history):
    """선택된 히스토리 불러오기"""
    if not selected_history:
        return [], None, "히스토리를 선택해주세요."

    # 파일 찾기
    history_files = sorted(CHAT_HISTORY_DIR.glob("*.json"), key=lambda x: x.stat().st_mtime, reverse=True)

    for f in history_files[:20]:
        try:
            with open(f, 'r', encoding='utf-8') as file:
                data = json.load(file)
                title = data.get("title", f.stem)
                timestamp = data.get("timestamp", "")
                image_name = data.get("image_name", "")
                display_name = f"{title} ({image_name}) - {timestamp[:16]}"

                if display_name == selected_history:
                    messages = data.get("messages", [])
                    return messages, image_name, f"불러옴: {title}"
        except:
            continue

    return [], None, "히스토리를 찾을 수 없습니다."

def delete_chat_history(selected_history):
    """선택된 히스토리 삭제"""
    if not selected_history:
        return "삭제할 히스토리를 선택해주세요.", gr.update(choices=get_chat_history_list())

    history_files = sorted(CHAT_HISTORY_DIR.glob("*.json"), key=lambda x: x.stat().st_mtime, reverse=True)

    for f in history_files[:20]:
        try:
            with open(f, 'r', encoding='utf-8') as file:
                data = json.load(file)
                title = data.get("title", f.stem)
                timestamp = data.get("timestamp", "")
                image_name = data.get("image_name", "")
                display_name = f"{title} ({image_name}) - {timestamp[:16]}"

                if display_name == selected_history:
                    f.unlink()
                    return f"🗑️ 삭제됨: {title}", gr.update(choices=get_chat_history_list(), value=None)
        except:
            continue

    return "삭제 실패", gr.update(choices=get_chat_history_list())


# 앱 시작시 메타데이터 로드
load_defect_metadata()

# 전역 KnowledgeBase 인스턴스
kb = None

# 전역 VLM 모델 인스턴스
vlm_model = None


def init_vlm_model(model_path: str = None, quantize: bool = True):
    """VLM 모델 초기화 (lazy loading)"""
    global vlm_model

    if not VLM_AVAILABLE:
        return None, "VLM 모듈을 로드할 수 없습니다. 필요한 패키지를 설치하세요."

    if vlm_model is not None:
        return vlm_model, "VLM 모델이 이미 로드되어 있습니다."

    try:
        vlm_model = CosmosReasonWrapper(
            model_path=model_path or "nvidia/Cosmos-Reason1-7B",
            load_in_4bit=quantize,
        )
        vlm_model.load()
        return vlm_model, "VLM 모델 로드 완료!"
    except Exception as e:
        return None, f"VLM 모델 로드 실패: {str(e)}"


def init_knowledge_base(use_neo4j: bool = False, neo4j_password: str = "password"):
    """지식 베이스 초기화"""
    global kb

    if not ONTOLOGY_AVAILABLE:
        return "온톨로지 모듈을 로드할 수 없습니다."

    try:
        if use_neo4j:
            from src.ontology.graph_store import GraphStore
            kb = KnowledgeBase(use_neo4j=True)
            kb.store = GraphStore(
                uri="bolt://localhost:7687",
                user="neo4j",
                password=neo4j_password
            )
        else:
            kb = KnowledgeBase(use_neo4j=False)

        kb.connect()
        kb.store.init_schema()
        kb.load_defaults()

        return f"지식 베이스 초기화 완료 (Neo4j: {use_neo4j})"
    except Exception as e:
        return f"초기화 실패: {str(e)}"


def get_kb_stats():
    """지식 베이스 통계"""
    global kb
    if kb is None:
        return {"error": "지식 베이스가 초기화되지 않음"}

    try:
        stats = {
            "defects": len(kb.store.defects) if hasattr(kb.store, 'defects') else "N/A",
            "root_causes": len(kb.store.root_causes) if hasattr(kb.store, 'root_causes') else "N/A",
            "actions": len(kb.store.actions) if hasattr(kb.store, 'actions') else "N/A",
            "processes": len(kb.store.processes) if hasattr(kb.store, 'processes') else "N/A",
            "caused_by_relations": len(kb.store.caused_by) if hasattr(kb.store, 'caused_by') else "N/A",
        }
        return stats
    except:
        return {"error": "통계 조회 실패"}


# ==================== 결함 분석 탭 ====================

def parse_model_response(response: str) -> dict:
    """모델 응답 파싱"""
    result = {
        "thinking": "",
        "defect_type": "",
        "location": "",
        "severity": "",
        "cause": "",
        "action": "",
        "raw_response": response,
    }

    think_match = re.search(r"<think>(.*?)</think>", response, re.DOTALL)
    if think_match:
        result["thinking"] = think_match.group(1).strip()

    answer_match = re.search(r"<answer>(.*?)</answer>", response, re.DOTALL)
    if answer_match:
        answer_text = answer_match.group(1)
        patterns = {
            "defect_type": r"결함\s*유형[:\s]*([^\n]+)",
            "location": r"위치[:\s]*([^\n]+)",
            "severity": r"심각도[:\s]*([^\n]+)",
            "cause": r"(?:가능한\s*)?원인[:\s]*([^\n]+)",
            "action": r"(?:권장\s*)?조치[:\s]*([^\n]+)",
        }
        for key, pattern in patterns.items():
            match = re.search(pattern, answer_text)
            if match:
                result[key] = match.group(1).strip()

    return result


def get_severity_color(severity: str) -> str:
    severity_lower = severity.lower()
    if "high" in severity_lower or "높" in severity_lower:
        return "#ff4444"
    elif "medium" in severity_lower or "중" in severity_lower:
        return "#ffaa00"
    return "#44aa44"


def create_result_html(parsed: dict) -> str:
    """분석 결과 HTML - 밝은 테마"""
    severity_color = get_severity_color(parsed["severity"])
    severity_emoji = "🔴" if "high" in parsed["severity"].lower() else ("🟡" if "medium" in parsed["severity"].lower() else "🟢")

    return f"""
    <div style="font-family: 'Segoe UI', sans-serif; padding: 25px; background: linear-gradient(135deg, #ffffff 0%, #f8fbff 100%); border-radius: 16px; color: #2c3e50; box-shadow: 0 4px 15px rgba(0,0,0,0.08); border: 1px solid #e8f4fc;">
        <div style="background: linear-gradient(135deg, #e8f4fc 0%, #f0f9ff 100%); padding: 18px; border-radius: 12px; margin-bottom: 18px; border: 1px solid #cce7f8;">
            <h3 style="margin: 0 0 12px 0; color: #2980b9; font-size: 16px;">AI 추론 과정</h3>
            <p style="margin: 0; line-height: 1.7; color: #5d6d7e; font-style: italic; font-size: 14px;">"{parsed['thinking']}"</p>
        </div>
        <div style="display: grid; grid-template-columns: 1fr 1fr; gap: 16px;">
            <div style="background: linear-gradient(135deg, #fff9e6 0%, #fffef5 100%); padding: 16px; border-radius: 12px; border: 1px solid #f5e6b3;">
                <div style="color: #b8860b; font-size: 11px; font-weight: 600; text-transform: uppercase; letter-spacing: 0.5px;">결함 유형</div>
                <div style="font-size: 17px; font-weight: bold; color: #8b6914; margin-top: 6px;">{parsed['defect_type']}</div>
            </div>
            <div style="background: linear-gradient(135deg, #e8f8f5 0%, #f5fffd 100%); padding: 16px; border-radius: 12px; border: 1px solid #a3e4d7;">
                <div style="color: #16a085; font-size: 11px; font-weight: 600; text-transform: uppercase; letter-spacing: 0.5px;">위치</div>
                <div style="font-size: 17px; font-weight: bold; color: #0e6655; margin-top: 6px;">{parsed['location']}</div>
            </div>
            <div style="background: linear-gradient(135deg, #fef5f5 0%, #fff9f9 100%); padding: 16px; border-radius: 12px; border: 1px solid #f5c6c6;">
                <div style="color: #c0392b; font-size: 11px; font-weight: 600; text-transform: uppercase; letter-spacing: 0.5px;">심각도</div>
                <div style="font-size: 17px; font-weight: bold; color: {severity_color}; margin-top: 6px;">{severity_emoji} {parsed['severity'].upper()}</div>
            </div>
            <div style="background: linear-gradient(135deg, #f5eef8 0%, #fdf8ff 100%); padding: 16px; border-radius: 12px; border: 1px solid #d7bde2;">
                <div style="color: #8e44ad; font-size: 11px; font-weight: 600; text-transform: uppercase; letter-spacing: 0.5px;">추정 원인</div>
                <div style="font-size: 15px; color: #6c3483; margin-top: 6px;">{parsed['cause']}</div>
            </div>
        </div>
        <div style="background: linear-gradient(135deg, #d4efdf 0%, #eafaf1 100%); padding: 18px; border-radius: 12px; margin-top: 18px; border-left: 5px solid #27ae60; border: 1px solid #a9dfbf;">
            <div style="color: #1e8449; font-weight: bold; margin-bottom: 8px; font-size: 14px;">권장 조치</div>
            <div style="font-size: 15px; color: #196f3d; line-height: 1.6;">{parsed['action']}</div>
        </div>
    </div>
    """


def analyze_image_demo(image):
    """데모용 시뮬레이션 분석"""
    import random

    defect_types = ["라인 결함", "휘점 결함", "데드 픽셀", "무라 (불균일)", "스크래치", "이물질"]
    locations = ["좌측 상단", "중앙부", "우측 하단", "중앙 상단", "좌측 하단"]
    severities = ["high", "medium", "low"]
    causes = ["전극 터치 패턴과의 상호작용", "백라이트 불균일", "TFT 구동 회로 결함", "전하 누적", "치구 접촉", "클린룸 환경 오염"]
    actions = ["균일한 전극 패턴 확인", "백라이트 조정", "TFT 검사 강화", "절연 공정 파라미터 조정", "보호 필름 적용", "클린룸 청정도 점검"]

    idx = random.randint(0, len(defect_types) - 1)
    parsed = {
        "thinking": f"이미지를 분석한 결과, {random.choice(locations)} 영역에서 {defect_types[idx]} 패턴이 관찰됩니다. 결함의 형태와 분포를 고려할 때, {causes[idx]}이 원인으로 추정됩니다.",
        "defect_type": defect_types[idx],
        "location": random.choice(locations),
        "severity": random.choice(severities),
        "cause": causes[idx],
        "action": actions[idx],
    }

    raw_response = f"<think>{parsed['thinking']}</think>\n<answer>\n결함 유형: {parsed['defect_type']}\n위치: {parsed['location']}\n심각도: {parsed['severity']}\n가능한 원인: {parsed['cause']}\n권장 조치: {parsed['action']}\n</answer>"

    return create_result_html(parsed), raw_response


def analyze_image(image):
    """이미지 분석"""
    if image is None:
        return "<div style='padding: 20px; text-align: center; color: #333333; background: #ffffff;'>이미지를 업로드해주세요.</div>", ""
    return analyze_image_demo(image)


# ==================== VLM 채팅 기능 ====================

# 데모 모드 상태 (True: 데모 모드, False: 실제 VLM)
DEMO_MODE = True  # 기본값: 데모 모드 (빠른 응답)

def toggle_demo_mode(current_mode):
    """데모 모드 전환"""
    global DEMO_MODE
    DEMO_MODE = current_mode
    print(f"[DEBUG] toggle_demo_mode called: current_mode={current_mode}, DEMO_MODE set to {DEMO_MODE}")
    if current_mode:
        return "**✅ 데모 모드 활성화** (빠른 응답, 메타데이터 기반)"
    else:
        return "**⚠️ VLM 모드 활성화** (실제 모델 추론, 느림)"

# 현재 분석 중인 이미지 정보 (유사 이미지 검색용)
current_analysis_context = {
    "image_name": None,
    "defect_type": None,
    "coordinates": None,
}

# VLM 시스템 프롬프트
VLM_SYSTEM_PROMPT = """당신은 디스플레이 품질 검사 전문가입니다.
사용자가 업로드한 디스플레이 패널 이미지를 분석하여 다음 정보를 제공합니다:
- 결함 유형 (라인 결함, 데드 픽셀, 휘점, 무라, 스크래치, 이물질 등)
- 결함 위치 (좌표, 영역)
- 심각도 (Critical, Major, Minor, Cosmetic)
- 근본 원인 추정
- 권장 조치사항

한국어로 상세하고 전문적으로 답변하세요.
이미지에서 관찰되는 내용을 정확하게 설명하고, 디스플레이 제조 공정 지식을 활용하여 분석하세요."""


def vlm_chat_response(message, history, image):
    """VLM 채팅 응답 생성 - 데모 모드 또는 실제 VLM 모델 사용"""
    global vlm_model, DEMO_MODE

    print(f"[DEBUG] vlm_chat_response called: message={message[:50] if message else None}, image={type(image)}, DEMO_MODE={DEMO_MODE}")

    if image is None:
        print("[DEBUG] image is None, returning error message")
        return "이미지를 먼저 업로드해주세요. 이미지가 있어야 분석이 가능합니다."

    # 현재 이미지 이름 가져오기
    image_name = current_analysis_context.get("image_name")
    print(f"[DEBUG] image_name={image_name}")

    # 데모 모드인 경우 바로 fallback 응답 반환 (빠른 응답)
    if DEMO_MODE:
        print("[DEBUG] DEMO_MODE is True, calling _vlm_fallback_response")
        response = _vlm_fallback_response(message, history, image, image_name)
        print(f"[DEBUG] fallback response: {response[:100] if response else None}...")
        return response

    # VLM 모델 초기화 (lazy loading)
    if vlm_model is None:
        if not VLM_AVAILABLE:
            return _vlm_fallback_response(message, history, image, image_name)

        try:
            model, status = init_vlm_model(quantize=True)
            if model is None:
                return f"VLM 모델 로드 실패: {status}\n\n데모 모드로 전환합니다.\n\n" + _vlm_fallback_response(message, history, image, image_name)
        except Exception as e:
            return f"VLM 모델 초기화 오류: {str(e)}\n\n데모 모드로 전환합니다.\n\n" + _vlm_fallback_response(message, history, image, image_name)

    # 대화 기록을 컨텍스트로 구성
    context = ""
    if history:
        for h in history[-3:]:  # 최근 3개 대화만 사용
            context += f"사용자: {h[0]}\n어시스턴트: {h[1]}\n\n"

    # 프롬프트 구성
    if context:
        prompt = f"이전 대화:\n{context}\n현재 질문: {message}\n\n위 이미지와 이전 대화 맥락을 고려하여 질문에 답변하세요."
    else:
        prompt = f"질문: {message}\n\n위 디스플레이 패널 이미지를 분석하여 질문에 답변하세요."

    try:
        # 실제 VLM 모델로 응답 생성
        response = vlm_model.generate(
            image=image,
            prompt=prompt,
            system_prompt=VLM_SYSTEM_PROMPT,
            max_new_tokens=1024,
            temperature=0.1,
            do_sample=False,
        )
        return response
    except Exception as e:
        return f"VLM 추론 오류: {str(e)}\n\n데모 모드로 전환합니다.\n\n" + _vlm_fallback_response(message, history, image, image_name)


def _vlm_fallback_response(message, history, image, image_name=None):
    """VLM 모델 사용 불가시 폴백 응답 (데모용) - 메타데이터 활용"""
    import random
    global current_analysis_context

    message_lower = message.lower()

    # 이미지 이름이 있으면 메타데이터에서 정보 가져오기
    defect_info = {}
    if image_name and image_name in DEFECT_METADATA:
        defect_info = DEFECT_METADATA[image_name]
        current_analysis_context["image_name"] = image_name
        current_analysis_context["defect_type"] = defect_info.get("defect_type", "")
        current_analysis_context["coordinates"] = defect_info.get("coordinates", {})

    # 좌표 쿼리 (x,y, 좌표 등)
    if any(word in message_lower for word in ["x,y", "x, y", "좌표", "coordinate", "bbox", "bounding"]):
        coords = current_analysis_context.get("coordinates") or defect_info.get("coordinates", {})
        if coords:
            x, y = coords.get("x", 128), coords.get("y", 128)
            w, h = coords.get("width", 20), coords.get("height", 20)
            return f"""**결함 좌표 정보:**

**Bounding Box:**
- 시작점 (x1, y1): ({x}, {y})
- 끝점 (x2, y2): ({x + w}, {y + h})
- 중심점: ({x + w//2}, {y + h//2})

**크기:**
- 너비: {w}px ({w * 0.1:.1f}mm)
- 높이: {h}px ({h * 0.1:.1f}mm)
- 면적: {w * h}px² ({w * h * 0.01:.2f}mm²)

**패널 위치:** {defect_info.get("location", "중앙")} 영역"""
        else:
            return """**결함 좌표 정보:**

**Bounding Box:**
- 시작점 (x1, y1): (128, 128)
- 끝점 (x2, y2): (156, 148)
- 중심점: (142, 138)

**크기:**
- 너비: 28px (2.8mm)
- 높이: 20px (2.0mm)
- 면적: 560px² (5.6mm²)

**패널 위치:** 중앙 영역"""

    # 시각화 요청
    if any(word in message_lower for word in ["시각화", "표시", "그려", "보여", "visualize", "draw", "show"]):
        coords = current_analysis_context.get("coordinates") or defect_info.get("coordinates", {})
        if coords:
            x, y = coords.get("x", 128), coords.get("y", 128)
            w, h = coords.get("width", 20), coords.get("height", 20)
            return f"""**결함 좌표 시각화 안내**

왼쪽 패널의 **'결함 좌표 시각화'** 버튼을 클릭하면 이미지에 결함 위치가 표시됩니다.

**현재 결함 정보:**
- 결함 유형: {defect_info.get("defect_type", "N/A")}
- 위치: {defect_info.get("location", "N/A")}
- 좌표: ({x}, {y}) - ({x + w}, {y + h})
- 심각도: {defect_info.get("severity", "N/A")}

시각화 이미지에는 다음 정보가 표시됩니다:
- **빨간색 박스**: 결함 영역 (Bounding Box)
- **녹색 십자**: 결함 중심점
- **라벨**: 결함 유형 및 심각도
- **좌표 텍스트**: 시작점-끝점 좌표"""
        else:
            return """**결함 좌표 시각화 안내**

왼쪽 패널의 **'결함 좌표 시각화'** 버튼을 클릭하면 이미지에 결함 위치가 표시됩니다.

시각화 이미지에는 다음 정보가 표시됩니다:
- **빨간색 박스**: 결함 영역 (Bounding Box)
- **녹색 십자**: 결함 중심점
- **라벨**: 결함 유형 및 심각도
- **좌표 텍스트**: 시작점-끝점 좌표"""

    # 유사 이미지 검색
    if any(word in message_lower for word in ["유사", "비슷", "similar", "같은", "동일", "찾아"]):
        defect_type = current_analysis_context.get("defect_type") or defect_info.get("defect_type", "")

        if defect_type:
            similar_images = find_similar_images(defect_type, image_name, max_results=5)
            if similar_images:
                result = f"**유사 결함 이미지 검색 결과 ({len(similar_images)}건):**\n\n"
                result += f"검색 조건: **{defect_type}** 유형\n\n"

                for i, sim in enumerate(similar_images, 1):
                    result += f"**{i}. {sim['image']}** (유사도: {sim['similarity']*100:.1f}%)\n"
                    result += f"   - 결함: {sim['defect_type']}\n"
                    result += f"   - 위치: {sim['location']}\n"
                    result += f"   - 심각도: {sim['severity']}\n"
                    result += f"   - 원인: {sim['cause']}\n\n"

                return result
            else:
                return f"**유사 결함 이미지 검색 결과:**\n\n'{defect_type}' 유형의 유사 이미지를 찾지 못했습니다."
        else:
            # 기본 유사 이미지 검색 (라인 결함 기준)
            similar_images = find_similar_images("라인 결함", None, max_results=3)
            result = "**유사 결함 이미지 검색 결과 (3건):**\n\n"
            for i, sim in enumerate(similar_images, 1):
                result += f"**{i}. {sim['image']}** (유사도: {sim['similarity']*100:.1f}%)\n"
                result += f"   - 결함: {sim['defect_type']}, 위치: {sim['location']}\n\n"
            return result

    # 질문 유형별 응답 생성
    if any(word in message_lower for word in ["결함", "defect", "문제", "이상", "뭐가", "무엇", "분석"]):
        if defect_info:
            return f"""**이미지 분석 결과:**

**결함 유형:** {defect_info.get("defect_type", "라인 결함")}
**위치:** {defect_info.get("location", "중앙")}
**심각도:** {defect_info.get("severity", "medium").upper()}
**추정 원인:** {defect_info.get("cause", "공정 불량")}
**권장 조치:** {defect_info.get("action", "설비 점검")}

이미지에서 {defect_info.get("location", "중앙")} 영역에 {defect_info.get("defect_type", "결함")} 패턴이 관찰됩니다."""
        else:
            responses = [
                "이미지를 분석한 결과, **라인 결함(Line Defect)**이 관찰됩니다. 화면 중앙부에서 수직 방향으로 약 2mm 길이의 밝은 선이 보입니다.",
                "분석 결과 **무라(Mura) 현상**이 검출되었습니다. 좌측 하단 영역에서 불균일한 밝기 분포가 나타나고 있습니다.",
                "**데드 픽셀(Dead Pixel)**이 발견되었습니다. 우측 상단 좌표 근처에 약 0.3mm 크기의 검은 점이 관찰됩니다.",
            ]
            return random.choice(responses)

    elif any(word in message_lower for word in ["원인", "왜", "이유", "cause", "why"]):
        cause = defect_info.get("cause", "") if defect_info else ""
        if cause:
            return f"""**결함 원인 분석:**

**직접 원인:** {cause}

**추가 분석:**
1. **공정 요인**: 해당 공정의 파라미터 편차 가능성
2. **설비 요인**: 관련 설비의 PM 이력 확인 필요
3. **재료 요인**: 입고 재료 LOT 품질 확인

**권장 조치:** {defect_info.get("action", "설비 점검 및 공정 조건 최적화")}"""
        else:
            return """**결함 원인 분석:**

**주요 원인:**
1. **공정 요인**: CVD 증착 공정에서의 온도 불균일
2. **설비 요인**: 에칭 장비의 RF 파워 불안정
3. **재료 요인**: 타겟 재료의 순도 저하

**권장 조치:** 설비 PM 및 공정 조건 재검토"""

    elif any(word in message_lower for word in ["위치", "어디", "location", "where"]) and "좌표" not in message_lower:
        location = defect_info.get("location", "중앙") if defect_info else "중앙"
        return f"""**결함 위치 분석:**

**영역:** {location}
**세부 위치:** Active Area 내부
**영향 범위:** 약 5mm x 3mm

해당 위치는 TFT 어레이의 데이터 라인 영역에 해당합니다."""

    elif any(word in message_lower for word in ["심각", "등급", "레벨", "severity", "critical", "major"]):
        severity = defect_info.get("severity", "medium") if defect_info else "medium"
        severity_upper = severity.upper()
        return f"""**심각도 평가 결과: {severity_upper}**

**판정 기준:**
- 결함 크기: 기준 대비 평가
- 발생 위치: Active Area 내부
- 시인성: 일반 조건에서 확인 가능 여부

**종합 판정:** {severity_upper} Defect
**조치 권고:** {defect_info.get("action", "설비 점검") if defect_info else "상세 분석 필요"}"""

    elif any(word in message_lower for word in ["조치", "해결", "대책", "action", "solution", "어떻게"]):
        action = defect_info.get("action", "") if defect_info else ""
        return f"""**권장 조치사항:**

**즉시 조치:**
1. {action if action else "설비 파라미터 점검"}
2. 해당 LOT 격리 및 샘플링 검사
3. 동일 시간대 생산품 추적

**재발 방지:**
1. PM 주기 단축 검토
2. 공정 모니터링 강화
3. 작업자 교육 실시"""

    elif any(word in message_lower for word in ["크기", "사이즈", "size", "넓이", "면적"]):
        coords = defect_info.get("coordinates", {}) if defect_info else {}
        w = coords.get("width", 28)
        h = coords.get("height", 20)
        return f"""**결함 크기 측정:**

**실제 크기:**
- 가로: {w * 0.1:.1f}mm
- 세로: {h * 0.1:.1f}mm
- 면적: {w * h * 0.01:.2f}mm²

**기준 대비:**
- 가로: 기준(0.3mm) 대비 {(w * 0.1 / 0.3) * 100:.0f}%
- 세로: 기준(0.5mm) 대비 {(h * 0.1 / 0.5) * 100:.0f}%"""

    elif any(word in message_lower for word in ["수량", "개수", "몇 개", "count", "how many"]):
        return """**결함 검출 개수:**

- Critical: 0개
- Major: 1개
- Minor: 2개
- Cosmetic: 1개

**총 결함 수: 4개**

판정: Minor 이상 결함 존재"""

    else:
        return f"""질문: "{message}"

이미지 분석이 완료되었습니다. 다음과 같은 질문을 해보세요:

- "어떤 결함이 있나요?"
- "결함의 원인은 뭔가요?"
- "**x,y 좌표 알려줘**" (결함 좌표 정보)
- "**유사 이미지 찾아줘**" (비슷한 결함 이미지 검색)
- "심각도는 어느 정도인가요?"
- "어떤 조치가 필요한가요?\""""


def vlm_chat(message, history, image, image_name=None):
    """VLM 채팅 핸들러"""
    global current_analysis_context

    print(f"[DEBUG] vlm_chat called: message={message}, image_name={image_name}, image_type={type(image)}")

    if not message.strip():
        print("[DEBUG] Empty message, returning")
        return history, ""

    # 이미지 이름이 있으면 컨텍스트에 저장
    if image_name:
        current_analysis_context["image_name"] = image_name
        if image_name in DEFECT_METADATA:
            current_analysis_context["defect_type"] = DEFECT_METADATA[image_name].get("defect_type", "")
            current_analysis_context["coordinates"] = DEFECT_METADATA[image_name].get("coordinates", {})

    # 응답 생성
    print("[DEBUG] Calling vlm_chat_response...")
    response = vlm_chat_response(message, history, image)
    print(f"[DEBUG] Got response: {response[:50] if response else None}...")

    # 히스토리에 추가 (Gradio 6.x 메시지 형식)
    history = history + [
        {"role": "user", "content": message},
        {"role": "assistant", "content": response}
    ]

    return history, ""


def clear_chat():
    """채팅 초기화"""
    global current_analysis_context
    current_analysis_context = {"image_name": None, "defect_type": None, "coordinates": None}
    return [], None


def select_sample_image(evt: gr.SelectData):
    """샘플 이미지 선택 핸들러"""
    global current_analysis_context

    if evt.index is not None and evt.index < len(SAMPLE_IMAGES):
        img_path = SAMPLE_IMAGES[evt.index]
        image_name = img_path.name

        # 컨텍스트 업데이트
        current_analysis_context["image_name"] = image_name
        if image_name in DEFECT_METADATA:
            info = DEFECT_METADATA[image_name]
            current_analysis_context["defect_type"] = info.get("defect_type", "")
            current_analysis_context["coordinates"] = info.get("coordinates", {})

        # 이미지 로드 및 반환
        img = Image.open(img_path)
        return img, image_name

    return None, None


def get_image_info(image_name):
    """이미지 정보 표시"""
    if not image_name:
        return ""

    if image_name in DEFECT_METADATA:
        info = DEFECT_METADATA[image_name]
        return f"""**선택된 이미지:** {image_name}
- 결함: {info.get("defect_type", "N/A")}
- 위치: {info.get("location", "N/A")}
- 심각도: {info.get("severity", "N/A")}"""
    return f"**선택된 이미지:** {image_name}"


# ==================== GraphRAG 관리 탭 ====================

def add_defect_node(defect_id, defect_type, korean_name, description, severity, visual_char):
    """결함 노드 추가"""
    global kb
    if kb is None:
        return "지식 베이스를 먼저 초기화해주세요."

    try:
        severity_list = [SeverityLevel(s.strip()) for s in severity.split(",") if s.strip()]
        defect = Defect(
            defect_id=defect_id,
            defect_type=DefectType(defect_type) if defect_type in [e.value for e in DefectType] else DefectType.CUSTOM,
            korean_name=korean_name,
            description=description,
            severity_levels=severity_list,
            visual_characteristics=visual_char
        )
        kb.store.create_defect(defect)
        return f"결함 노드 '{korean_name}' 추가 완료!"
    except Exception as e:
        return f"오류: {str(e)}"


def add_cause_node(cause_id, cause_type, korean_name, description, category):
    """원인 노드 추가"""
    global kb
    if kb is None:
        return "지식 베이스를 먼저 초기화해주세요."

    try:
        cause = RootCause(
            cause_id=cause_id,
            cause_type=cause_type,
            korean_name=korean_name,
            description=description,
            category=CauseCategory(category)
        )
        kb.store.create_root_cause(cause)
        return f"원인 노드 '{korean_name}' 추가 완료!"
    except Exception as e:
        return f"오류: {str(e)}"


def add_action_node(action_id, action_type, korean_name, description, priority):
    """조치 노드 추가"""
    global kb
    if kb is None:
        return "지식 베이스를 먼저 초기화해주세요."

    try:
        action = Action(
            action_id=action_id,
            action_type=action_type,
            korean_name=korean_name,
            description=description,
            priority=ActionPriority(priority)
        )
        kb.store.create_action(action)
        return f"조치 노드 '{korean_name}' 추가 완료!"
    except Exception as e:
        return f"오류: {str(e)}"


def add_caused_by_relation(defect_id, cause_id, probability, evidence):
    """CAUSED_BY 관계 추가"""
    global kb
    if kb is None:
        return "지식 베이스를 먼저 초기화해주세요."

    try:
        relation = CausedByRelation(
            defect_id=defect_id,
            cause_id=cause_id,
            probability=float(probability),
            evidence=evidence
        )
        kb.store.create_caused_by(relation)
        return f"관계 추가 완료: {defect_id} --[CAUSED_BY]--> {cause_id}"
    except Exception as e:
        return f"오류: {str(e)}"


def add_requires_relation(cause_id, action_id, effectiveness):
    """REQUIRES 관계 추가"""
    global kb
    if kb is None:
        return "지식 베이스를 먼저 초기화해주세요."

    try:
        relation = RequiresRelation(
            cause_id=cause_id,
            action_id=action_id,
            effectiveness=float(effectiveness)
        )
        kb.store.create_requires(relation)
        return f"관계 추가 완료: {cause_id} --[REQUIRES]--> {action_id}"
    except Exception as e:
        return f"오류: {str(e)}"


def get_all_nodes():
    """모든 노드 목록 조회"""
    global kb
    if kb is None:
        return "지식 베이스를 먼저 초기화해주세요."

    try:
        result = []

        if hasattr(kb.store, 'defects'):
            result.append("=== 결함 노드 ===")
            for d in kb.store.defects.values():
                result.append(f"  [{d.defect_id}] {d.korean_name} ({d.defect_type.value})")

        if hasattr(kb.store, 'root_causes'):
            result.append("\n=== 원인 노드 ===")
            for c in kb.store.root_causes.values():
                result.append(f"  [{c.cause_id}] {c.korean_name} ({c.category.value})")

        if hasattr(kb.store, 'actions'):
            result.append("\n=== 조치 노드 ===")
            for a in kb.store.actions.values():
                result.append(f"  [{a.action_id}] {a.korean_name} ({a.priority.value})")

        if hasattr(kb.store, 'processes'):
            result.append("\n=== 공정 노드 ===")
            for p in kb.store.processes.values():
                result.append(f"  [{p.process_id}] {p.korean_name}")

        return "\n".join(result) if result else "노드가 없습니다."
    except Exception as e:
        return f"오류: {str(e)}"


def get_all_relations():
    """모든 관계 목록 조회"""
    global kb
    if kb is None:
        return "지식 베이스를 먼저 초기화해주세요."

    try:
        result = []

        if hasattr(kb.store, 'caused_by'):
            result.append("=== CAUSED_BY 관계 ===")
            for r in kb.store.caused_by:
                defect = kb.store.defects.get(r.defect_id)
                cause = kb.store.root_causes.get(r.cause_id)
                d_name = defect.korean_name if defect else r.defect_id
                c_name = cause.korean_name if cause else r.cause_id
                result.append(f"  {d_name} --[{r.probability:.0%}]--> {c_name}")

        if hasattr(kb.store, 'requires'):
            result.append("\n=== REQUIRES 관계 ===")
            for r in kb.store.requires:
                cause = kb.store.root_causes.get(r.cause_id)
                action = kb.store.actions.get(r.action_id)
                c_name = cause.korean_name if cause else r.cause_id
                a_name = action.korean_name if action else r.action_id
                result.append(f"  {c_name} --[{r.effectiveness:.0%}]--> {a_name}")

        return "\n".join(result) if result else "관계가 없습니다."
    except Exception as e:
        return f"오류: {str(e)}"


def query_defect_analysis(defect_type):
    """결함 분석 쿼리"""
    global kb
    if kb is None:
        return "지식 베이스를 먼저 초기화해주세요."

    try:
        analysis = kb.analyze_defect(defect_type)

        result = [f"=== '{defect_type}' 분석 결과 ===\n"]

        result.append("【추정 원인】")
        for i, cause in enumerate(analysis.get("root_causes", []), 1):
            result.append(f"  {i}. {cause.get('cause')} ({cause.get('probability', 0)*100:.0f}%)")
            if cause.get('evidence'):
                result.append(f"     근거: {cause.get('evidence')}")

        result.append("\n【권장 조치】")
        for i, action in enumerate(analysis.get("recommended_actions", []), 1):
            result.append(f"  {i}. {action.get('action')} (효과: {action.get('effectiveness', 0)*100:.0f}%)")
            result.append(f"     대상: {action.get('for_cause')}")

        return "\n".join(result)
    except Exception as e:
        return f"오류: {str(e)}"


# ==================== 엑셀 업로드 기능 ====================

def create_excel_template():
    """엑셀 템플릿 생성"""
    if not PANDAS_AVAILABLE:
        return None, "pandas가 설치되지 않았습니다. pip install pandas openpyxl"

    try:
        # 결함 시트
        defects_df = pd.DataFrame({
            'defect_id': ['DEF001', 'DEF002'],
            'defect_type': ['dead_pixel', 'bright_spot'],
            'korean_name': ['데드 픽셀', '휘점 결함'],
            'description': ['화면에 검은 점으로 나타나는 비활성 픽셀', '화면에 밝은 점으로 나타나는 결함'],
            'severity_levels': ['low,medium,high', 'low,medium,high'],
            'visual_characteristics': ['검은색 또는 어두운 점', '밝은 점, 흰색 또는 색상 점']
        })

        # 원인 시트
        causes_df = pd.DataFrame({
            'cause_id': ['RC001', 'RC002'],
            'cause_type': ['tft_manufacturing_defect', 'contamination'],
            'korean_name': ['TFT 제조 결함', '오염'],
            'description': ['TFT 어레이 제조 공정 중 발생한 결함', '클린룸 환경 오염'],
            'category': ['process', 'environment']
        })

        # 조치 시트
        actions_df = pd.DataFrame({
            'action_id': ['ACT001', 'ACT002'],
            'action_type': ['equipment_inspection', 'process_adjustment'],
            'korean_name': ['장비 점검', '공정 조건 조정'],
            'description': ['관련 장비의 상태 점검 및 유지보수', '공정 파라미터 재설정'],
            'priority': ['high', 'medium']
        })

        # CAUSED_BY 관계 시트
        caused_by_df = pd.DataFrame({
            'defect_id': ['DEF001', 'DEF002'],
            'cause_id': ['RC001', 'RC002'],
            'probability': [0.7, 0.6],
            'evidence': ['TFT 트랜지스터 제조 결함으로 인한 픽셀 비활성화', '클린룸 오염으로 인한 결함']
        })

        # REQUIRES 관계 시트
        requires_df = pd.DataFrame({
            'cause_id': ['RC001', 'RC002'],
            'action_id': ['ACT001', 'ACT002'],
            'effectiveness': [0.8, 0.9]
        })

        # 엑셀 파일 생성
        output_path = Path('/tmp/graphrag_template.xlsx')
        with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
            defects_df.to_excel(writer, sheet_name='Defects', index=False)
            causes_df.to_excel(writer, sheet_name='RootCauses', index=False)
            actions_df.to_excel(writer, sheet_name='Actions', index=False)
            caused_by_df.to_excel(writer, sheet_name='CAUSED_BY', index=False)
            requires_df.to_excel(writer, sheet_name='REQUIRES', index=False)

        return str(output_path), "템플릿 생성 완료! 다운로드 후 데이터를 입력하세요."

    except Exception as e:
        return None, f"템플릿 생성 오류: {str(e)}"


def upload_excel_data(file):
    """엑셀 파일 업로드 및 데이터 추가"""
    global kb

    if not PANDAS_AVAILABLE:
        return "pandas가 설치되지 않았습니다. pip install pandas openpyxl"

    if kb is None:
        return "지식 베이스를 먼저 초기화해주세요."

    if file is None:
        return "파일을 선택해주세요."

    try:
        results = []

        # 엑셀 파일 읽기
        xlsx = pd.ExcelFile(file.name)
        sheet_names = xlsx.sheet_names

        # Defects 시트 처리
        if 'Defects' in sheet_names:
            df = pd.read_excel(xlsx, sheet_name='Defects')
            count = 0
            for _, row in df.iterrows():
                try:
                    severity_list = []
                    if pd.notna(row.get('severity_levels')):
                        severity_list = [SeverityLevel(s.strip()) for s in str(row['severity_levels']).split(',') if s.strip()]

                    defect = Defect(
                        defect_id=str(row['defect_id']),
                        defect_type=DefectType(row['defect_type']) if row['defect_type'] in [e.value for e in DefectType] else DefectType.CUSTOM,
                        korean_name=str(row['korean_name']),
                        description=str(row.get('description', '')),
                        severity_levels=severity_list,
                        visual_characteristics=str(row.get('visual_characteristics', ''))
                    )
                    kb.store.create_defect(defect)
                    count += 1
                except Exception as e:
                    results.append(f"  - 결함 '{row.get('defect_id')}' 오류: {e}")
            results.append(f"결함 노드: {count}개 추가")

        # RootCauses 시트 처리
        if 'RootCauses' in sheet_names:
            df = pd.read_excel(xlsx, sheet_name='RootCauses')
            count = 0
            for _, row in df.iterrows():
                try:
                    cause = RootCause(
                        cause_id=str(row['cause_id']),
                        cause_type=str(row['cause_type']),
                        korean_name=str(row['korean_name']),
                        description=str(row.get('description', '')),
                        category=CauseCategory(row.get('category', 'process'))
                    )
                    kb.store.create_root_cause(cause)
                    count += 1
                except Exception as e:
                    results.append(f"  - 원인 '{row.get('cause_id')}' 오류: {e}")
            results.append(f"원인 노드: {count}개 추가")

        # Actions 시트 처리
        if 'Actions' in sheet_names:
            df = pd.read_excel(xlsx, sheet_name='Actions')
            count = 0
            for _, row in df.iterrows():
                try:
                    action = Action(
                        action_id=str(row['action_id']),
                        action_type=str(row['action_type']),
                        korean_name=str(row['korean_name']),
                        description=str(row.get('description', '')),
                        priority=ActionPriority(row.get('priority', 'medium'))
                    )
                    kb.store.create_action(action)
                    count += 1
                except Exception as e:
                    results.append(f"  - 조치 '{row.get('action_id')}' 오류: {e}")
            results.append(f"조치 노드: {count}개 추가")

        # CAUSED_BY 관계 시트 처리
        if 'CAUSED_BY' in sheet_names:
            df = pd.read_excel(xlsx, sheet_name='CAUSED_BY')
            count = 0
            for _, row in df.iterrows():
                try:
                    relation = CausedByRelation(
                        defect_id=str(row['defect_id']),
                        cause_id=str(row['cause_id']),
                        probability=float(row.get('probability', 0.5)),
                        evidence=str(row.get('evidence', ''))
                    )
                    kb.store.create_caused_by(relation)
                    count += 1
                except Exception as e:
                    results.append(f"  - CAUSED_BY 관계 오류: {e}")
            results.append(f"CAUSED_BY 관계: {count}개 추가")

        # REQUIRES 관계 시트 처리
        if 'REQUIRES' in sheet_names:
            df = pd.read_excel(xlsx, sheet_name='REQUIRES')
            count = 0
            for _, row in df.iterrows():
                try:
                    relation = RequiresRelation(
                        cause_id=str(row['cause_id']),
                        action_id=str(row['action_id']),
                        effectiveness=float(row.get('effectiveness', 0.5))
                    )
                    kb.store.create_requires(relation)
                    count += 1
                except Exception as e:
                    results.append(f"  - REQUIRES 관계 오류: {e}")
            results.append(f"REQUIRES 관계: {count}개 추가")

        return "=== 엑셀 업로드 완료 ===\n" + "\n".join(results)

    except Exception as e:
        return f"엑셀 처리 오류: {str(e)}"


# ==================== 빅데이터 분석 탭 ====================

# 빅데이터 모듈 임포트
try:
    from src.bigdata.datalake.extractor import DefectDataExtractor, DefectCase, ExtractedData
    from src.bigdata.pipeline.parquet_converter import ParquetConverter
    from src.bigdata.pipeline.data_pipeline import DataPipeline, PipelineStatus
    from src.bigdata.spark.processor import SparkProcessor
    BIGDATA_AVAILABLE = True
except ImportError:
    BIGDATA_AVAILABLE = False

# 목업 데이터 생성기 임포트
try:
    from src.bigdata.mockdata import MockDataGenerator
    MOCKDATA_AVAILABLE = True
except ImportError:
    MOCKDATA_AVAILABLE = False

# 전역 목업 데이터 생성기
mock_generator = None
mock_data_result = None

# 전역 파이프라인 인스턴스
pipeline = None
last_pipeline_result = None


def init_bigdata_pipeline():
    """빅데이터 파이프라인 초기화"""
    global pipeline
    try:
        extractor = DefectDataExtractor()
        converter = ParquetConverter()
        pipeline = DataPipeline(extractor=extractor, converter=converter)
        return "빅데이터 파이프라인 초기화 완료"
    except Exception as e:
        return f"초기화 실패: {str(e)}"


def run_defect_analysis_pipeline(case_id, cell_id, defect_type, defect_date, customer, severity, description):
    """불량 분석 파이프라인 실행"""
    global pipeline, last_pipeline_result

    if not BIGDATA_AVAILABLE:
        return create_pipeline_result_html({
            "status": "error",
            "message": "빅데이터 모듈을 로드할 수 없습니다."
        }), ""

    if not pipeline:
        init_bigdata_pipeline()

    try:
        from datetime import datetime

        # DefectCase 생성
        defect_case = DefectCase(
            case_id=case_id or f"CASE{datetime.now().strftime('%Y%m%d%H%M%S')}",
            cell_id=cell_id or "CELL001",
            defect_type=defect_type or "unknown",
            defect_date=datetime.strptime(defect_date, "%Y-%m-%d") if defect_date else datetime.now(),
            customer=customer or "Unknown",
            severity=severity or "MEDIUM",
            description=description or ""
        )

        # 파이프라인 실행
        result = pipeline.run(defect_case)
        last_pipeline_result = result

        # 결과 요약
        summary = pipeline.get_pipeline_summary(result)

        return create_pipeline_result_html(summary), json.dumps(summary, indent=2, ensure_ascii=False, default=str)

    except Exception as e:
        return create_pipeline_result_html({
            "status": "error",
            "message": str(e)
        }), str(e)


def create_pipeline_result_html(summary: dict) -> str:
    """파이프라인 결과 HTML 생성"""
    status = summary.get("status", "unknown")
    status_color = "#27ae60" if status == "completed" else ("#e74c3c" if status in ["failed", "error"] else "#f39c12")
    status_icon = "✅" if status == "completed" else ("❌" if status in ["failed", "error"] else "⏳")

    steps_html = ""
    for step in summary.get("steps", []):
        step_status = step.get("status", "pending")
        step_icon = "✅" if step_status == "completed" else ("❌" if step_status == "failed" else "⏳")
        step_color = "#27ae60" if step_status == "completed" else ("#e74c3c" if step_status == "failed" else "#95a5a6")
        duration = f"{step.get('duration', 0):.2f}초" if step.get('duration') else "-"
        steps_html += f"""
        <div style="display: flex; align-items: center; padding: 8px; margin: 4px 0; background: #f8f9fa; border-radius: 6px; border-left: 3px solid {step_color};">
            <span style="margin-right: 10px;">{step_icon}</span>
            <span style="flex: 1; color: #333;">{step.get('name', '')}</span>
            <span style="color: #666; font-size: 12px;">{duration}</span>
        </div>
        """

    files_html = ""
    for name, path in summary.get("parquet_files", {}).items():
        files_html += f"""
        <div style="padding: 6px 10px; margin: 2px 0; background: #e8f4fc; border-radius: 4px; font-size: 12px; color: #333;">
            📁 {name}: {path}
        </div>
        """

    return f"""
    <div style="font-family: 'Segoe UI', sans-serif; padding: 20px; background: #ffffff; border-radius: 12px; border: 1px solid #e0e0e0;">
        <div style="display: flex; align-items: center; margin-bottom: 15px; padding-bottom: 15px; border-bottom: 1px solid #eee;">
            <span style="font-size: 24px; margin-right: 10px;">{status_icon}</span>
            <div>
                <div style="font-size: 18px; font-weight: bold; color: #333;">케이스: {summary.get('case_id', 'N/A')}</div>
                <div style="font-size: 13px; color: {status_color}; font-weight: 500;">{status.upper()}</div>
            </div>
            <div style="margin-left: auto; text-align: right;">
                <div style="font-size: 12px; color: #666;">처리 시간</div>
                <div style="font-size: 16px; font-weight: bold; color: #333;">{summary.get('duration_seconds', 0):.2f}초</div>
            </div>
        </div>

        <div style="margin-bottom: 15px;">
            <div style="font-size: 14px; font-weight: bold; color: #333; margin-bottom: 8px;">처리 단계</div>
            {steps_html}
        </div>

        <div style="margin-bottom: 15px;">
            <div style="font-size: 14px; font-weight: bold; color: #333; margin-bottom: 8px;">생성된 Parquet 파일</div>
            {files_html if files_html else '<div style="color: #999; font-size: 12px;">파일 없음</div>'}
        </div>

        {f'<div style="padding: 10px; background: #d4efdf; border-radius: 6px; color: #196f3d;"><strong>데이터마트:</strong> {summary.get("datamart_path", "N/A")}</div>' if summary.get("datamart_path") else ''}
        {f'<div style="padding: 10px; background: #fadbd8; border-radius: 6px; color: #943126; margin-top: 10px;"><strong>오류:</strong> {summary.get("error", "")}</div>' if summary.get("error") else ''}
    </div>
    """


def get_extracted_data_summary():
    """추출된 데이터 요약"""
    global last_pipeline_result
    if not last_pipeline_result:
        return "파이프라인을 먼저 실행해주세요."

    summary = []
    for step in last_pipeline_result.steps:
        if step.name == "데이터 추출" and step.result:
            data = step.result
            summary.append(f"=== 추출 데이터 요약 ===")
            summary.append(f"케이스 ID: {data.case_id}")
            summary.append(f"셀 ID: {data.cell_id}")
            summary.append(f"추출 시간: {data.extraction_time}")
            summary.append(f"\n제품 이력: {len(data.product_history)}건")
            summary.append(f"개발 이력: {len(data.dev_history)}건")
            summary.append(f"변경점: {len(data.change_points)}건")
            summary.append(f"설비 마스터: {len(data.equipment_master)}건")
            summary.append(f"유지보수 이력: {len(data.maintenance_history)}건")
            summary.append(f"FDC 파라미터: {len(data.fdc_parameters)}건")
            break

    return "\n".join(summary) if summary else "추출된 데이터가 없습니다."


# ==================== 목업 데이터 생성 탭 ====================

def init_mock_generator(output_dir: str = "/tmp/mockdata"):
    """목업 데이터 생성기 초기화"""
    global mock_generator
    if not MOCKDATA_AVAILABLE:
        return "목업 데이터 모듈을 로드할 수 없습니다."
    try:
        mock_generator = MockDataGenerator(output_dir=output_dir)
        return f"목업 데이터 생성기 초기화 완료\n출력 디렉토리: {output_dir}"
    except Exception as e:
        return f"초기화 실패: {str(e)}"


def generate_mock_data(num_lots: int, num_cells: int, num_days: int):
    """목업 데이터 생성"""
    global mock_generator, mock_data_result

    if not MOCKDATA_AVAILABLE:
        return create_mock_result_html({"status": "error", "message": "목업 데이터 모듈을 로드할 수 없습니다."}), ""

    if not mock_generator:
        init_mock_generator()

    try:
        # 데이터 생성
        result_paths = mock_generator.generate_all(
            num_lots=int(num_lots),
            num_cells_per_lot=int(num_cells),
            num_days=int(num_days)
        )
        mock_data_result = result_paths

        # 요약 정보 생성
        summary = mock_generator.get_summary()

        result = {
            "status": "completed",
            "num_lots": num_lots,
            "num_cells": num_cells,
            "num_days": num_days,
            "categories": {},
            "total_records": 0,
            "total_files": 0
        }

        for category, info in summary.items():
            result["categories"][category] = {
                "files": len(info.get("files", [])),
                "records": info.get("total_records", 0),
                "directory": info.get("directory", "")
            }
            result["total_records"] += info.get("total_records", 0)
            result["total_files"] += len(info.get("files", []))

        return create_mock_result_html(result), json.dumps(result, indent=2, ensure_ascii=False, default=str)

    except Exception as e:
        return create_mock_result_html({"status": "error", "message": str(e)}), str(e)


def create_mock_result_html(result: dict) -> str:
    """목업 데이터 생성 결과 HTML"""
    status = result.get("status", "unknown")
    status_color = "#27ae60" if status == "completed" else "#e74c3c"
    status_icon = "✅" if status == "completed" else "❌"

    if status == "error":
        return f"""
        <div style="font-family: 'Segoe UI', sans-serif; padding: 20px; background: #ffffff; border-radius: 12px; border: 1px solid #e74c3c;">
            <div style="display: flex; align-items: center; color: #e74c3c;">
                <span style="font-size: 24px; margin-right: 10px;">❌</span>
                <span style="font-size: 16px; font-weight: bold;">오류 발생</span>
            </div>
            <div style="margin-top: 10px; padding: 10px; background: #fadbd8; border-radius: 6px; color: #943126;">
                {result.get('message', '알 수 없는 오류')}
            </div>
        </div>
        """

    categories_html = ""
    for cat, info in result.get("categories", {}).items():
        cat_icon = {
            "development": "🔬", "equipment": "⚙️", "material": "📦",
            "inspection": "🔍", "quality": "✅", "manufacturing": "🏭",
            "mes": "📊", "traceability": "🔗", "parquet": "📁"
        }.get(cat, "📄")

        categories_html += f"""
        <div style="display: flex; align-items: center; padding: 10px; margin: 4px 0; background: #f8f9fa; border-radius: 6px; border-left: 3px solid #4a90d9;">
            <span style="margin-right: 10px; font-size: 18px;">{cat_icon}</span>
            <div style="flex: 1;">
                <div style="font-weight: bold; color: #333;">{cat.upper()}</div>
                <div style="font-size: 12px; color: #666;">{info.get('directory', '')}</div>
            </div>
            <div style="text-align: right;">
                <div style="font-size: 14px; font-weight: bold; color: #4a90d9;">{info.get('records', 0):,}건</div>
                <div style="font-size: 11px; color: #888;">{info.get('files', 0)}파일</div>
            </div>
        </div>
        """

    return f"""
    <div style="font-family: 'Segoe UI', sans-serif; padding: 20px; background: #ffffff; border-radius: 12px; border: 1px solid #e0e0e0;">
        <div style="display: flex; align-items: center; margin-bottom: 15px; padding-bottom: 15px; border-bottom: 1px solid #eee;">
            <span style="font-size: 24px; margin-right: 10px;">{status_icon}</span>
            <div>
                <div style="font-size: 18px; font-weight: bold; color: #333;">목업 데이터 생성 완료</div>
                <div style="font-size: 13px; color: {status_color}; font-weight: 500;">{status.upper()}</div>
            </div>
            <div style="margin-left: auto; text-align: right;">
                <div style="font-size: 12px; color: #666;">총 레코드</div>
                <div style="font-size: 20px; font-weight: bold; color: #4a90d9;">{result.get('total_records', 0):,}건</div>
            </div>
        </div>

        <div style="display: grid; grid-template-columns: repeat(3, 1fr); gap: 10px; margin-bottom: 15px;">
            <div style="background: linear-gradient(135deg, #e8f4fc 0%, #f0f9ff 100%); padding: 15px; border-radius: 8px; text-align: center;">
                <div style="font-size: 11px; color: #666; text-transform: uppercase;">로트 수</div>
                <div style="font-size: 24px; font-weight: bold; color: #2980b9;">{result.get('num_lots', 0)}</div>
            </div>
            <div style="background: linear-gradient(135deg, #d4efdf 0%, #eafaf1 100%); padding: 15px; border-radius: 8px; text-align: center;">
                <div style="font-size: 11px; color: #666; text-transform: uppercase;">셀/로트</div>
                <div style="font-size: 24px; font-weight: bold; color: #27ae60;">{result.get('num_cells', 0)}</div>
            </div>
            <div style="background: linear-gradient(135deg, #fef5e7 0%, #fefbf3 100%); padding: 15px; border-radius: 8px; text-align: center;">
                <div style="font-size: 11px; color: #666; text-transform: uppercase;">기간 (일)</div>
                <div style="font-size: 24px; font-weight: bold; color: #f39c12;">{result.get('num_days', 0)}</div>
            </div>
        </div>

        <div style="margin-bottom: 10px;">
            <div style="font-size: 14px; font-weight: bold; color: #333; margin-bottom: 8px;">생성된 데이터 카테고리</div>
            {categories_html}
        </div>

        <div style="font-size: 12px; color: #888; text-align: center; margin-top: 15px;">
            총 {result.get('total_files', 0)}개 파일 생성됨
        </div>
    </div>
    """


def get_mock_data_preview(category: str):
    """목업 데이터 미리보기"""
    global mock_generator

    if not mock_generator:
        return "목업 데이터 생성기를 먼저 초기화해주세요."

    try:
        summary = mock_generator.get_summary()
        cat_info = summary.get(category, {})

        if not cat_info.get("files"):
            return f"{category} 카테고리에 파일이 없습니다."

        # 첫 번째 파일 읽기
        import os
        first_file = cat_info["files"][0]["filename"]
        filepath = os.path.join(cat_info["directory"], first_file)

        if filepath.endswith(".json"):
            with open(filepath, "r", encoding="utf-8") as f:
                data = json.load(f)

            # 처음 3개 레코드만 표시
            preview_data = data[:3] if len(data) > 3 else data
            result = [
                f"=== {category.upper()} 데이터 미리보기 ===",
                f"파일: {first_file}",
                f"총 레코드: {len(data)}건",
                f"\n처음 {len(preview_data)}개 레코드:",
                "-" * 50
            ]

            for i, record in enumerate(preview_data, 1):
                result.append(f"\n[{i}]")
                for key, value in list(record.items())[:8]:  # 처음 8개 필드만
                    if isinstance(value, (dict, list)):
                        value = json.dumps(value, ensure_ascii=False)[:50] + "..." if len(json.dumps(value)) > 50 else json.dumps(value, ensure_ascii=False)
                    result.append(f"  {key}: {value}")

            return "\n".join(result)
        else:
            return f"미리보기 지원 안 됨: {first_file}"

    except Exception as e:
        return f"미리보기 오류: {str(e)}"


# ==================== 품질 대시보드 차트 생성 ====================

def generate_quality_dashboard():
    """품질 분석 대시보드 차트 생성"""
    global mock_generator

    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
        import matplotlib.font_manager as fm
        import numpy as np
        from pathlib import Path
        import os

        # 한글 폰트 설정
        font_paths = [
            '/usr/share/fonts/truetype/nanum/NanumGothic.ttf',
            '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc',
        ]
        for fpath in font_paths:
            if os.path.exists(fpath):
                fm.fontManager.addfont(fpath)
                prop = fm.FontProperties(fname=fpath)
                plt.rcParams['font.family'] = prop.get_name()
                plt.rcParams['axes.unicode_minus'] = False
                break

        # 목업 데이터 로드
        mock_dir = "/tmp/mockdata_large"
        if not os.path.exists(mock_dir):
            mock_dir = "/tmp/mockdata"

        def load_json_data(category):
            cat_dir = os.path.join(mock_dir, category)
            if os.path.exists(cat_dir):
                files = [f for f in os.listdir(cat_dir) if f.endswith('.json')]
                if files:
                    with open(os.path.join(cat_dir, files[0]), 'r', encoding='utf-8') as f:
                        return json.load(f)
            return []

        quality_data = load_json_data('quality')
        manufacturing_data = load_json_data('manufacturing')
        mes_data = load_json_data('mes')

        if not quality_data:
            return None, "목업 데이터가 없습니다. '목업 데이터' 탭에서 먼저 데이터를 생성해주세요."

        # 데이터 분석
        fail_cases = [q for q in quality_data if q.get('inspection_result') == 'FAIL']
        defect_types = {}
        for case in fail_cases:
            dt = case.get('defect_type', 'UNKNOWN')
            defect_types[dt] = defect_types.get(dt, 0) + 1

        severity_counts = {}
        for case in fail_cases:
            sev = case.get('severity', 'UNKNOWN')
            severity_counts[sev] = severity_counts.get(sev, 0) + 1

        equipment_stats = {}
        for case in quality_data:
            eq = case.get('equipment_id', 'UNKNOWN')
            if eq not in equipment_stats:
                equipment_stats[eq] = {'total': 0, 'fail': 0}
            equipment_stats[eq]['total'] += 1
            if case.get('inspection_result') == 'FAIL':
                equipment_stats[eq]['fail'] += 1

        customer_stats = {}
        for mfg in manufacturing_data:
            cust = mfg.get('customer', 'UNKNOWN')
            if cust not in customer_stats:
                customer_stats[cust] = {'total': 0, 'pass': 0, 'yields': []}
            customer_stats[cust]['total'] += 1
            if mfg.get('final_result') == 'PASS':
                customer_stats[cust]['pass'] += 1
            customer_stats[cust]['yields'].append(mfg.get('yield_rate', 0))

        line_stats = {}
        for mes in mes_data:
            line = mes.get('line_id', 'UNKNOWN')
            if line not in line_stats:
                line_stats[line] = {'good': 0, 'ng': 0}
            line_stats[line]['good'] += mes.get('good_qty', 0)
            line_stats[line]['ng'] += mes.get('ng_qty', 0)

        # 색상 팔레트
        colors = ['#3498db', '#e74c3c', '#2ecc71', '#f39c12', '#9b59b6', '#1abc9c', '#e67e22', '#34495e']

        # 종합 대시보드 생성
        fig = plt.figure(figsize=(18, 12))

        # 1. 불량 유형 파이차트
        ax1 = fig.add_subplot(2, 3, 1)
        labels = list(defect_types.keys())
        sizes = list(defect_types.values())
        if labels and sizes:
            ax1.pie(sizes, labels=labels, autopct='%1.1f%%', colors=colors[:len(labels)], startangle=90)
        ax1.set_title('불량 유형별 분포', fontsize=13, fontweight='bold', pad=10)

        # 2. 심각도 막대
        ax2 = fig.add_subplot(2, 3, 2)
        severities = ['CRITICAL', 'MAJOR', 'MINOR', 'COSMETIC']
        counts = [severity_counts.get(s, 0) for s in severities]
        sev_colors = ['#e74c3c', '#f39c12', '#3498db', '#2ecc71']
        bars = ax2.barh(severities, counts, color=sev_colors, height=0.6)
        ax2.set_xlabel('건수', fontsize=11)
        ax2.set_title('심각도별 불량 분포', fontsize=13, fontweight='bold')
        for bar, count in zip(bars, counts):
            ax2.text(bar.get_width() + 10, bar.get_y() + bar.get_height()/2, f'{count:,}', va='center', fontsize=10)
        if counts:
            ax2.set_xlim(0, max(counts) * 1.15)

        # 3. 설비별 불량률 TOP 10
        ax3 = fig.add_subplot(2, 3, 3)
        top_equipment = sorted(
            [(eq, info['fail']/info['total']*100, info['total'], info['fail'])
             for eq, info in equipment_stats.items() if info['total'] >= 50],
            key=lambda x: -x[1]
        )[:10]
        if top_equipment:
            eq_names = [x[0] for x in top_equipment]
            eq_rates = [x[1] for x in top_equipment]
            bars = ax3.bar(eq_names, eq_rates, color='#e74c3c', alpha=0.8)
            avg_rate = sum(eq_rates) / len(eq_rates) if eq_rates else 0
            ax3.axhline(y=avg_rate, color='#3498db', linestyle='--', linewidth=2, label=f'평균: {avg_rate:.1f}%')
            ax3.set_ylabel('불량률 (%)', fontsize=11)
            ax3.set_title('설비별 불량률 TOP 10', fontsize=13, fontweight='bold')
            ax3.legend(loc='upper right', fontsize=9)
            plt.setp(ax3.xaxis.get_majorticklabels(), rotation=45, ha='right', fontsize=8)

        # 4. 고객사별 수율
        ax4 = fig.add_subplot(2, 3, 4)
        if customer_stats:
            customers = list(customer_stats.keys())
            avg_yields = [sum(customer_stats[c]['yields'])/len(customer_stats[c]['yields']) if customer_stats[c]['yields'] else 0 for c in customers]
            pass_rates = [customer_stats[c]['pass']/customer_stats[c]['total']*100 if customer_stats[c]['total'] > 0 else 0 for c in customers]
            x = np.arange(len(customers))
            width = 0.35
            ax4.bar(x - width/2, avg_yields, width, label='평균 수율', color='#3498db')
            ax4.bar(x + width/2, pass_rates, width, label='통과율', color='#2ecc71')
            ax4.set_xticks(x)
            ax4.set_xticklabels(customers, rotation=45, ha='right', fontsize=8)
            ax4.set_ylabel('비율 (%)', fontsize=11)
            ax4.set_title('고객사별 품질 현황', fontsize=13, fontweight='bold')
            ax4.legend(loc='lower right', fontsize=9)
            ax4.set_ylim(85, 100)

        # 5. 라인별 생산
        ax5 = fig.add_subplot(2, 3, 5)
        if line_stats:
            lines = sorted(line_stats.keys())
            good_qty = [line_stats[l]['good'] for l in lines]
            ng_qty = [line_stats[l]['ng'] for l in lines]
            x = np.arange(len(lines))
            ax5.bar(x, good_qty, 0.6, label='양품', color='#2ecc71')
            ax5.bar(x, ng_qty, 0.6, bottom=good_qty, label='불량', color='#e74c3c')
            ax5.set_xticks(x)
            ax5.set_xticklabels(lines, fontsize=10)
            ax5.set_ylabel('수량', fontsize=11)
            ax5.set_title('라인별 생산 실적', fontsize=13, fontweight='bold')
            ax5.legend(loc='upper right', fontsize=9)
            for i, (g, n) in enumerate(zip(good_qty, ng_qty)):
                total = g + n
                yield_rate = g / total * 100 if total > 0 else 0
                ax5.text(i, total + 100, f'{yield_rate:.1f}%', ha='center', fontsize=9, fontweight='bold')

        # 6. KPI 요약
        ax6 = fig.add_subplot(2, 3, 6)
        ax6.axis('off')
        total_records = len(quality_data)
        total_defects = len(fail_cases)
        defect_rate = total_defects / total_records * 100 if total_records > 0 else 0
        avg_yield_all = sum([sum(cs['yields'])/len(cs['yields']) if cs['yields'] else 0 for cs in customer_stats.values()]) / len(customer_stats) if customer_stats else 0

        kpi_data = [
            ('총 검사 건수', f'{total_records:,}건', '#3498db'),
            ('총 불량 건수', f'{total_defects:,}건', '#e74c3c'),
            ('전체 불량률', f'{defect_rate:.1f}%', '#f39c12'),
            ('평균 수율', f'{avg_yield_all:.1f}%', '#2ecc71'),
        ]

        for i, (label, value, color) in enumerate(kpi_data):
            y_pos = 0.82 - i * 0.21
            ax6.add_patch(plt.Rectangle((0.1, y_pos - 0.07), 0.8, 0.16,
                                          facecolor=color, alpha=0.15, transform=ax6.transAxes))
            ax6.text(0.5, y_pos + 0.015, value, transform=ax6.transAxes, fontsize=20,
                     verticalalignment='center', horizontalalignment='center', fontweight='bold', color=color)
            ax6.text(0.5, y_pos - 0.04, label, transform=ax6.transAxes, fontsize=10,
                     verticalalignment='center', horizontalalignment='center', color='#555')

        ax6.set_title('주요 품질 지표 (KPI)', fontsize=13, fontweight='bold', pad=15)

        plt.suptitle('디스플레이 품질 분석 대시보드', fontsize=18, fontweight='bold', y=0.98)
        plt.tight_layout(rect=[0, 0, 1, 0.95])

        # 이미지로 변환
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=120, bbox_inches='tight', facecolor='white')
        buf.seek(0)
        plt.close(fig)

        img = Image.open(buf)

        summary = f"총 {total_records:,}건 분석 | 불량 {total_defects:,}건 ({defect_rate:.1f}%) | 평균 수율 {avg_yield_all:.1f}%"
        return img, summary

    except Exception as e:
        return None, f"차트 생성 오류: {str(e)}"


def generate_defect_chart():
    """불량 유형별 차트 생성"""
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
        import matplotlib.font_manager as fm
        import os

        # 한글 폰트
        for fpath in ['/usr/share/fonts/truetype/nanum/NanumGothic.ttf', '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc']:
            if os.path.exists(fpath):
                fm.fontManager.addfont(fpath)
                plt.rcParams['font.family'] = fm.FontProperties(fname=fpath).get_name()
                plt.rcParams['axes.unicode_minus'] = False
                break

        mock_dir = "/tmp/mockdata_large" if os.path.exists("/tmp/mockdata_large") else "/tmp/mockdata"
        quality_file = None
        qual_dir = os.path.join(mock_dir, 'quality')
        if os.path.exists(qual_dir):
            files = [f for f in os.listdir(qual_dir) if f.endswith('.json')]
            if files:
                quality_file = os.path.join(qual_dir, files[0])

        if not quality_file:
            return None

        with open(quality_file, 'r', encoding='utf-8') as f:
            quality_data = json.load(f)

        fail_cases = [q for q in quality_data if q.get('inspection_result') == 'FAIL']
        defect_types = {}
        for case in fail_cases:
            dt = case.get('defect_type', 'UNKNOWN')
            defect_types[dt] = defect_types.get(dt, 0) + 1

        colors = ['#3498db', '#e74c3c', '#2ecc71', '#f39c12', '#9b59b6', '#1abc9c', '#e67e22', '#34495e']

        fig, ax = plt.subplots(figsize=(10, 8))
        labels = list(defect_types.keys())
        sizes = list(defect_types.values())
        wedges, texts, autotexts = ax.pie(sizes, labels=labels, autopct='%1.1f%%',
                                           colors=colors[:len(labels)], startangle=90, explode=[0.02]*len(labels))
        ax.set_title('불량 유형별 분포', fontsize=16, fontweight='bold', pad=20)
        ax.legend(wedges, [f'{l}: {s:,}건' for l, s in zip(labels, sizes)],
                  title="불량 유형", loc="center left", bbox_to_anchor=(1, 0, 0.5, 1))
        plt.tight_layout()

        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=120, bbox_inches='tight', facecolor='white')
        buf.seek(0)
        plt.close(fig)

        return Image.open(buf)
    except:
        return None


def generate_equipment_chart():
    """설비별 불량률 차트 생성"""
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
        import matplotlib.font_manager as fm
        import os

        for fpath in ['/usr/share/fonts/truetype/nanum/NanumGothic.ttf', '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc']:
            if os.path.exists(fpath):
                fm.fontManager.addfont(fpath)
                plt.rcParams['font.family'] = fm.FontProperties(fname=fpath).get_name()
                plt.rcParams['axes.unicode_minus'] = False
                break

        mock_dir = "/tmp/mockdata_large" if os.path.exists("/tmp/mockdata_large") else "/tmp/mockdata"
        qual_dir = os.path.join(mock_dir, 'quality')
        if not os.path.exists(qual_dir):
            return None

        files = [f for f in os.listdir(qual_dir) if f.endswith('.json')]
        if not files:
            return None

        with open(os.path.join(qual_dir, files[0]), 'r', encoding='utf-8') as f:
            quality_data = json.load(f)

        equipment_stats = {}
        for case in quality_data:
            eq = case.get('equipment_id', 'UNKNOWN')
            if eq not in equipment_stats:
                equipment_stats[eq] = {'total': 0, 'fail': 0}
            equipment_stats[eq]['total'] += 1
            if case.get('inspection_result') == 'FAIL':
                equipment_stats[eq]['fail'] += 1

        top_equipment = sorted(
            [(eq, info['fail']/info['total']*100) for eq, info in equipment_stats.items() if info['total'] >= 50],
            key=lambda x: -x[1]
        )[:10]

        fig, ax = plt.subplots(figsize=(12, 6))
        eq_names = [x[0] for x in top_equipment]
        eq_rates = [x[1] for x in top_equipment]
        bars = ax.bar(eq_names, eq_rates, color='#e74c3c', alpha=0.8)
        avg_rate = sum(eq_rates) / len(eq_rates) if eq_rates else 0
        ax.axhline(y=avg_rate, color='#3498db', linestyle='--', linewidth=2, label=f'평균: {avg_rate:.1f}%')
        ax.set_ylabel('불량률 (%)', fontsize=12)
        ax.set_title('설비별 불량률 TOP 10', fontsize=16, fontweight='bold')
        ax.legend()
        plt.xticks(rotation=45, ha='right')
        for bar, rate in zip(bars, eq_rates):
            ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.15, f'{rate:.1f}%', ha='center', fontsize=9)
        plt.tight_layout()

        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=120, bbox_inches='tight', facecolor='white')
        buf.seek(0)
        plt.close(fig)

        return Image.open(buf)
    except:
        return None


def generate_customer_chart():
    """고객사별 품질 차트 생성"""
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
        import matplotlib.font_manager as fm
        import numpy as np
        import os

        for fpath in ['/usr/share/fonts/truetype/nanum/NanumGothic.ttf', '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc']:
            if os.path.exists(fpath):
                fm.fontManager.addfont(fpath)
                plt.rcParams['font.family'] = fm.FontProperties(fname=fpath).get_name()
                plt.rcParams['axes.unicode_minus'] = False
                break

        mock_dir = "/tmp/mockdata_large" if os.path.exists("/tmp/mockdata_large") else "/tmp/mockdata"
        mfg_dir = os.path.join(mock_dir, 'manufacturing')
        if not os.path.exists(mfg_dir):
            return None

        files = [f for f in os.listdir(mfg_dir) if f.endswith('.json')]
        if not files:
            return None

        with open(os.path.join(mfg_dir, files[0]), 'r', encoding='utf-8') as f:
            mfg_data = json.load(f)

        customer_stats = {}
        for mfg in mfg_data:
            cust = mfg.get('customer', 'UNKNOWN')
            if cust not in customer_stats:
                customer_stats[cust] = {'total': 0, 'pass': 0, 'yields': []}
            customer_stats[cust]['total'] += 1
            if mfg.get('final_result') == 'PASS':
                customer_stats[cust]['pass'] += 1
            customer_stats[cust]['yields'].append(mfg.get('yield_rate', 0))

        fig, ax = plt.subplots(figsize=(12, 6))
        customers = list(customer_stats.keys())
        avg_yields = [sum(customer_stats[c]['yields'])/len(customer_stats[c]['yields']) for c in customers]
        pass_rates = [customer_stats[c]['pass']/customer_stats[c]['total']*100 for c in customers]
        x = np.arange(len(customers))
        width = 0.35
        ax.bar(x - width/2, avg_yields, width, label='평균 수율', color='#3498db')
        ax.bar(x + width/2, pass_rates, width, label='통과율', color='#2ecc71')
        ax.set_xticks(x)
        ax.set_xticklabels(customers, rotation=45, ha='right')
        ax.set_ylabel('비율 (%)', fontsize=12)
        ax.set_title('고객사별 품질 현황', fontsize=16, fontweight='bold')
        ax.legend()
        ax.set_ylim(85, 100)
        plt.tight_layout()

        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=120, bbox_inches='tight', facecolor='white')
        buf.seek(0)
        plt.close(fig)

        return Image.open(buf)
    except:
        return None


# ==================== 그래프 시각화 탭 ====================

def create_graph_visualization():
    """지식 그래프 시각화"""
    global kb

    if not GRAPH_AVAILABLE:
        return None, "networkx/matplotlib가 설치되지 않았습니다."

    if kb is None:
        return None, "지식 베이스를 먼저 초기화해주세요."

    try:
        G = nx.DiGraph()

        # 노드 추가 (밝은 파스텔 색상)
        node_colors = []
        node_labels = {}

        if hasattr(kb.store, 'defects'):
            for d in kb.store.defects.values():
                G.add_node(d.defect_id, type='defect')
                node_labels[d.defect_id] = d.korean_name
                node_colors.append('#ffb3b3')  # 밝은 빨간색

        if hasattr(kb.store, 'root_causes'):
            for c in kb.store.root_causes.values():
                G.add_node(c.cause_id, type='cause')
                node_labels[c.cause_id] = c.korean_name
                node_colors.append('#b3e6e0')  # 밝은 청록색

        if hasattr(kb.store, 'actions'):
            for a in kb.store.actions.values():
                G.add_node(a.action_id, type='action')
                node_labels[a.action_id] = a.korean_name
                node_colors.append('#b3d9ff')  # 밝은 파란색

        if hasattr(kb.store, 'processes'):
            for p in kb.store.processes.values():
                G.add_node(p.process_id, type='process')
                node_labels[p.process_id] = p.korean_name
                node_colors.append('#c8e6c9')  # 밝은 녹색

        # 엣지 추가
        if hasattr(kb.store, 'caused_by'):
            for r in kb.store.caused_by:
                G.add_edge(r.defect_id, r.cause_id, relation='CAUSED_BY', weight=r.probability)

        if hasattr(kb.store, 'requires'):
            for r in kb.store.requires:
                G.add_edge(r.cause_id, r.action_id, relation='REQUIRES', weight=r.effectiveness)

        if hasattr(kb.store, 'occurs_in'):
            for r in kb.store.occurs_in:
                G.add_edge(r.defect_id, r.process_id, relation='OCCURS_IN')

        # 그래프 그리기 (흰색 배경)
        fig, ax = plt.subplots(1, 1, figsize=(14, 10))
        fig.patch.set_facecolor('#ffffff')
        ax.set_facecolor('#ffffff')

        pos = nx.spring_layout(G, k=2, iterations=50, seed=42)

        # 엣지 그리기
        edge_colors = []
        for u, v, data in G.edges(data=True):
            if data.get('relation') == 'CAUSED_BY':
                edge_colors.append('#e53935')
            elif data.get('relation') == 'REQUIRES':
                edge_colors.append('#1e88e5')
            else:
                edge_colors.append('#43a047')

        nx.draw_networkx_edges(G, pos, edge_color=edge_colors, arrows=True,
                               arrowsize=20, arrowstyle='->', alpha=0.8, ax=ax)

        # 노드 그리기
        nx.draw_networkx_nodes(G, pos, node_color=node_colors, node_size=2000, alpha=0.9, ax=ax)

        # 한글 폰트 설정
        font_path = '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc'
        if not Path(font_path).exists():
            font_path = '/usr/share/fonts/truetype/nanum/NanumGothic.ttf'

        if Path(font_path).exists():
            font_prop = font_manager.FontProperties(fname=font_path, size=9)
            legend_font = font_manager.FontProperties(fname=font_path, size=10)
            title_font = font_manager.FontProperties(fname=font_path, size=16, weight='bold')
        else:
            font_prop = None
            legend_font = None
            title_font = None

        # 레이블 (한글 폰트 적용, 흰색 배경에 맞게 어두운 색상)
        for node, (x, y) in pos.items():
            label = node_labels.get(node, node)
            ax.text(x, y, label, fontsize=9, color='#333333', fontweight='bold',
                    ha='center', va='center',
                    fontproperties=font_prop if font_prop else None)

        # 범례 (밝은 파스텔 색상)
        legend_elements = [
            plt.Line2D([0], [0], marker='o', color='w', markerfacecolor='#ffb3b3', markersize=15, label='Defect (결함)', markeredgecolor='#e57373'),
            plt.Line2D([0], [0], marker='o', color='w', markerfacecolor='#b3e6e0', markersize=15, label='RootCause (원인)', markeredgecolor='#4db6ac'),
            plt.Line2D([0], [0], marker='o', color='w', markerfacecolor='#b3d9ff', markersize=15, label='Action (조치)', markeredgecolor='#64b5f6'),
            plt.Line2D([0], [0], marker='o', color='w', markerfacecolor='#c8e6c9', markersize=15, label='Process (공정)', markeredgecolor='#81c784'),
        ]
        legend = ax.legend(handles=legend_elements, loc='upper left', facecolor='#ffffff',
                           labelcolor='#333333', fontsize=10, prop=legend_font if legend_font else None,
                           edgecolor='#e0e0e0', framealpha=0.95)

        ax.set_title('Display Defect Knowledge Graph', color='#333333', fontsize=16, fontweight='bold',
                     fontproperties=title_font if title_font else None)
        ax.axis('off')

        plt.tight_layout()

        # 이미지를 파일로 저장 (흰색 배경)
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=120, facecolor='#ffffff', edgecolor='none')
        buf.seek(0)
        plt.close(fig)

        return Image.open(buf), f"노드: {G.number_of_nodes()}, 엣지: {G.number_of_edges()}"

    except Exception as e:
        return None, f"시각화 오류: {str(e)}"


def create_subgraph_visualization(defect_type):
    """특정 결함 중심 서브그래프 시각화"""
    global kb

    if not GRAPH_AVAILABLE:
        return None, "networkx/matplotlib가 설치되지 않았습니다."

    if kb is None:
        return None, "지식 베이스를 먼저 초기화해주세요."

    try:
        G = nx.DiGraph()

        # 해당 결함 찾기
        target_defect = None
        if hasattr(kb.store, 'defects'):
            for d in kb.store.defects.values():
                if d.defect_type.value == defect_type:
                    target_defect = d
                    break

        if not target_defect:
            return None, f"'{defect_type}' 결함을 찾을 수 없습니다."

        # 결함 노드 추가
        G.add_node(target_defect.defect_id, type='defect', label=target_defect.korean_name)

        # 관련 원인 찾기
        related_causes = []
        if hasattr(kb.store, 'caused_by'):
            for r in kb.store.caused_by:
                if r.defect_id == target_defect.defect_id:
                    cause = kb.store.root_causes.get(r.cause_id)
                    if cause:
                        G.add_node(cause.cause_id, type='cause', label=cause.korean_name)
                        G.add_edge(target_defect.defect_id, cause.cause_id,
                                   label=f'{r.probability:.0%}', relation='CAUSED_BY')
                        related_causes.append(cause.cause_id)

        # 관련 조치 찾기
        if hasattr(kb.store, 'requires'):
            for r in kb.store.requires:
                if r.cause_id in related_causes:
                    action = kb.store.actions.get(r.action_id)
                    if action:
                        G.add_node(action.action_id, type='action', label=action.korean_name)
                        G.add_edge(r.cause_id, action.action_id,
                                   label=f'{r.effectiveness:.0%}', relation='REQUIRES')

        # 그래프 그리기 (흰색 배경)
        fig, ax = plt.subplots(1, 1, figsize=(12, 8))
        fig.patch.set_facecolor('#ffffff')
        ax.set_facecolor('#ffffff')

        pos = nx.spring_layout(G, k=3, iterations=50, seed=42)

        # 노드 색상 (밝은 파스텔)
        node_colors = []
        for node in G.nodes():
            node_type = G.nodes[node].get('type')
            if node_type == 'defect':
                node_colors.append('#ffb3b3')  # 밝은 빨간색
            elif node_type == 'cause':
                node_colors.append('#b3e6e0')  # 밝은 청록색
            else:
                node_colors.append('#b3d9ff')  # 밝은 파란색

        # 엣지 (선명한 색상)
        edge_colors = ['#ef5350' if G.edges[e].get('relation') == 'CAUSED_BY' else '#42a5f5'
                       for e in G.edges()]

        nx.draw_networkx_edges(G, pos, edge_color=edge_colors, arrows=True,
                               arrowsize=25, arrowstyle='->', width=2, ax=ax)
        nx.draw_networkx_nodes(G, pos, node_color=node_colors, node_size=3000, alpha=0.9, ax=ax)

        # 한글 폰트 설정
        font_path = '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc'
        if not Path(font_path).exists():
            font_path = '/usr/share/fonts/truetype/nanum/NanumGothic.ttf'

        if Path(font_path).exists():
            font_prop = font_manager.FontProperties(fname=font_path, size=10)
            edge_font = font_manager.FontProperties(fname=font_path, size=9)
            title_font = font_manager.FontProperties(fname=font_path, size=14, weight='bold')
        else:
            font_prop = None
            edge_font = None
            title_font = None

        # 노드 레이블 (한글 폰트 적용, 어두운 색상)
        labels = {n: G.nodes[n].get('label', n) for n in G.nodes()}
        for node, (x, y) in pos.items():
            label = labels.get(node, node)
            ax.text(x, y, label, fontsize=10, color='#333333', fontweight='bold',
                    ha='center', va='center',
                    fontproperties=font_prop if font_prop else None)

        # 엣지 레이블 (흰색 배경에 맞게 진한 파란색)
        edge_labels = {(u, v): d.get('label', '') for u, v, d in G.edges(data=True)}
        for (u, v), label in edge_labels.items():
            x = (pos[u][0] + pos[v][0]) / 2
            y = (pos[u][1] + pos[v][1]) / 2
            ax.text(x, y, label, fontsize=9, color='#0066cc',
                    ha='center', va='center',
                    fontproperties=edge_font if edge_font else None)

        ax.set_title(f'{target_defect.korean_name} 분석 그래프', color='#333333', fontsize=14, fontweight='bold',
                     fontproperties=title_font if title_font else None)
        ax.axis('off')

        plt.tight_layout()

        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=120, facecolor='#ffffff', edgecolor='none')
        buf.seek(0)
        plt.close(fig)

        return Image.open(buf), f"결함: {target_defect.korean_name}, 연결 노드: {G.number_of_nodes()-1}"

    except Exception as e:
        return None, f"시각화 오류: {str(e)}"


# ==================== Gradio 인터페이스 ====================

LIGHT_CSS = """
/* ===== 분석 탭 메인 레이아웃 ===== */
#analysis-main-row {
    align-items: stretch !important;
}
#history-sidebar {
    display: flex;
    flex-direction: column;
    height: 100%;
    min-height: 700px;
    background: #f8f9fa;
    border-radius: 8px;
    padding: 12px;
    border: 1px solid #e0e0e0;
}

/* ===== 채팅 히스토리 스크롤 리스트 ===== */
#history-list-container {
    flex: 1;
    max-height: 600px;
    overflow-y: auto;
    border: 1px solid #e0e0e0;
    border-radius: 8px;
    padding: 8px;
    background: #ffffff;
}
#history-radio-list {
    max-height: none !important;
}
#history-radio-list label {
    padding: 8px 10px;
    margin: 4px 0;
    border-radius: 6px;
    cursor: pointer;
    transition: background 0.2s;
    font-size: 13px;
    display: block;
    white-space: nowrap;
    overflow: hidden;
    text-overflow: ellipsis;
}
#history-radio-list label:hover {
    background: #e8f0fe !important;
}

/* ===== 전체 배경 흰색 + 글자 검정 ===== */
*, *::before, *::after {
    color: #333333 !important;
}
html, body, .gradio-container, .main, .app {
    background: #ffffff !important;
    background-color: #ffffff !important;
    color: #333333 !important;
}
.dark, [data-theme="dark"], .dark *, [data-theme="dark"] * {
    background: #ffffff !important;
    background-color: #ffffff !important;
    color: #333333 !important;
}

/* ===== 모든 블록/패널 흰색 ===== */
.block, .form, .container, .wrap, .panel, .group, .box,
div[class*="block"], div[class*="panel"], div[class*="group"] {
    background: #ffffff !important;
    background-color: #ffffff !important;
    color: #333333 !important;
}

/* ===== 입력 필드 ===== */
input, textarea, select,
.gr-input, .gr-textbox, .gr-dropdown,
input[type="text"], input[type="password"], input[type="number"] {
    background: #ffffff !important;
    background-color: #ffffff !important;
    color: #333333 !important;
    border: 1px solid #d0d5dd !important;
}
input::placeholder, textarea::placeholder {
    color: #999999 !important;
}

/* ===== 버튼 ===== */
button, .gr-button, .gr-button-secondary, button.secondary {
    background: #f5f7fa !important;
    color: #333333 !important;
    border: 1px solid #d0d5dd !important;
}
button.primary, .gr-button-primary, button[class*="primary"] {
    background: #4a90d9 !important;
    color: #ffffff !important;
}

/* ===== 탭 ===== */
.tabs, .tab-nav, .tabitem, .tab-content {
    background: #ffffff !important;
    background-color: #ffffff !important;
}
.tab-nav button, button[role="tab"] {
    background: #f0f4f8 !important;
    color: #333333 !important;
}
.tab-nav button.selected, button[role="tab"][aria-selected="true"] {
    background: #4a90d9 !important;
    color: #ffffff !important;
}

/* ===== 아코디언 ===== */
.accordion, .accordion-header, details, summary {
    background: #f8fafc !important;
    color: #333333 !important;
}

/* ===== 레이블/텍스트 ===== */
label, .label-wrap, span, p, h1, h2, h3, h4, h5, h6,
.prose, .markdown-body, .md, .text {
    color: #333333 !important;
}

/* ===== 체크박스/라디오 ===== */
.gr-checkbox, .checkbox-container, input[type="checkbox"],
.gr-radio, input[type="radio"] {
    background: #ffffff !important;
}

/* ===== 슬라이더 ===== */
.gr-slider, input[type="range"] {
    background: #ffffff !important;
}

/* ===== 파일 업로드 ===== */
.file-upload, .upload-container, .upload-button,
div[class*="upload"], div[class*="file"] {
    background: #f8fafc !important;
    border-color: #d0d5dd !important;
    color: #333333 !important;
}

/* ===== 이미지 컨테이너 (회색 배경) ===== */
.image-container, .gallery, .gr-image, .gr-gallery,
div[class*="image"], div[class*="gallery"],
.image-frame, .upload-container, .image-preview {
    background: #f0f0f0 !important;
    background-color: #f0f0f0 !important;
}

/* ===== 테이블 ===== */
table, th, td, .dataframe {
    background: #ffffff !important;
    color: #333333 !important;
    border-color: #e0e0e0 !important;
}

/* ===== 코드 블록 ===== */
pre, code, .code {
    background: #f5f5f5 !important;
    color: #333333 !important;
}

/* ===== 스크롤바 ===== */
::-webkit-scrollbar {
    background: #f0f0f0 !important;
}
::-webkit-scrollbar-thumb {
    background: #cccccc !important;
}

/* ===== SVG 아이콘 ===== */
svg, svg path, svg circle, svg rect {
    fill: #333333 !important;
    stroke: #333333 !important;
}
button svg, button svg path {
    fill: currentColor !important;
    stroke: currentColor !important;
}

/* ===== Row/Column 정렬 ===== */
.contain {
    max-width: 100% !important;
    width: 100% !important;
}
.gradio-container {
    max-width: 100% !important;
    width: 100% !important;
    margin: 0 auto !important;
    padding: 20px !important;
}

/* ===== 라이트박스/모달 팝업 스타일 ===== */
.image-popup-overlay {
    position: fixed !important;
    top: 0 !important;
    left: 0 !important;
    width: 100vw !important;
    height: 100vh !important;
    background: rgba(0, 0, 0, 0.85) !important;
    z-index: 9999 !important;
    display: flex !important;
    justify-content: center !important;
    align-items: center !important;
    padding: 20px !important;
}
.image-popup-content {
    max-width: 90vw !important;
    max-height: 90vh !important;
    background: #ffffff !important;
    border-radius: 8px !important;
    padding: 10px !important;
    box-shadow: 0 10px 50px rgba(0, 0, 0, 0.5) !important;
}
.image-popup-content img {
    max-width: 85vw !important;
    max-height: 80vh !important;
    object-fit: contain !important;
}
.popup-close-btn {
    position: absolute !important;
    top: 20px !important;
    right: 30px !important;
    font-size: 40px !important;
    color: #ffffff !important;
    cursor: pointer !important;
    z-index: 10000 !important;
    background: rgba(0, 0, 0, 0.5) !important;
    border-radius: 50% !important;
    width: 50px !important;
    height: 50px !important;
    display: flex !important;
    justify-content: center !important;
    align-items: center !important;
    line-height: 1 !important;
}
.popup-close-btn:hover {
    background: rgba(255, 0, 0, 0.7) !important;
}
.clickable-image {
    cursor: zoom-in !important;
}
.clickable-image:hover {
    opacity: 0.9 !important;
    box-shadow: 0 0 10px rgba(74, 144, 217, 0.5) !important;
}
"""


# ===== CS 워크플로우 함수 =====
def init_cs_workflow():
    """CS 워크플로우 시스템 초기화"""
    global cs_manager, quality_analyzer, report_generator

    if not CS_WORKFLOW_AVAILABLE:
        return "CS 워크플로우 모듈을 로드할 수 없습니다."

    try:
        cs_manager = CSComplaintManager(data_dir="/tmp/cs_complaints")
        quality_analyzer = QualityAnalyzer(data_dir="/tmp/quality_analysis")
        report_generator = ReportGenerator(output_dir="/tmp/cs_reports")

        # 샘플 과거 사례 생성
        report_generator.generate_sample_past_cases(count=5)

        return "CS 워크플로우 시스템 초기화 완료\n- 불만 관리자\n- 품질 분석기\n- 보고서 생성기\n- 과거 사례 5건 생성"
    except Exception as e:
        return f"초기화 오류: {str(e)}"


def create_cs_complaint(customer, product_model, lot_id, cell_id, defect_type, defect_description, severity):
    """CS 불만 접수 생성"""
    global cs_manager

    if cs_manager is None:
        init_cs_workflow()

    if cs_manager is None:
        return "CS 워크플로우를 먼저 초기화하세요.", None

    try:
        complaint = cs_manager.create_complaint(
            customer=customer,
            product_model=product_model,
            lot_id=lot_id,
            cell_id=cell_id,
            defect_type=defect_type,
            defect_description=defect_description,
            severity=severity
        )

        result = f"""### 불만 접수 완료

| 항목 | 값 |
|------|-----|
| 접수번호 | {complaint.complaint_id} |
| 고객사 | {complaint.customer} |
| 제품모델 | {complaint.product_model} |
| LOT ID | {complaint.lot_id} |
| CELL ID | {complaint.cell_id} |
| 결함유형 | {complaint.defect_type} |
| 심각도 | {complaint.severity} |
| 상태 | {complaint.status} |
"""
        return result, complaint.complaint_id
    except Exception as e:
        return f"접수 오류: {str(e)}", None


def perform_first_analysis(complaint_id):
    """1차 기본 분석 수행"""
    global cs_manager, quality_analyzer

    if quality_analyzer is None:
        init_cs_workflow()

    if quality_analyzer is None or cs_manager is None:
        return "CS 워크플로우를 먼저 초기화하세요.", None

    try:
        complaint = cs_manager.get_complaint(complaint_id)
        if not complaint:
            return f"불만 접수 ID를 찾을 수 없습니다: {complaint_id}", None

        analysis = quality_analyzer.perform_first_analysis(
            complaint_id=complaint.complaint_id,
            defect_type=complaint.defect_type,
            lot_id=complaint.lot_id,
            cell_id=complaint.cell_id,
            product_model=complaint.product_model,
            analyst="QA_System"
        )

        # 불만 데이터 업데이트
        cs_manager.update_first_analysis(
            complaint_id=complaint.complaint_id,
            result=analysis.to_dict(),
            analyst="QA_System"
        )

        bigdata = analysis.bigdata_result
        equipment_list = ", ".join(bigdata.get("related_equipment", [])[:3])
        cases_list = ", ".join(analysis.similar_cases[:3])

        result = f"""### 1차 기본 분석 완료 (품질부서)

#### 분석 결과
| 항목 | 값 |
|------|-----|
| 분석 ID | {analysis.analysis_id} |
| 귀책 부서 추정 | **{analysis.responsible_dept}** |
| 신뢰도 | {analysis.confidence_score*100:.1f}% |

#### 빅데이터 분석
- 결함 통계: 총 {bigdata.get('defect_statistics', {}).get('total_count', 'N/A')}건
- 고위험 설비: {equipment_list}
- 연관 LOT: {len(bigdata.get('related_lots', []))}건

#### 유사 사례
{cases_list}

---
**권장 조치**: 귀책 부서({analysis.responsible_dept})에 2차 상세 분석 요청
"""
        return result, analysis.analysis_id
    except Exception as e:
        return f"분석 오류: {str(e)}", None


def perform_second_analysis(complaint_id, first_analysis_id):
    """2차 상세 분석 수행"""
    global cs_manager, quality_analyzer

    if quality_analyzer is None:
        return "CS 워크플로우를 먼저 초기화하세요.", None

    try:
        complaint = cs_manager.get_complaint(complaint_id)
        if not complaint:
            return f"불만 접수 ID를 찾을 수 없습니다: {complaint_id}", None

        analysis = quality_analyzer.perform_second_analysis(
            complaint_id=complaint.complaint_id,
            first_analysis_id=first_analysis_id,
            defect_image_path=complaint.defect_image_path,
            analyst="Dept_Analyst"
        )

        # 불만 데이터 업데이트
        cs_manager.update_second_analysis(
            complaint_id=complaint.complaint_id,
            result=analysis.to_dict(),
            root_cause=analysis.root_cause,
            countermeasure=", ".join(analysis.countermeasures[:3])
        )

        # GraphRAG 결과
        graphrag = analysis.graphrag_result
        cause_analysis = graphrag.get("cause_analysis", {})

        # 과거 사례
        past_cases = analysis.past_case_result.get("cases_analyzed", [])

        result = f"""### 2차 상세 분석 완료 (귀책부서)

#### 근본 원인
**{analysis.root_cause}**

#### 이미지 분석 (Cosmos VLM)
- 검출 결함: {len(analysis.image_result.get('detected_defects', []))}건
- 신뢰도: {analysis.image_result.get('vlm_confidence', 0)*100:.1f}%
- 특성: {analysis.image_result.get('defect_characteristics', {}).get('pattern', 'N/A')}

#### GraphRAG 분석
- 주요 원인: {cause_analysis.get('primary_cause', 'N/A')}
- 부가 원인: {', '.join(cause_analysis.get('secondary_causes', [])[:2])}

#### 과거 사례 참조
"""
        for case in past_cases[:2]:
            result += f"- **{case.get('case_id', 'N/A')}**: {case.get('root_cause', 'N/A')} (유사도 {case.get('similarity_score', 0)*100:.0f}%)\n"

        result += f"""
#### 대책 수립
**즉각 대책:**
"""
        for i, measure in enumerate(analysis.countermeasures[:3], 1):
            result += f"{i}. {measure}\n"

        result += """
**재발 방지 대책:**
"""
        for i, measure in enumerate(analysis.prevention_measures[:3], 1):
            result += f"{i}. {measure}\n"

        return result, analysis.analysis_id
    except Exception as e:
        return f"분석 오류: {str(e)}", None


def generate_final_report(complaint_id):
    """최종 보고서 생성"""
    global cs_manager, quality_analyzer, report_generator

    if report_generator is None:
        return "CS 워크플로우를 먼저 초기화하세요.", None

    try:
        complaint = cs_manager.get_complaint(complaint_id)
        if not complaint:
            return f"불만 접수 ID를 찾을 수 없습니다: {complaint_id}", None

        # 분석 결과 가져오기
        analyses = quality_analyzer.get_results_by_complaint(complaint_id)
        first_analysis = None
        second_analysis = None

        for a in analyses:
            if "1차" in a.analysis_type:
                first_analysis = a.to_dict()
            elif "2차" in a.analysis_type:
                second_analysis = a.to_dict()

        if not first_analysis:
            first_analysis = {}
        if not second_analysis:
            second_analysis = {}

        # 보고서 생성
        report_path = report_generator.generate_final_report(
            complaint_data=complaint.to_dict(),
            first_analysis=first_analysis,
            second_analysis=second_analysis
        )

        # 불만 처리 완료
        cs_manager.complete_complaint(complaint_id)

        summary = report_generator.get_report_summary(report_path)

        result = f"""### 최종 보고서 생성 완료

| 항목 | 값 |
|------|-----|
| 파일 경로 | {report_path} |
| 파일 형식 | {summary.get('format', 'Unknown')} |
| 파일 크기 | {summary.get('size_kb', 0):.2f} KB |
| 생성 시간 | {summary.get('created', 'N/A')} |

**불만 처리 상태**: 완료

---
보고서를 다운로드하여 Office Copilot에서 추가 편집할 수 있습니다.
"""
        return result, report_path
    except Exception as e:
        return f"보고서 생성 오류: {str(e)}", None


def get_complaints_list():
    """불만 접수 목록 조회"""
    global cs_manager

    if cs_manager is None:
        init_cs_workflow()

    if cs_manager is None:
        return "CS 워크플로우를 먼저 초기화하세요."

    try:
        complaints = cs_manager.get_all_complaints()

        if not complaints:
            return "등록된 불만 접수가 없습니다."

        result = "### 불만 접수 목록\n\n"
        result += "| 접수번호 | 고객사 | 결함유형 | 상태 | 귀책부서 |\n"
        result += "|----------|--------|----------|------|----------|\n"

        for c in complaints[-10:]:  # 최근 10건
            result += f"| {c.complaint_id} | {c.customer} | {c.defect_type} | {c.status} | {c.responsible_dept} |\n"

        return result
    except Exception as e:
        return f"조회 오류: {str(e)}"


# ===== 이메일 생성 기능 =====
import requests

# 고객사별 담당자 정보
CUSTOMER_CONTACTS = {
    "APPLE": {"name": "Apple Quality Team", "email": "quality@apple.com", "title": "Quality Manager"},
    "SAMSUNG_MOBILE": {"name": "삼성 품질관리팀", "email": "quality@samsung.com", "title": "품질관리팀장"},
    "LG_MOBILE": {"name": "LG 품질관리팀", "email": "quality@lge.com", "title": "품질담당"},
    "GOOGLE": {"name": "Google Quality Team", "email": "quality@google.com", "title": "Quality Lead"},
    "XIAOMI": {"name": "Xiaomi Quality Team", "email": "quality@xiaomi.com", "title": "Quality Director"},
    "HUAWEI": {"name": "Huawei Quality Team", "email": "quality@huawei.com", "title": "Quality Manager"},
    "META": {"name": "Meta Quality Team", "email": "quality@meta.com", "title": "Quality Lead"},
}

# 사내 부서별 담당자 정보
INTERNAL_DEPARTMENTS = {
    "CS팀": {"name": "CS팀", "email": "cs-team@sdc.com", "manager": "김민수"},
    "품질관리팀": {"name": "품질관리팀", "email": "quality-team@sdc.com", "manager": "이정호"},
    "TFT공정": {"name": "TFT공정팀", "email": "tft-process@sdc.com", "manager": "박성준"},
    "CF공정": {"name": "CF공정팀", "email": "cf-process@sdc.com", "manager": "최영희"},
    "OLED공정": {"name": "OLED공정팀", "email": "oled-process@sdc.com", "manager": "정우진"},
    "Module공정": {"name": "Module공정팀", "email": "module-process@sdc.com", "manager": "한미영"},
    "Cell공정": {"name": "Cell공정팀", "email": "cell-process@sdc.com", "manager": "강동원"},
    "검사팀": {"name": "검사팀", "email": "inspection@sdc.com", "manager": "윤서연"},
}

# 결함 유형별 예상 귀책 부서
DEFECT_RESPONSIBLE_DEPT = {
    "DEAD_PIXEL": "TFT공정",
    "BRIGHT_SPOT": "OLED공정",
    "LINE_DEFECT": "TFT공정",
    "MURA": "CF공정",
    "SCRATCH": "Module공정",
    "TOUCH_FAIL": "Cell공정",
}

# ===== 내부 결재 시스템 =====
from datetime import datetime, timedelta
import uuid

# 결재 상태
APPROVAL_STATUS = {
    "PENDING": "대기",
    "APPROVED": "승인",
    "REJECTED": "반려",
    "CANCELLED": "취소"
}

# 결재 유형
APPROVAL_TYPES = {
    "COMPLAINT_EMAIL": "고객 접수 확인 메일",
    "MEETING_REQUEST": "미팅 요청 메일",
    "FIRST_ANALYSIS": "1차 분석 결과",
    "TASK_ASSIGNMENT": "업무 할당",
    "FINAL_REPORT": "최종 보고서",
    "CUSTOMER_REPLY": "고객 회신 메일"
}

# 결재 유형별 기본 기한 (시간 단위)
APPROVAL_DEADLINES = {
    "COMPLAINT_EMAIL": 4,      # 4시간 이내
    "MEETING_REQUEST": 8,      # 8시간 이내
    "FIRST_ANALYSIS": 24,      # 24시간 이내
    "TASK_ASSIGNMENT": 8,      # 8시간 이내
    "FINAL_REPORT": 48,        # 48시간 이내
    "CUSTOMER_REPLY": 24,      # 24시간 이내
}

# 알림 임계값 (기한까지 남은 시간, 시간 단위)
DEADLINE_WARNING_THRESHOLD = 2  # 2시간 전 경고
DEADLINE_URGENT_THRESHOLD = 1   # 1시간 전 긴급

# 결재 라인 (결재자 정보)
APPROVAL_LINE = {
    "COMPLAINT_EMAIL": [
        {"position": "팀장", "name": "이정호", "email": "jhlee@sdc.com", "dept": "품질관리팀"},
    ],
    "MEETING_REQUEST": [
        {"position": "팀장", "name": "이정호", "email": "jhlee@sdc.com", "dept": "품질관리팀"},
    ],
    "FIRST_ANALYSIS": [
        {"position": "팀장", "name": "이정호", "email": "jhlee@sdc.com", "dept": "품질관리팀"},
        {"position": "부장", "name": "김대영", "email": "dykim@sdc.com", "dept": "품질본부"},
    ],
    "TASK_ASSIGNMENT": [
        {"position": "팀장", "name": "이정호", "email": "jhlee@sdc.com", "dept": "품질관리팀"},
    ],
    "FINAL_REPORT": [
        {"position": "팀장", "name": "이정호", "email": "jhlee@sdc.com", "dept": "품질관리팀"},
        {"position": "부장", "name": "김대영", "email": "dykim@sdc.com", "dept": "품질본부"},
        {"position": "상무", "name": "박철수", "email": "cspark@sdc.com", "dept": "품질담당"},
    ],
    "CUSTOMER_REPLY": [
        {"position": "팀장", "name": "이정호", "email": "jhlee@sdc.com", "dept": "품질관리팀"},
        {"position": "부장", "name": "김대영", "email": "dykim@sdc.com", "dept": "품질본부"},
    ],
}

# 결재 저장소 (메모리 기반)
approval_storage = {}

def create_approval_request(approval_type, complaint_id, title, content, requester="담당자", custom_deadline_hours=None):
    """결재 요청 생성"""
    approval_id = f"APR-{datetime.now().strftime('%Y%m%d%H%M%S')}-{uuid.uuid4().hex[:4].upper()}"

    approvers = APPROVAL_LINE.get(approval_type, [])

    # 기한 계산
    deadline_hours = custom_deadline_hours if custom_deadline_hours else APPROVAL_DEADLINES.get(approval_type, 24)
    created_at = datetime.now()
    deadline = created_at + timedelta(hours=deadline_hours)

    approval_request = {
        "approval_id": approval_id,
        "type": approval_type,
        "type_name": APPROVAL_TYPES.get(approval_type, approval_type),
        "complaint_id": complaint_id,
        "title": title,
        "content": content,
        "requester": requester,
        "created_at": created_at.isoformat(),
        "deadline": deadline.isoformat(),
        "deadline_hours": deadline_hours,
        "status": "PENDING",
        "approvers": [],
        "current_step": 0,
        "total_steps": len(approvers),
        "history": []
    }

    # 결재자 목록 초기화
    for i, approver in enumerate(approvers):
        approval_request["approvers"].append({
            "step": i + 1,
            "position": approver["position"],
            "name": approver["name"],
            "email": approver["email"],
            "dept": approver["dept"],
            "status": "PENDING" if i == 0 else "WAITING",
            "approved_at": None,
            "comment": ""
        })

    approval_storage[approval_id] = approval_request

    # 결재 요청 로그 저장
    save_approval_log(approval_request, "CREATE")

    return approval_id, approval_request

def get_approval_status(approval_id):
    """결재 상태 조회"""
    if approval_id not in approval_storage:
        return None
    return approval_storage[approval_id]

def get_pending_approvals(approver_email=None):
    """대기중인 결재 목록 조회"""
    pending = []
    for approval_id, approval in approval_storage.items():
        if approval["status"] == "PENDING":
            if approver_email:
                current_step = approval["current_step"]
                if current_step < len(approval["approvers"]):
                    current_approver = approval["approvers"][current_step]
                    if current_approver["email"] == approver_email:
                        pending.append(approval)
            else:
                pending.append(approval)
    return pending

def process_approval(approval_id, action, comment="", approver_name=""):
    """결재 처리 (승인/반려)"""
    if approval_id not in approval_storage:
        return False, "결재 요청을 찾을 수 없습니다."

    approval = approval_storage[approval_id]

    if approval["status"] != "PENDING":
        return False, f"이미 처리된 결재입니다. (상태: {APPROVAL_STATUS.get(approval['status'], approval['status'])})"

    current_step = approval["current_step"]

    if current_step >= len(approval["approvers"]):
        return False, "모든 결재가 완료되었습니다."

    current_approver = approval["approvers"][current_step]

    if action == "APPROVE":
        current_approver["status"] = "APPROVED"
        current_approver["approved_at"] = datetime.now().isoformat()
        current_approver["comment"] = comment

        approval["history"].append({
            "step": current_step + 1,
            "action": "APPROVE",
            "approver": current_approver["name"],
            "position": current_approver["position"],
            "timestamp": datetime.now().isoformat(),
            "comment": comment
        })

        # 다음 결재자로 이동
        approval["current_step"] += 1

        # 모든 결재 완료 확인
        if approval["current_step"] >= approval["total_steps"]:
            approval["status"] = "APPROVED"
            save_approval_log(approval, "COMPLETE")
            return True, "최종 승인이 완료되었습니다."
        else:
            # 다음 결재자 상태 변경
            approval["approvers"][approval["current_step"]]["status"] = "PENDING"
            save_approval_log(approval, "STEP")
            next_approver = approval["approvers"][approval["current_step"]]
            return True, f"승인 완료. 다음 결재자: {next_approver['position']} {next_approver['name']}"

    elif action == "REJECT":
        current_approver["status"] = "REJECTED"
        current_approver["approved_at"] = datetime.now().isoformat()
        current_approver["comment"] = comment

        approval["status"] = "REJECTED"
        approval["history"].append({
            "step": current_step + 1,
            "action": "REJECT",
            "approver": current_approver["name"],
            "position": current_approver["position"],
            "timestamp": datetime.now().isoformat(),
            "comment": comment
        })

        save_approval_log(approval, "REJECT")
        return True, f"반려되었습니다. 사유: {comment}"

    return False, "잘못된 액션입니다."

def save_approval_log(approval, action_type):
    """결재 로그 저장"""
    import os

    log_dir = "/tmp/cs_approvals"
    os.makedirs(log_dir, exist_ok=True)

    log_file = f"{log_dir}/approval_{approval['approval_id']}.json"

    with open(log_file, 'w', encoding='utf-8') as f:
        json.dump(approval, f, ensure_ascii=False, indent=2)

    # 전체 로그 파일에도 기록
    all_log_file = f"{log_dir}/approval_history.log"
    with open(all_log_file, 'a', encoding='utf-8') as f:
        f.write(f"[{datetime.now().isoformat()}] {action_type}: {approval['approval_id']} - {approval['type_name']} - {APPROVAL_STATUS.get(approval['status'], approval['status'])}\n")

def format_approval_status(approval):
    """결재 상태를 마크다운으로 포맷팅"""
    if not approval:
        return "결재 정보가 없습니다."

    status_icon = {
        "PENDING": "🕐",
        "APPROVED": "✅",
        "REJECTED": "❌",
        "WAITING": "⏳",
        "CANCELLED": "🚫"
    }

    result = f"""### 결재 정보
**결재번호:** {approval['approval_id']}
**결재유형:** {approval['type_name']}
**관련 접수번호:** {approval['complaint_id']}
**요청일시:** {approval['created_at'][:19].replace('T', ' ')}
**상태:** {status_icon.get(approval['status'], '')} {APPROVAL_STATUS.get(approval['status'], approval['status'])}

---
#### 결재 라인
| 순서 | 직급 | 결재자 | 부서 | 상태 | 처리일시 | 의견 |
|:----:|:----:|:------:|:----:|:----:|:--------:|:----:|
"""

    for approver in approval["approvers"]:
        status = status_icon.get(approver["status"], "") + " " + APPROVAL_STATUS.get(approver["status"], approver["status"])
        approved_at = approver["approved_at"][:19].replace('T', ' ') if approver["approved_at"] else "-"
        comment = approver["comment"] if approver["comment"] else "-"
        result += f"| {approver['step']} | {approver['position']} | {approver['name']} | {approver['dept']} | {status} | {approved_at} | {comment} |\n"

    return result

def get_all_approvals_for_complaint(complaint_id):
    """특정 불만 접수에 대한 모든 결재 조회"""
    approvals = []
    for approval_id, approval in approval_storage.items():
        if approval["complaint_id"] == complaint_id:
            approvals.append(approval)
    return sorted(approvals, key=lambda x: x["created_at"])


# ===== 결재 기한 알림 함수 =====
def get_deadline_status(approval):
    """결재 기한 상태 확인"""
    if approval["status"] != "PENDING":
        return "completed", 0, ""

    deadline_str = approval.get("deadline")
    if not deadline_str:
        return "no_deadline", 0, ""

    try:
        deadline = datetime.fromisoformat(deadline_str)
        now = datetime.now()
        remaining = deadline - now
        remaining_hours = remaining.total_seconds() / 3600

        if remaining_hours < 0:
            return "overdue", abs(remaining_hours), f"⚠️ 기한 초과 ({abs(remaining_hours):.1f}시간)"
        elif remaining_hours <= DEADLINE_URGENT_THRESHOLD:
            return "urgent", remaining_hours, f"🔴 긴급 (잔여 {remaining_hours:.1f}시간)"
        elif remaining_hours <= DEADLINE_WARNING_THRESHOLD:
            return "warning", remaining_hours, f"🟡 임박 (잔여 {remaining_hours:.1f}시간)"
        else:
            return "normal", remaining_hours, f"🟢 정상 (잔여 {remaining_hours:.1f}시간)"
    except:
        return "error", 0, ""


def get_overdue_approvals():
    """기한 초과 결재 목록"""
    overdue = []
    for approval in approval_storage.values():
        if approval["status"] == "PENDING":
            status, hours, _ = get_deadline_status(approval)
            if status == "overdue":
                overdue.append((approval, hours))
    return sorted(overdue, key=lambda x: x[1], reverse=True)


def get_urgent_approvals():
    """긴급 결재 목록 (기한 임박)"""
    urgent = []
    for approval in approval_storage.values():
        if approval["status"] == "PENDING":
            status, hours, _ = get_deadline_status(approval)
            if status in ["urgent", "warning"]:
                urgent.append((approval, hours))
    return sorted(urgent, key=lambda x: x[1])


def get_deadline_alerts():
    """기한 알림 조회"""
    overdue = get_overdue_approvals()
    urgent = get_urgent_approvals()

    if not overdue and not urgent:
        return """### ⏰ 기한 알림

✅ 기한 임박하거나 초과된 결재가 없습니다.
"""

    result = "### ⏰ 기한 알림\n\n"

    if overdue:
        result += f"""#### ⚠️ 기한 초과 ({len(overdue)}건)
| 결재번호 | 유형 | 접수ID | 현재 결재자 | 초과시간 |
|:---------|:-----|:-------|:------------|:---------|
"""
        for approval, hours in overdue:
            current_step = approval["current_step"]
            if current_step < len(approval["approvers"]):
                current_approver = approval["approvers"][current_step]
                approver_info = f"{current_approver['position']} {current_approver['name']}"
            else:
                approver_info = "-"
            result += f"| {approval['approval_id']} | {approval['type_name']} | {approval['complaint_id']} | {approver_info} | **{hours:.1f}시간** |\n"
        result += "\n"

    if urgent:
        result += f"""#### 🔴 기한 임박 ({len(urgent)}건)
| 결재번호 | 유형 | 접수ID | 현재 결재자 | 잔여시간 |
|:---------|:-----|:-------|:------------|:---------|
"""
        for approval, hours in urgent:
            current_step = approval["current_step"]
            if current_step < len(approval["approvers"]):
                current_approver = approval["approvers"][current_step]
                approver_info = f"{current_approver['position']} {current_approver['name']}"
            else:
                approver_info = "-"
            icon = "🔴" if hours <= DEADLINE_URGENT_THRESHOLD else "🟡"
            result += f"| {approval['approval_id']} | {approval['type_name']} | {approval['complaint_id']} | {approver_info} | {icon} **{hours:.1f}시간** |\n"

    return result


def get_deadline_summary():
    """기한 현황 요약"""
    overdue_count = len(get_overdue_approvals())
    urgent_count = len(get_urgent_approvals())
    pending_count = sum(1 for a in approval_storage.values() if a["status"] == "PENDING")
    normal_count = pending_count - overdue_count - urgent_count

    return f"""| 상태 | 건수 |
|:----:|:----:|
| ⚠️ 기한초과 | **{overdue_count}** |
| 🔴 긴급 | **{urgent_count}** |
| 🟢 정상 | **{normal_count}** |
"""


def format_approval_with_deadline(approval):
    """기한 정보를 포함한 결재 상태 포맷팅"""
    base_status = format_approval_status(approval)

    deadline_str = approval.get("deadline")
    if deadline_str:
        try:
            deadline = datetime.fromisoformat(deadline_str)
            deadline_display = deadline.strftime("%Y-%m-%d %H:%M")
            status, hours, status_msg = get_deadline_status(approval)

            deadline_info = f"""
---
#### ⏰ 기한 정보
- **결재 기한:** {deadline_display}
- **기한 상태:** {status_msg}
- **설정 기한:** {approval.get('deadline_hours', 'N/A')}시간
"""
            return base_status + deadline_info
        except:
            pass

    return base_status


# ===== 기한 초과 알림 이메일 함수 =====
# 알림 발송 이력 저장소
notification_history = {}

def send_deadline_notification_email(approval, notification_type="overdue"):
    """기한 초과/임박 알림 이메일 발송"""
    import os

    approval_id = approval["approval_id"]
    current_step = approval["current_step"]

    if current_step >= len(approval["approvers"]):
        return False, "결재자 정보가 없습니다."

    current_approver = approval["approvers"][current_step]
    recipient_email = current_approver["email"]
    recipient_name = current_approver["name"]
    recipient_position = current_approver["position"]

    # 기한 정보
    deadline_str = approval.get("deadline", "")
    try:
        deadline = datetime.fromisoformat(deadline_str)
        deadline_display = deadline.strftime("%Y-%m-%d %H:%M")
    except:
        deadline_display = "N/A"

    status, hours, status_msg = get_deadline_status(approval)

    # 알림 유형에 따른 메시지
    if notification_type == "overdue":
        subject = f"[긴급] 결재 기한 초과 알림 - {approval_id}"
        urgency = "⚠️ 긴급: 결재 기한이 초과되었습니다!"
        time_info = f"기한 초과: {hours:.1f}시간"
    elif notification_type == "urgent":
        subject = f"[알림] 결재 기한 임박 - {approval_id}"
        urgency = "🔴 주의: 결재 기한이 임박했습니다!"
        time_info = f"잔여 시간: {hours:.1f}시간"
    else:
        subject = f"[알림] 결재 요청 - {approval_id}"
        urgency = "결재 요청이 있습니다."
        time_info = f"잔여 시간: {hours:.1f}시간"

    email_content = f"""
{recipient_position} {recipient_name}님께,

{urgency}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
■ 결재 정보
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
• 결재번호: {approval_id}
• 결재유형: {approval['type_name']}
• 관련 접수번호: {approval['complaint_id']}
• 요청자: {approval['requester']}
• 요청일시: {approval['created_at'][:19].replace('T', ' ')}
• 결재기한: {deadline_display}
• {time_info}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
■ 결재 제목
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
{approval['title']}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

결재 시스템에 접속하여 결재를 처리해 주시기 바랍니다.
http://localhost:7860 → CS 워크플로우 → 결재 현황

감사합니다.
SDC 품질관리시스템
"""

    # 이메일 로그 저장
    email_log_dir = "/tmp/cs_approval_notifications"
    os.makedirs(email_log_dir, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    log_file = f"{email_log_dir}/{approval_id}_{notification_type}_{timestamp}.txt"

    with open(log_file, 'w', encoding='utf-8') as f:
        f.write(f"To: {recipient_name} <{recipient_email}>\n")
        f.write(f"Subject: {subject}\n")
        f.write(f"Date: {datetime.now().isoformat()}\n")
        f.write(f"Notification Type: {notification_type}\n")
        f.write("-" * 50 + "\n")
        f.write(email_content)

    # 알림 이력 저장
    if approval_id not in notification_history:
        notification_history[approval_id] = []

    notification_history[approval_id].append({
        "type": notification_type,
        "sent_at": datetime.now().isoformat(),
        "recipient": recipient_email,
        "recipient_name": recipient_name
    })

    return True, f"알림 이메일 발송 완료: {recipient_name} ({recipient_email})"


def send_all_overdue_notifications():
    """모든 기한 초과 결재에 대해 알림 발송"""
    overdue = get_overdue_approvals()
    urgent = get_urgent_approvals()

    sent_count = 0
    results = []

    # 기한 초과 건 알림
    for approval, hours in overdue:
        approval_id = approval["approval_id"]

        # 최근 1시간 내 동일 알림 발송 여부 확인
        if not should_send_notification(approval_id, "overdue", hours_threshold=1):
            results.append(f"⏭️ {approval_id}: 최근 알림 발송됨 (스킵)")
            continue

        success, message = send_deadline_notification_email(approval, "overdue")
        if success:
            sent_count += 1
            results.append(f"✅ {approval_id}: {message}")
        else:
            results.append(f"❌ {approval_id}: {message}")

    # 긴급 건 알림 (1시간 이내)
    for approval, hours in urgent:
        if hours > DEADLINE_URGENT_THRESHOLD:
            continue  # 긴급 임계값 초과는 스킵

        approval_id = approval["approval_id"]

        if not should_send_notification(approval_id, "urgent", hours_threshold=0.5):
            results.append(f"⏭️ {approval_id}: 최근 알림 발송됨 (스킵)")
            continue

        success, message = send_deadline_notification_email(approval, "urgent")
        if success:
            sent_count += 1
            results.append(f"✅ {approval_id}: {message}")
        else:
            results.append(f"❌ {approval_id}: {message}")

    return sent_count, results


def should_send_notification(approval_id, notification_type, hours_threshold=1):
    """알림 발송 여부 확인 (중복 발송 방지)"""
    if approval_id not in notification_history:
        return True

    for notification in notification_history[approval_id]:
        if notification["type"] == notification_type:
            try:
                sent_at = datetime.fromisoformat(notification["sent_at"])
                hours_since = (datetime.now() - sent_at).total_seconds() / 3600
                if hours_since < hours_threshold:
                    return False
            except:
                pass

    return True


def get_notification_history_for_approval(approval_id):
    """특정 결재의 알림 발송 이력 조회"""
    if approval_id not in notification_history:
        return "알림 발송 이력이 없습니다."

    history = notification_history[approval_id]
    result = f"### 알림 발송 이력 ({len(history)}건)\n\n"
    result += "| 유형 | 발송일시 | 수신자 |\n"
    result += "|:----:|:---------|:-------|\n"

    for h in sorted(history, key=lambda x: x["sent_at"], reverse=True):
        type_icon = "⚠️" if h["type"] == "overdue" else "🔴" if h["type"] == "urgent" else "📧"
        type_name = "기한초과" if h["type"] == "overdue" else "긴급" if h["type"] == "urgent" else "일반"
        sent_time = h["sent_at"][:19].replace("T", " ")
        result += f"| {type_icon} {type_name} | {sent_time} | {h['recipient_name']} |\n"

    return result


def format_notification_results(sent_count, results):
    """알림 발송 결과 포맷팅"""
    result = f"### 📧 알림 발송 결과\n\n"
    result += f"**발송 완료:** {sent_count}건\n\n"

    if results:
        result += "| 상태 | 결재번호 | 결과 |\n"
        result += "|:----:|:---------|:-----|\n"

        for r in results:
            if r.startswith("✅"):
                status = "✅"
                content = r[2:]
            elif r.startswith("❌"):
                status = "❌"
                content = r[2:]
            else:
                status = "⏭️"
                content = r[3:]

            parts = content.split(":", 1)
            approval_id = parts[0].strip()
            message = parts[1].strip() if len(parts) > 1 else ""
            result += f"| {status} | {approval_id} | {message} |\n"
    else:
        result += "_발송할 알림이 없습니다._\n"

    return result


# ===== 결재 대시보드 함수 =====
def get_approval_dashboard_summary():
    """결재 현황 요약 통계"""
    total = len(approval_storage)
    pending = sum(1 for a in approval_storage.values() if a["status"] == "PENDING")
    approved = sum(1 for a in approval_storage.values() if a["status"] == "APPROVED")
    rejected = sum(1 for a in approval_storage.values() if a["status"] == "REJECTED")

    # 결재 유형별 통계
    type_stats = {}
    for approval in approval_storage.values():
        type_name = approval["type_name"]
        if type_name not in type_stats:
            type_stats[type_name] = {"total": 0, "pending": 0, "approved": 0, "rejected": 0}
        type_stats[type_name]["total"] += 1
        if approval["status"] == "PENDING":
            type_stats[type_name]["pending"] += 1
        elif approval["status"] == "APPROVED":
            type_stats[type_name]["approved"] += 1
        elif approval["status"] == "REJECTED":
            type_stats[type_name]["rejected"] += 1

    summary = f"""## 결재 현황 대시보드

### 전체 현황
| 구분 | 건수 |
|:----:|:----:|
| 🕐 대기 | **{pending}** |
| ✅ 승인 | **{approved}** |
| ❌ 반려 | **{rejected}** |
| 📊 전체 | **{total}** |

---
### 결재 유형별 현황
| 유형 | 대기 | 승인 | 반려 | 전체 |
|:-----|:----:|:----:|:----:|:----:|
"""
    for type_name, stats in type_stats.items():
        summary += f"| {type_name} | {stats['pending']} | {stats['approved']} | {stats['rejected']} | {stats['total']} |\n"

    if not type_stats:
        summary += "| (결재 내역 없음) | - | - | - | - |\n"

    return summary


def get_pending_approvals_list():
    """대기 중인 결재 목록"""
    pending = [a for a in approval_storage.values() if a["status"] == "PENDING"]
    pending = sorted(pending, key=lambda x: x["created_at"], reverse=True)

    if not pending:
        return "### 🕐 대기 중인 결재\n\n대기 중인 결재가 없습니다."

    result = """### 🕐 대기 중인 결재

| 결재번호 | 유형 | 접수ID | 현재 결재자 | 요청일시 |
|:---------|:-----|:-------|:------------|:---------|
"""
    for approval in pending[:20]:  # 최대 20건
        current_step = approval["current_step"]
        if current_step < len(approval["approvers"]):
            current_approver = approval["approvers"][current_step]
            approver_info = f"{current_approver['position']} {current_approver['name']}"
        else:
            approver_info = "-"

        created = approval["created_at"][:16].replace("T", " ")
        result += f"| {approval['approval_id']} | {approval['type_name']} | {approval['complaint_id']} | {approver_info} | {created} |\n"

    return result


def get_recent_approvals_list():
    """최근 처리된 결재 목록"""
    processed = [a for a in approval_storage.values() if a["status"] in ["APPROVED", "REJECTED"]]
    processed = sorted(processed, key=lambda x: x["created_at"], reverse=True)

    if not processed:
        return "### 📋 최근 처리된 결재\n\n처리된 결재가 없습니다."

    result = """### 📋 최근 처리된 결재

| 결재번호 | 유형 | 접수ID | 상태 | 처리일시 |
|:---------|:-----|:-------|:----:|:---------|
"""
    for approval in processed[:20]:  # 최대 20건
        status_icon = "✅" if approval["status"] == "APPROVED" else "❌"
        status_text = APPROVAL_STATUS.get(approval["status"], approval["status"])

        # 마지막 처리 일시 조회
        last_time = approval["created_at"]
        if approval["history"]:
            last_time = approval["history"][-1]["timestamp"]
        last_time = last_time[:16].replace("T", " ")

        result += f"| {approval['approval_id']} | {approval['type_name']} | {approval['complaint_id']} | {status_icon} {status_text} | {last_time} |\n"

    return result


def get_approval_detail(approval_id):
    """결재 상세 정보 조회"""
    approval = get_approval_status(approval_id)
    if not approval:
        return "결재 정보를 찾을 수 없습니다."

    return format_approval_status(approval) + f"\n\n---\n#### 결재 내용\n```\n{approval['content'][:1000]}{'...' if len(approval['content']) > 1000 else ''}\n```"


def search_approvals(search_type, search_value):
    """결재 검색"""
    results = []

    if search_type == "결재번호":
        if search_value in approval_storage:
            results = [approval_storage[search_value]]
    elif search_type == "접수ID":
        results = [a for a in approval_storage.values() if a["complaint_id"] == search_value]
    elif search_type == "상태":
        status_map = {"대기": "PENDING", "승인": "APPROVED", "반려": "REJECTED"}
        target_status = status_map.get(search_value, search_value)
        results = [a for a in approval_storage.values() if a["status"] == target_status]
    elif search_type == "결재유형":
        results = [a for a in approval_storage.values() if search_value in a["type_name"]]

    if not results:
        return "검색 결과가 없습니다."

    result_md = f"### 검색 결과 ({len(results)}건)\n\n"
    result_md += "| 결재번호 | 유형 | 접수ID | 상태 | 요청일시 |\n"
    result_md += "|:---------|:-----|:-------|:----:|:---------|\n"

    for approval in sorted(results, key=lambda x: x["created_at"], reverse=True):
        status_icon = {"PENDING": "🕐", "APPROVED": "✅", "REJECTED": "❌"}.get(approval["status"], "")
        status_text = APPROVAL_STATUS.get(approval["status"], approval["status"])
        created = approval["created_at"][:16].replace("T", " ")
        result_md += f"| {approval['approval_id']} | {approval['type_name']} | {approval['complaint_id']} | {status_icon} {status_text} | {created} |\n"

    return result_md


def call_ollama_llm(prompt, system_prompt=None):
    """Ollama LLM 호출"""
    try:
        full_prompt = f"{system_prompt}\n\n{prompt}" if system_prompt else prompt
        response = requests.post(
            "http://localhost:11434/api/generate",
            json={
                "model": "llama4:scout",
                "prompt": full_prompt,
                "stream": False
            },
            timeout=120
        )
        response.raise_for_status()
        return response.json().get("response", "").strip()
    except Exception as e:
        return f"[LLM 오류] {str(e)}"


def generate_complaint_email(customer, product_model, lot_id, cell_id, defect_type, defect_description, severity, complaint_id):
    """불만 접수 확인 이메일 생성 (LLM 사용)"""
    contact = CUSTOMER_CONTACTS.get(customer, {"name": "고객 담당자", "email": "customer@example.com", "title": "담당자"})

    system_prompt = """당신은 디스플레이 제조회사의 품질관리 담당자입니다.
고객사로부터 불만 접수가 들어왔을 때 보내는 공식 이메일을 작성해야 합니다.
이메일은 정중하고 전문적인 톤으로 작성하세요.
한국어로 작성하되, 고객사가 해외인 경우에도 한국어로 작성합니다.
이메일 형식으로만 출력하세요."""

    prompt = f"""다음 불만 접수 정보를 바탕으로 고객사에게 보낼 접수 확인 이메일을 작성하세요:

불만 접수 정보:
- 접수번호: {complaint_id}
- 고객사: {customer}
- 담당자: {contact['name']} ({contact['title']})
- 제품모델: {product_model}
- LOT ID: {lot_id}
- CELL ID: {cell_id}
- 결함유형: {defect_type}
- 결함설명: {defect_description}
- 심각도: {severity}

이메일에는 다음 내용을 포함하세요:
1. 불만 접수 확인 및 접수번호 안내
2. 접수된 불만 내용 요약
3. 향후 분석 일정 (1차 분석 3일 이내, 2차 분석 5일 이내)
4. 담당자 연락처 안내
5. 정중한 마무리

발신자: SDC 품질관리팀 (quality@sdc.com)
"""

    email_content = call_ollama_llm(prompt, system_prompt)

    # 이메일 메타 정보
    email_meta = f"""**수신자:** {contact['name']} <{contact['email']}>
**발신자:** SDC 품질관리팀 <quality@sdc.com>
**제목:** [SDC] 품질 불만 접수 확인 - {complaint_id}

---

"""
    return email_meta + email_content, contact['email'], contact['name']


def translate_email(email_content, target_language):
    """이메일 번역 (LLM 사용)"""
    language_prompts = {
        "영어": "Translate to professional English",
        "일본어": "日本語に翻訳してください (Translate to professional Japanese)",
        "중국어": "翻译成专业中文 (Translate to professional Chinese)",
        "한국어": "전문적인 한국어로 번역하세요",
    }

    system_prompt = f"""당신은 전문 비즈니스 번역가입니다.
이메일을 {target_language}로 번역하세요.
전문적이고 정중한 톤을 유지하세요.
이메일 형식을 그대로 유지하면서 번역만 수행하세요.
수신자/발신자 정보는 번역하지 마세요."""

    prompt = f"""{language_prompts.get(target_language, f"Translate to {target_language}")}

원본 이메일:
{email_content}
"""

    translated = call_ollama_llm(prompt, system_prompt)
    return translated


def send_complaint_email(recipient_email, recipient_name, email_content, complaint_id):
    """이메일 전송 (시뮬레이션)"""
    # 실제 환경에서는 SMTP 또는 이메일 API 사용
    # 여기서는 시뮬레이션으로 로그 저장

    import os
    from datetime import datetime

    email_log_dir = "/tmp/cs_emails"
    os.makedirs(email_log_dir, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    log_file = os.path.join(email_log_dir, f"email_{complaint_id}_{timestamp}.txt")

    with open(log_file, 'w', encoding='utf-8') as f:
        f.write(f"전송 시간: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write(f"수신자: {recipient_name} <{recipient_email}>\n")
        f.write(f"접수번호: {complaint_id}\n")
        f.write("="*50 + "\n\n")
        f.write(email_content)

    return f"""### ✅ 이메일 전송 완료

| 항목 | 값 |
|------|-----|
| 수신자 | {recipient_name} |
| 이메일 | {recipient_email} |
| 접수번호 | {complaint_id} |
| 전송시간 | {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} |
| 로그파일 | {log_file} |

> 📧 이메일이 성공적으로 전송되었습니다. (시뮬레이션)
"""


def generate_meeting_request_email(customer, product_model, defect_type, defect_description, severity, complaint_id):
    """사내 미팅 요청 이메일 생성 (LLM 사용)"""
    from datetime import datetime, timedelta

    # 예상 귀책 부서 결정
    responsible_dept = DEFECT_RESPONSIBLE_DEPT.get(defect_type, "품질관리팀")

    # 참석 부서 목록
    attendee_depts = ["CS팀", "품질관리팀", responsible_dept]
    attendee_depts = list(set(attendee_depts))  # 중복 제거

    # 참석자 정보 수집
    attendees = []
    attendee_emails = []
    for dept in attendee_depts:
        dept_info = INTERNAL_DEPARTMENTS.get(dept, {})
        if dept_info:
            attendees.append(f"- {dept_info['name']}: {dept_info['manager']} ({dept_info['email']})")
            attendee_emails.append(dept_info['email'])

    attendees_str = "\n".join(attendees)
    attendee_emails_str = "; ".join(attendee_emails)

    # 미팅 일정 제안 (다음 영업일 오후 2시)
    meeting_date = datetime.now() + timedelta(days=1)
    if meeting_date.weekday() >= 5:  # 주말이면 월요일로
        meeting_date += timedelta(days=(7 - meeting_date.weekday()))
    meeting_time = meeting_date.strftime("%Y-%m-%d") + " 14:00"

    system_prompt = """당신은 디스플레이 제조회사의 CS팀 담당자입니다.
고객 불만 접수 건에 대한 사내 긴급 미팅을 요청하는 이메일을 작성해야 합니다.
이메일은 공식적이고 간결하게 작성하세요.
한국어로 작성하세요."""

    prompt = f"""다음 불만 접수 정보를 바탕으로 사내 미팅 요청 이메일을 작성하세요:

불만 접수 정보:
- 접수번호: {complaint_id}
- 고객사: {customer}
- 제품모델: {product_model}
- 결함유형: {defect_type}
- 결함설명: {defect_description}
- 심각도: {severity}
- 예상 귀책부서: {responsible_dept}

미팅 정보:
- 일시: {meeting_time}
- 장소: 본관 3층 회의실 A
- 참석 대상:
{attendees_str}

이메일에는 다음 내용을 포함하세요:
1. 미팅 목적 (고객 불만 대응 긴급 회의)
2. 불만 접수 요약
3. 미팅 일시 및 장소
4. 참석 요청 부서 및 담당자
5. 회신 요청

발신자: CS팀 (cs-team@sdc.com)
"""

    email_content = call_ollama_llm(prompt, system_prompt)

    # 이메일 메타 정보
    email_meta = f"""**수신자:** {attendee_emails_str}
**발신자:** CS팀 <cs-team@sdc.com>
**제목:** [긴급] 고객 불만 대응 미팅 요청 - {complaint_id}

**참석 대상:**
{attendees_str}

**예상 귀책부서:** {responsible_dept}
**미팅 일시:** {meeting_time}

---

"""
    return email_meta + email_content, attendee_emails_str, responsible_dept, meeting_time


def send_meeting_request_email(attendee_emails, email_content, complaint_id, meeting_time):
    """미팅 요청 이메일 전송 (시뮬레이션)"""
    import os
    from datetime import datetime

    email_log_dir = "/tmp/cs_emails"
    os.makedirs(email_log_dir, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    log_file = os.path.join(email_log_dir, f"meeting_{complaint_id}_{timestamp}.txt")

    with open(log_file, 'w', encoding='utf-8') as f:
        f.write(f"전송 시간: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write(f"수신자: {attendee_emails}\n")
        f.write(f"접수번호: {complaint_id}\n")
        f.write(f"미팅일시: {meeting_time}\n")
        f.write("="*50 + "\n\n")
        f.write(email_content)

    return f"""### ✅ 미팅 요청 이메일 전송 완료

| 항목 | 값 |
|------|-----|
| 수신자 | {attendee_emails} |
| 접수번호 | {complaint_id} |
| 미팅일시 | {meeting_time} |
| 전송시간 | {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} |
| 로그파일 | {log_file} |

> 📅 미팅 요청 이메일이 성공적으로 전송되었습니다. (시뮬레이션)
"""


def generate_meeting_summary(complaint_id, first_analysis_result, meeting_notes):
    """1차 미팅 결과 정리 (LLM 사용)"""
    system_prompt = """당신은 품질관리 회의록 작성 전문가입니다.
회의 내용을 바탕으로 체계적인 회의록을 작성하세요.
한국어로 작성하세요."""

    prompt = f"""다음 1차 분석 결과와 미팅 노트를 바탕으로 회의록을 작성하세요:

접수번호: {complaint_id}

1차 분석 결과:
{first_analysis_result}

미팅 노트:
{meeting_notes}

회의록에 포함할 내용:
1. 회의 개요 (일시, 참석자)
2. 논의 내용 요약
3. 결정 사항
4. 액션 아이템 (담당자, 내용, 납기)
5. 다음 단계
"""

    summary = call_ollama_llm(prompt, system_prompt)
    return summary


def generate_task_assignment_email(complaint_id, tasks_data):
    """업무 할당 이메일 생성 (LLM 사용)"""
    system_prompt = """당신은 프로젝트 관리자입니다.
업무 할당 내용을 바탕으로 각 담당자에게 보낼 공식 이메일을 작성하세요.
한국어로 작성하세요."""

    prompt = f"""다음 업무 할당 내용을 바탕으로 업무 할당 이메일을 작성하세요:

접수번호: {complaint_id}

할당된 업무:
{tasks_data}

이메일에 포함할 내용:
1. 업무 배경 (고객 불만 대응)
2. 할당된 업무 내용
3. 납기일
4. 산출물 요구사항
5. 문의처
"""

    email_content = call_ollama_llm(prompt, system_prompt)
    return email_content


def send_task_assignment_email(recipient_dept, recipient_email, email_content, complaint_id, task_desc, deadline):
    """업무 할당 이메일 전송 (시뮬레이션)"""
    import os
    from datetime import datetime

    email_log_dir = "/tmp/cs_emails"
    os.makedirs(email_log_dir, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    log_file = os.path.join(email_log_dir, f"task_{complaint_id}_{timestamp}.txt")

    with open(log_file, 'w', encoding='utf-8') as f:
        f.write(f"전송 시간: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write(f"수신자: {recipient_dept} <{recipient_email}>\n")
        f.write(f"접수번호: {complaint_id}\n")
        f.write(f"업무: {task_desc}\n")
        f.write(f"납기: {deadline}\n")
        f.write("="*50 + "\n\n")
        f.write(email_content)

    return f"""### ✅ 업무 할당 이메일 전송 완료

| 항목 | 값 |
|------|-----|
| 수신 부서 | {recipient_dept} |
| 이메일 | {recipient_email} |
| 업무 | {task_desc} |
| 납기 | {deadline} |
| 전송시간 | {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} |

> 📋 업무 할당 이메일이 전송되었습니다. (시뮬레이션)
"""


def register_first_output(complaint_id, output_type, output_description, file_path=None):
    """1차 산출물 등록"""
    import os
    from datetime import datetime
    import json

    output_dir = "/tmp/cs_outputs"
    os.makedirs(output_dir, exist_ok=True)

    output_data = {
        "complaint_id": complaint_id,
        "output_type": output_type,
        "description": output_description,
        "file_path": file_path,
        "registered_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "status": "등록완료"
    }

    output_file = os.path.join(output_dir, f"output_{complaint_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json")
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(output_data, f, ensure_ascii=False, indent=2)

    return f"""### ✅ 1차 산출물 등록 완료

| 항목 | 값 |
|------|-----|
| 접수번호 | {complaint_id} |
| 산출물 유형 | {output_type} |
| 설명 | {output_description} |
| 등록시간 | {output_data['registered_at']} |
| 저장경로 | {output_file} |

> 📁 1차 산출물이 등록되었습니다. 2차 분석에서 활용됩니다.
"""


def generate_first_output_report(complaint_id, meeting_summary, tasks_summary, outputs_summary):
    """1차 산출물 보고서 생성"""
    from datetime import datetime

    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # 제목
        title = doc.add_heading('1차 분석 산출물 보고서', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 기본 정보
        doc.add_heading('1. 기본 정보', level=1)
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        cells = table.rows[0].cells
        cells[0].text = '접수번호'
        cells[1].text = complaint_id
        cells = table.rows[1].cells
        cells[0].text = '작성일'
        cells[1].text = datetime.now().strftime('%Y-%m-%d')
        cells = table.rows[2].cells
        cells[0].text = '작성자'
        cells[1].text = '품질관리팀'

        # 미팅 결과
        doc.add_heading('2. 1차 미팅 결과', level=1)
        doc.add_paragraph(meeting_summary)

        # 업무 할당 현황
        doc.add_heading('3. 업무 할당 현황', level=1)
        doc.add_paragraph(tasks_summary)

        # 산출물 목록
        doc.add_heading('4. 1차 산출물 목록', level=1)
        doc.add_paragraph(outputs_summary)

        # 다음 단계
        doc.add_heading('5. 다음 단계', level=1)
        doc.add_paragraph("- 2차 상세 분석 진행")
        doc.add_paragraph("- 귀책 부서 상세 조사")
        doc.add_paragraph("- 개선 대책 수립")

        # 저장
        import os
        output_dir = "/tmp/cs_reports"
        os.makedirs(output_dir, exist_ok=True)
        report_path = os.path.join(output_dir, f"1차산출물보고서_{complaint_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
        doc.save(report_path)

        return f"""### ✅ 1차 산출물 보고서 생성 완료

| 항목 | 값 |
|------|-----|
| 접수번호 | {complaint_id} |
| 보고서 유형 | 1차 산출물 보고서 |
| 생성시간 | {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} |
| 저장경로 | {report_path} |

> 📄 1차 산출물 보고서가 생성되었습니다.
""", report_path

    except Exception as e:
        return f"보고서 생성 오류: {str(e)}", None


def register_second_output(complaint_id, output_type, output_description, file_path=None):
    """2차 산출물 등록"""
    import os
    from datetime import datetime
    import json

    output_dir = "/tmp/cs_outputs"
    os.makedirs(output_dir, exist_ok=True)

    output_data = {
        "complaint_id": complaint_id,
        "phase": "2차 분석",
        "output_type": output_type,
        "description": output_description,
        "file_path": file_path,
        "registered_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "status": "등록완료"
    }

    output_file = os.path.join(output_dir, f"output2_{complaint_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json")
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(output_data, f, ensure_ascii=False, indent=2)

    return f"""### ✅ 2차 산출물 등록 완료

| 항목 | 값 |
|------|-----|
| 접수번호 | {complaint_id} |
| 단계 | 2차 분석 |
| 산출물 유형 | {output_type} |
| 설명 | {output_description} |
| 등록시간 | {output_data['registered_at']} |
| 저장경로 | {output_file} |

> 📁 2차 산출물이 등록되었습니다. 최종 보고서에 포함됩니다.
"""


def generate_second_output_report(complaint_id, second_result, outputs_summary):
    """2차 산출물 보고서 생성"""
    from datetime import datetime

    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # 제목
        title = doc.add_heading('2차 분석 산출물 보고서', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 기본 정보
        doc.add_heading('1. 기본 정보', level=1)
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        cells = table.rows[0].cells
        cells[0].text = '접수번호'
        cells[1].text = complaint_id
        cells = table.rows[1].cells
        cells[0].text = '작성일'
        cells[1].text = datetime.now().strftime('%Y-%m-%d')
        cells = table.rows[2].cells
        cells[0].text = '작성부서'
        cells[1].text = '귀책부서'

        # 2차 분석 결과
        doc.add_heading('2. 2차 분석 결과', level=1)
        doc.add_paragraph(second_result if second_result else "2차 분석 결과 없음")

        # 산출물 목록
        doc.add_heading('3. 2차 산출물 목록', level=1)
        doc.add_paragraph(outputs_summary if outputs_summary else "등록된 산출물 없음")

        # 조치 사항
        doc.add_heading('4. 조치 사항', level=1)
        doc.add_paragraph("- 상세 원인 분석 완료")
        doc.add_paragraph("- 재발 방지 대책 수립")
        doc.add_paragraph("- 공정 개선 방안 도출")
        doc.add_paragraph("- 품질 관리 강화 방안 마련")

        # 다음 단계
        doc.add_heading('5. 다음 단계', level=1)
        doc.add_paragraph("- 최종 보고서 작성")
        doc.add_paragraph("- 고객사 회신")
        doc.add_paragraph("- 개선 조치 이행")

        # 저장
        import os
        output_dir = "/tmp/cs_reports"
        os.makedirs(output_dir, exist_ok=True)
        report_path = os.path.join(output_dir, f"2차산출물보고서_{complaint_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
        doc.save(report_path)

        return f"""### ✅ 2차 산출물 보고서 생성 완료

| 항목 | 값 |
|------|-----|
| 접수번호 | {complaint_id} |
| 보고서 유형 | 2차 산출물 보고서 |
| 생성시간 | {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} |
| 저장경로 | {report_path} |

> 📄 2차 산출물 보고서가 생성되었습니다.
""", report_path

    except Exception as e:
        return f"보고서 생성 오류: {str(e)}", None


def generate_customer_reply_email(customer, complaint_id, defect_type, report_path):
    """고객 회신 이메일 생성 (LLM 사용)"""
    contact = CUSTOMER_CONTACTS.get(customer, {"name": "고객 담당자", "email": "customer@example.com", "title": "담당자"})

    system_prompt = """당신은 디스플레이 제조회사의 품질관리 담당자입니다.
고객 불만에 대한 분석이 완료되어 결과를 회신하는 공식 이메일을 작성해야 합니다.
이메일은 정중하고 전문적인 톤으로 작성하세요.
한국어로 작성하세요."""

    prompt = f"""다음 정보를 바탕으로 고객사에게 보낼 분석 완료 회신 이메일을 작성하세요:

분석 완료 정보:
- 접수번호: {complaint_id}
- 고객사: {customer}
- 담당자: {contact['name']} ({contact['title']})
- 결함유형: {defect_type}
- 최종 보고서: 첨부됨

이메일에는 다음 내용을 포함하세요:
1. 불만 접수에 대한 감사 인사
2. 분석 완료 안내
3. 분석 결과 요약 (원인 파악, 개선 조치)
4. 첨부 보고서 안내
5. 재발 방지 약속
6. 추가 문의 연락처
7. 정중한 마무리

발신자: SDC 품질관리팀 (quality@sdc.com)
"""

    email_content = call_ollama_llm(prompt, system_prompt)

    # 이메일 메타 정보
    email_meta = f"""**수신자:** {contact['name']} <{contact['email']}>
**발신자:** SDC 품질관리팀 <quality@sdc.com>
**제목:** [SDC] 품질 불만 분석 완료 회신 - {complaint_id}
**첨부:** 최종분석보고서.docx

---

"""
    return email_meta + email_content, contact['email'], contact['name']


def send_customer_reply_email(recipient_email, recipient_name, email_content, complaint_id, report_path):
    """고객 회신 이메일 전송 (시뮬레이션)"""
    import os
    from datetime import datetime

    email_log_dir = "/tmp/cs_emails"
    os.makedirs(email_log_dir, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    log_file = os.path.join(email_log_dir, f"reply_{complaint_id}_{timestamp}.txt")

    with open(log_file, 'w', encoding='utf-8') as f:
        f.write(f"전송 시간: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write(f"수신자: {recipient_name} <{recipient_email}>\n")
        f.write(f"접수번호: {complaint_id}\n")
        f.write(f"첨부파일: {report_path}\n")
        f.write("="*50 + "\n\n")
        f.write(email_content)

    return f"""### ✅ 고객 회신 이메일 전송 완료

| 항목 | 값 |
|------|-----|
| 수신자 | {recipient_name} |
| 이메일 | {recipient_email} |
| 접수번호 | {complaint_id} |
| 첨부파일 | {report_path} |
| 전송시간 | {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} |
| 로그파일 | {log_file} |

> 📧 고객 회신 이메일이 성공적으로 전송되었습니다. (시뮬레이션)
> 📎 최종 보고서가 첨부되어 발송되었습니다.
"""


def create_demo():
    with gr.Blocks(
        title="디스플레이 결함 분석 시스템",
    ) as demo:

        # 헤더 (배경 없음, 높이 50% 감소)
        gr.HTML("""
        <div style="text-align: center; padding: 10px 20px; margin-bottom: 15px; border-bottom: 2px solid #e0e0e0; position: relative;">
            <h1 style="color: #333333; margin: 0; font-size: 24px;">디스플레이 결함 분석 시스템</h1>
            <p style="color: #666666; margin: 5px 0 0 0; font-size: 13px;">Cosmos Reason VLM + GraphRAG 기반 지능형 품질 검사</p>
            <a href="http://localhost:3000" target="_blank"
               style="position: absolute; right: 20px; top: 50%; transform: translateY(-50%);
                      background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                      color: white; padding: 8px 16px; border-radius: 6px; text-decoration: none;
                      font-size: 13px; font-weight: 500; box-shadow: 0 2px 4px rgba(0,0,0,0.1);
                      transition: all 0.2s ease;">
                🔗 고객품질분석시스템
            </a>
        </div>
        """)

        # 탭 구성
        with gr.Tabs(selected="cs_workflow"):

            # ===== 탭 1: CS 워크플로우 =====
            with gr.TabItem("1. CS 워크플로우", id="cs_workflow"):
                gr.Markdown("""
                ### CS 품질 불량 분석 워크플로우
                **프로세스**: 불만 접수 → 1차 기본분석 → 2차 상세분석 → 최종 보고서
                """)

                # 공유 상태 변수
                cs_shared_complaint_id = gr.State(value="")
                cs_shared_first_analysis_id = gr.State(value="")
                cs_shared_second_analysis_id = gr.State(value="")

                with gr.Tabs():
                    # ===== 페이지 1: 불만 접수 =====
                    with gr.TabItem("Step 1. 불만 접수", id="cs_step1"):
                        gr.Markdown("""
                        ## 📋 CS 불만 접수
                        고객사로부터 접수된 품질 불량 정보를 입력합니다.
                        """)

                        with gr.Row():
                            with gr.Column(scale=1):
                                gr.Markdown("#### 시스템 초기화")
                                cs_init_btn = gr.Button("CS 워크플로우 초기화", variant="primary", size="lg")
                                cs_init_result = gr.Textbox(label="초기화 결과", lines=2)

                            with gr.Column(scale=1):
                                gr.Markdown("#### 접수 목록")
                                cs_list_btn = gr.Button("불만 접수 목록 조회", variant="secondary")

                        gr.Markdown("---")
                        gr.Markdown("#### 불만 정보 입력")

                        with gr.Row():
                            with gr.Column(scale=1):
                                cs_customer = gr.Dropdown(
                                    choices=["APPLE", "SAMSUNG_MOBILE", "LG_MOBILE", "GOOGLE", "XIAOMI", "HUAWEI", "META"],
                                    label="고객사",
                                    value="APPLE"
                                )
                                cs_product = gr.Dropdown(
                                    choices=["OLED_67_FHD", "OLED_61_FHD", "LTPO_68_QHD", "LCD_109_2K", "OLED_76_FOLD"],
                                    label="제품모델",
                                    value="OLED_67_FHD"
                                )
                                cs_lot_id = gr.Textbox(label="LOT ID", value="LOT20241203001")
                                cs_cell_id = gr.Textbox(label="CELL ID", value="CELL12345")

                            with gr.Column(scale=1):
                                cs_defect_type = gr.Dropdown(
                                    choices=["DEAD_PIXEL", "BRIGHT_SPOT", "LINE_DEFECT", "MURA", "SCRATCH", "TOUCH_FAIL"],
                                    label="결함유형",
                                    value="DEAD_PIXEL"
                                )
                                cs_severity = gr.Dropdown(
                                    choices=["LOW", "MEDIUM", "HIGH", "CRITICAL"],
                                    label="심각도",
                                    value="HIGH"
                                )
                                cs_description = gr.Textbox(
                                    label="결함 설명",
                                    value="화면 중앙부에 검은색 점 발견. 크기 약 0.3mm",
                                    lines=3
                                )

                        cs_create_btn = gr.Button("불만 접수 등록", variant="primary", size="lg")

                        gr.Markdown("---")
                        gr.Markdown("#### 접수 결과")
                        cs_complaint_id_display = gr.Textbox(label="생성된 불만 접수 ID", interactive=False)
                        cs_create_result = gr.Markdown(value="시스템 초기화 후 불만 정보를 입력하고 접수하세요.")

                        gr.Markdown("---")
                        gr.Markdown("#### 📧 고객 확인 이메일")

                        with gr.Row():
                            with gr.Column(scale=2):
                                cs_email_content = gr.Textbox(
                                    label="이메일 내용 (수정 가능)",
                                    lines=15,
                                    placeholder="불만 접수 후 이메일이 자동 생성됩니다...",
                                    interactive=True
                                )

                            with gr.Column(scale=1):
                                gr.Markdown("**이메일 옵션**")
                                cs_recipient_email = gr.Textbox(label="수신자 이메일", interactive=False)
                                cs_recipient_name = gr.Textbox(label="수신자 이름", interactive=False)

                                cs_email_language = gr.Dropdown(
                                    choices=["한국어", "영어", "일본어", "중국어"],
                                    label="이메일 언어",
                                    value="한국어"
                                )
                                cs_translate_btn = gr.Button("언어 변환", variant="secondary")

                                gr.Markdown("---")
                                gr.Markdown("**결재 요청**")
                                cs_email_approval_btn = gr.Button("📋 결재 요청", variant="secondary")
                                cs_email_approval_id = gr.Textbox(label="결재번호", interactive=False)

                        gr.Markdown("---")
                        gr.Markdown("#### 결재 처리 (고객 확인 이메일)")

                        with gr.Row():
                            with gr.Column(scale=1):
                                cs_email_approval_status = gr.Markdown(value="결재 요청을 먼저 생성하세요.")
                            with gr.Column(scale=1):
                                cs_email_approval_comment = gr.Textbox(
                                    label="결재 의견",
                                    placeholder="승인/반려 의견을 입력하세요",
                                    lines=2
                                )
                                with gr.Row():
                                    cs_email_approve_btn = gr.Button("✅ 승인", variant="primary")
                                    cs_email_reject_btn = gr.Button("❌ 반려", variant="stop")
                                cs_email_approval_result = gr.Markdown()

                        with gr.Row():
                            cs_send_email_btn = gr.Button("📤 고객 이메일 전송 (결재완료 후)", variant="primary", size="lg")
                        cs_email_result = gr.Markdown()

                        gr.Markdown("---")
                        gr.Markdown("#### 📅 사내 미팅 요청 이메일")

                        with gr.Row():
                            with gr.Column(scale=2):
                                cs_meeting_email_content = gr.Textbox(
                                    label="미팅 요청 이메일 내용 (수정 가능)",
                                    lines=15,
                                    placeholder="불만 접수 후 '미팅 요청 메일 생성' 버튼을 클릭하세요...",
                                    interactive=True
                                )

                            with gr.Column(scale=1):
                                gr.Markdown("**미팅 정보**")
                                cs_meeting_attendees = gr.Textbox(label="참석자 이메일", interactive=False)
                                cs_meeting_responsible_dept = gr.Textbox(label="예상 귀책부서", interactive=False)
                                cs_meeting_time = gr.Textbox(label="미팅 일시", interactive=False)

                                cs_generate_meeting_btn = gr.Button("📝 미팅 요청 메일 생성", variant="secondary", size="lg")

                                gr.Markdown("---")
                                cs_send_meeting_btn = gr.Button("📤 미팅 요청 전송", variant="primary", size="lg")
                                cs_meeting_result = gr.Markdown()

                        cs_list_result = gr.Markdown()

                    # ===== 페이지 2: 1차 분석 =====
                    with gr.TabItem("Step 2. 1차 분석", id="cs_step2"):
                        gr.Markdown("""
                        ## 🔍 1차 기본 분석 (품질부서)
                        품질부서에서 결함 유형을 분석하고 귀책 부서를 결정합니다.
                        """)

                        with gr.Row():
                            with gr.Column(scale=1):
                                gr.Markdown("#### 분석 대상")
                                cs_first_complaint_id = gr.Textbox(
                                    label="불만 접수 ID",
                                    placeholder="Step 1에서 생성된 ID를 입력하세요",
                                    info="Step 1에서 생성된 불만 접수 ID를 입력합니다"
                                )
                                cs_first_analysis_btn = gr.Button("1차 분석 실행", variant="primary", size="lg")

                            with gr.Column(scale=1):
                                gr.Markdown("#### 분석 정보")
                                gr.Markdown("""
                                **1차 분석 내용:**
                                - 결함 유형 분류
                                - 원인 추정
                                - 귀책 부서 결정
                                - 긴급도 평가
                                """)

                        gr.Markdown("---")
                        gr.Markdown("#### 1차 분석 결과")
                        cs_first_analysis_id_display = gr.Textbox(label="1차 분석 ID", interactive=False)
                        cs_first_result = gr.Markdown(value="불만 접수 ID를 입력하고 1차 분석을 실행하세요.")

                        gr.Markdown("---")
                        gr.Markdown("#### 📋 1차 미팅 결과 정리")

                        with gr.Row():
                            with gr.Column(scale=1):
                                cs_meeting_notes = gr.Textbox(
                                    label="미팅 노트 입력",
                                    lines=5,
                                    placeholder="미팅에서 논의된 내용을 입력하세요...\n예: 원인 분석 결과, 결정 사항, 추가 조사 필요 항목 등"
                                )
                                cs_generate_summary_btn = gr.Button("📝 미팅 결과 정리 (LLM)", variant="secondary")

                            with gr.Column(scale=1):
                                cs_meeting_summary = gr.Textbox(
                                    label="미팅 결과 (수정 가능)",
                                    lines=8,
                                    interactive=True
                                )

                        gr.Markdown("---")
                        gr.Markdown("#### 👥 업무 할당 및 납기 등록")

                        with gr.Row():
                            with gr.Column(scale=1):
                                cs_task_dept = gr.Dropdown(
                                    choices=["TFT공정", "CF공정", "OLED공정", "Module공정", "Cell공정", "검사팀", "품질관리팀"],
                                    label="담당 부서",
                                    value="TFT공정"
                                )
                                cs_task_description = gr.Textbox(
                                    label="업무 내용",
                                    placeholder="할당할 업무 내용을 입력하세요",
                                    lines=2
                                )
                                cs_task_deadline = gr.Textbox(
                                    label="납기일",
                                    placeholder="예: 2026-01-10",
                                    value=""
                                )
                                cs_generate_task_email_btn = gr.Button("📝 업무 할당 메일 생성", variant="secondary")

                            with gr.Column(scale=1):
                                cs_task_email_content = gr.Textbox(
                                    label="업무 할당 이메일 (수정 가능)",
                                    lines=8,
                                    interactive=True
                                )
                                cs_send_task_email_btn = gr.Button("📤 업무 할당 메일 전송", variant="primary")
                                cs_task_email_result = gr.Markdown()

                        gr.Markdown("---")
                        gr.Markdown("#### 📁 1차 산출물 등록")

                        with gr.Row():
                            with gr.Column(scale=1):
                                cs_output_type = gr.Dropdown(
                                    choices=["분석 보고서", "원인 조사서", "공정 데이터", "검사 결과", "미팅 회의록", "기타"],
                                    label="산출물 유형",
                                    value="분석 보고서"
                                )
                                cs_output_description = gr.Textbox(
                                    label="산출물 설명",
                                    placeholder="산출물에 대한 설명을 입력하세요",
                                    lines=2
                                )
                                cs_register_output_btn = gr.Button("📁 산출물 등록", variant="secondary")
                                cs_output_result = gr.Markdown()

                            with gr.Column(scale=1):
                                gr.Markdown("**등록된 산출물 요약**")
                                cs_outputs_summary = gr.Textbox(
                                    label="산출물 목록",
                                    lines=5,
                                    interactive=True,
                                    placeholder="등록된 산출물이 여기에 표시됩니다..."
                                )

                        gr.Markdown("---")
                        gr.Markdown("#### 📄 1차 산출물 보고서 생성")

                        with gr.Row():
                            with gr.Column(scale=1):
                                cs_generate_output_report_btn = gr.Button("📄 1차 산출물 보고서 생성 (Word)", variant="primary", size="lg")
                            with gr.Column(scale=1):
                                cs_output_report_path = gr.Textbox(label="보고서 경로", interactive=False)

                        cs_output_report_result = gr.Markdown()

                        gr.Markdown("---")
                        gr.Markdown("#### 📋 1차 분석 결과 결재")

                        with gr.Row():
                            with gr.Column(scale=1):
                                cs_first_approval_btn = gr.Button("📋 1차 분석 결재 요청", variant="secondary", size="lg")
                                cs_first_approval_id = gr.Textbox(label="결재번호", interactive=False)

                            with gr.Column(scale=1):
                                cs_first_approval_status = gr.Markdown(value="1차 분석 완료 후 결재를 요청하세요.")

                        with gr.Row():
                            with gr.Column(scale=1):
                                cs_first_approval_comment = gr.Textbox(
                                    label="결재 의견",
                                    placeholder="승인/반려 의견을 입력하세요",
                                    lines=2
                                )
                            with gr.Column(scale=1):
                                with gr.Row():
                                    cs_first_approve_btn = gr.Button("✅ 승인", variant="primary")
                                    cs_first_reject_btn = gr.Button("❌ 반려", variant="stop")
                                cs_first_approval_result = gr.Markdown()

                    # ===== 페이지 3: 2차 분석 =====
                    with gr.TabItem("Step 3. 2차 분석", id="cs_step3"):
                        gr.Markdown("""
                        ## 🔬 2차 상세 분석 (귀책부서)
                        귀책 부서에서 상세 원인을 분석하고 개선 대책을 수립합니다.
                        """)

                        with gr.Row():
                            with gr.Column(scale=1):
                                gr.Markdown("#### 분석 대상")
                                cs_second_complaint_id = gr.Textbox(
                                    label="불만 접수 ID",
                                    placeholder="Step 1에서 생성된 ID를 입력하세요"
                                )
                                cs_second_first_analysis_id = gr.Textbox(
                                    label="1차 분석 ID",
                                    placeholder="Step 2에서 생성된 ID를 입력하세요"
                                )
                                cs_second_analysis_btn = gr.Button("2차 분석 실행", variant="primary", size="lg")

                            with gr.Column(scale=1):
                                gr.Markdown("#### 분석 정보")
                                gr.Markdown("""
                                **2차 분석 내용:**
                                - 상세 원인 분석
                                - 재발 방지 대책
                                - 공정 개선 방안
                                - 품질 관리 강화 방안
                                """)

                        gr.Markdown("---")
                        gr.Markdown("#### 2차 분석 결과")
                        cs_second_analysis_id_display = gr.Textbox(label="2차 분석 ID", interactive=False)
                        cs_second_result = gr.Markdown(value="불만 접수 ID와 1차 분석 ID를 입력하고 2차 분석을 실행하세요.")

                        gr.Markdown("---")
                        gr.Markdown("#### 📁 2차 산출물 등록")

                        with gr.Row():
                            with gr.Column(scale=1):
                                cs_second_output_type = gr.Dropdown(
                                    choices=["상세 원인 분석서", "재발 방지 대책서", "공정 개선 계획서", "품질 관리 방안", "시정 조치서", "기타"],
                                    label="산출물 유형",
                                    value="상세 원인 분석서"
                                )
                                cs_second_output_description = gr.Textbox(
                                    label="산출물 설명",
                                    placeholder="산출물에 대한 설명을 입력하세요",
                                    lines=2
                                )
                                cs_second_register_output_btn = gr.Button("📁 2차 산출물 등록", variant="secondary")
                                cs_second_output_result = gr.Markdown()

                            with gr.Column(scale=1):
                                gr.Markdown("**등록된 2차 산출물 요약**")
                                cs_second_outputs_summary = gr.Textbox(
                                    label="2차 산출물 목록",
                                    lines=5,
                                    interactive=True,
                                    placeholder="등록된 2차 산출물이 여기에 표시됩니다..."
                                )

                        gr.Markdown("---")
                        gr.Markdown("#### 📄 2차 산출물 보고서 생성")

                        with gr.Row():
                            with gr.Column(scale=1):
                                cs_generate_second_output_report_btn = gr.Button("📄 2차 산출물 보고서 생성 (Word)", variant="primary", size="lg")
                            with gr.Column(scale=1):
                                cs_second_output_report_path = gr.Textbox(label="보고서 경로", interactive=False)

                        cs_second_output_report_result = gr.Markdown()

                    # ===== 페이지 4: 보고서 생성 =====
                    with gr.TabItem("Step 4. 보고서 생성", id="cs_step4"):
                        gr.Markdown("""
                        ## 📄 최종 보고서 생성
                        분석 결과를 종합하여 Word 형식의 최종 보고서를 생성합니다.
                        """)

                        with gr.Row():
                            with gr.Column(scale=1):
                                gr.Markdown("#### 보고서 생성 대상")
                                cs_report_complaint_id = gr.Textbox(
                                    label="불만 접수 ID",
                                    placeholder="Step 1에서 생성된 ID를 입력하세요"
                                )
                                cs_report_defect_type = gr.Dropdown(
                                    choices=["DEAD_PIXEL", "BRIGHT_SPOT", "LINE_DEFECT", "MURA", "SCRATCH", "TOUCH_FAIL"],
                                    label="결함 유형",
                                    value="DEAD_PIXEL"
                                )
                                cs_report_customer = gr.Dropdown(
                                    choices=["APPLE", "SAMSUNG_MOBILE", "SAMSUNG_VD", "LG", "XIAOMI", "HUAWEI", "GOOGLE", "SONY", "HONDA", "TOYOTA"],
                                    label="고객사",
                                    value="APPLE"
                                )
                                cs_report_btn = gr.Button("최종 보고서 생성 (Word)", variant="primary", size="lg")

                            with gr.Column(scale=1):
                                gr.Markdown("#### 보고서 정보")
                                gr.Markdown("""
                                **보고서 포함 내용:**
                                - 불만 접수 정보
                                - 1차 분석 결과
                                - 2차 분석 결과
                                - 개선 대책 및 향후 계획
                                """)

                        gr.Markdown("---")
                        gr.Markdown("#### 생성 결과")
                        cs_report_path = gr.Textbox(label="보고서 저장 경로", interactive=False)
                        cs_report_result = gr.Markdown(value="불만 접수 ID를 입력하고 보고서를 생성하세요.")

                        with gr.Row():
                            cs_view_report_btn = gr.Button("📥 보고서 다운로드", variant="secondary", size="lg")
                            cs_report_file = gr.File(label="보고서 파일", visible=True)

                        gr.Markdown("---")
                        gr.Markdown("#### 📋 최종 보고서 결재")

                        with gr.Row():
                            with gr.Column(scale=1):
                                cs_final_approval_btn = gr.Button("📋 최종 보고서 결재 요청", variant="secondary", size="lg")
                                cs_final_approval_id = gr.Textbox(label="결재번호", interactive=False)

                            with gr.Column(scale=1):
                                cs_final_approval_status = gr.Markdown(value="보고서 생성 후 결재를 요청하세요.")

                        with gr.Row():
                            with gr.Column(scale=1):
                                cs_final_approval_comment = gr.Textbox(
                                    label="결재 의견",
                                    placeholder="승인/반려 의견을 입력하세요",
                                    lines=2
                                )
                            with gr.Column(scale=1):
                                with gr.Row():
                                    cs_final_approve_btn = gr.Button("✅ 승인", variant="primary")
                                    cs_final_reject_btn = gr.Button("❌ 반려", variant="stop")
                                cs_final_approval_result = gr.Markdown()

                        gr.Markdown("---")
                        gr.Markdown("#### 고객 회신 메일")

                        with gr.Row():
                            with gr.Column(scale=1):
                                cs_reply_generate_btn = gr.Button("고객 회신 메일 생성", variant="secondary", size="lg")
                                gr.Markdown("**수신자 정보**")
                                cs_reply_recipient_name = gr.Textbox(label="수신자 이름", interactive=False)
                                cs_reply_recipient_email = gr.Textbox(label="수신자 이메일", interactive=False)

                                cs_reply_language = gr.Dropdown(
                                    choices=["한국어", "English", "日本語", "中文"],
                                    label="메일 언어",
                                    value="한국어"
                                )
                                cs_reply_translate_btn = gr.Button("언어 변환", variant="secondary")

                            with gr.Column(scale=2):
                                cs_reply_email_content = gr.Textbox(
                                    label="회신 메일 내용 (수정 가능)",
                                    lines=12,
                                    interactive=True,
                                    placeholder="보고서 생성 후 '고객 회신 메일 생성' 버튼을 클릭하세요..."
                                )

                        gr.Markdown("---")
                        gr.Markdown("#### 📋 고객 회신 메일 결재")

                        with gr.Row():
                            with gr.Column(scale=1):
                                cs_reply_approval_btn = gr.Button("📋 회신 메일 결재 요청", variant="secondary")
                                cs_reply_approval_id = gr.Textbox(label="결재번호", interactive=False)
                            with gr.Column(scale=1):
                                cs_reply_approval_status = gr.Markdown(value="회신 메일 생성 후 결재를 요청하세요.")

                        with gr.Row():
                            with gr.Column(scale=1):
                                cs_reply_approval_comment = gr.Textbox(
                                    label="결재 의견",
                                    placeholder="승인/반려 의견을 입력하세요",
                                    lines=2
                                )
                            with gr.Column(scale=1):
                                with gr.Row():
                                    cs_reply_approve_btn = gr.Button("✅ 승인", variant="primary")
                                    cs_reply_reject_btn = gr.Button("❌ 반려", variant="stop")
                                cs_reply_approval_result = gr.Markdown()

                        with gr.Row():
                            cs_reply_send_btn = gr.Button("📤 고객 회신 메일 전송 (결재완료 후)", variant="primary", size="lg")
                        cs_reply_result = gr.Markdown()

                    # ===== 페이지 5: 결재 현황 대시보드 =====
                    with gr.TabItem("결재 현황", id="cs_approval_dashboard"):
                        gr.Markdown("""
                        ## 📊 결재 현황 대시보드
                        모든 결재 요청의 현황을 확인하고 관리합니다.
                        """)

                        with gr.Row():
                            cs_dashboard_refresh_btn = gr.Button("🔄 새로고침", variant="secondary", size="lg")

                        with gr.Row():
                            with gr.Column(scale=1):
                                # 요약 통계
                                cs_dashboard_summary = gr.Markdown(value="새로고침 버튼을 클릭하여 결재 현황을 조회하세요.")

                            with gr.Column(scale=1):
                                # 기한 현황
                                cs_deadline_summary = gr.Markdown(value="")

                        gr.Markdown("---")

                        # 기한 알림 섹션
                        cs_deadline_alerts = gr.Markdown(value="")

                        gr.Markdown("---")

                        with gr.Row():
                            with gr.Column(scale=1):
                                # 대기 중인 결재
                                cs_pending_list = gr.Markdown(value="")

                            with gr.Column(scale=1):
                                # 최근 처리된 결재
                                cs_recent_list = gr.Markdown(value="")

                        gr.Markdown("---")
                        gr.Markdown("#### 🔍 결재 검색")

                        with gr.Row():
                            with gr.Column(scale=1):
                                cs_search_type = gr.Dropdown(
                                    choices=["결재번호", "접수ID", "상태", "결재유형"],
                                    label="검색 조건",
                                    value="접수ID"
                                )
                                cs_search_value = gr.Textbox(
                                    label="검색어",
                                    placeholder="검색어를 입력하세요"
                                )
                                cs_search_btn = gr.Button("검색", variant="primary")

                            with gr.Column(scale=2):
                                cs_search_result = gr.Markdown(value="검색 조건을 입력하고 검색 버튼을 클릭하세요.")

                        gr.Markdown("---")
                        gr.Markdown("#### 📄 결재 상세 조회")

                        with gr.Row():
                            with gr.Column(scale=1):
                                cs_detail_approval_id = gr.Textbox(
                                    label="결재번호",
                                    placeholder="상세 조회할 결재번호를 입력하세요"
                                )
                                cs_detail_btn = gr.Button("상세 조회", variant="secondary")

                            with gr.Column(scale=2):
                                cs_detail_result = gr.Markdown(value="결재번호를 입력하고 상세 조회 버튼을 클릭하세요.")

                        gr.Markdown("---")
                        gr.Markdown("#### ⚡ 빠른 결재 처리")

                        with gr.Row():
                            with gr.Column(scale=1):
                                cs_quick_approval_id = gr.Textbox(
                                    label="결재번호",
                                    placeholder="처리할 결재번호"
                                )
                                cs_quick_comment = gr.Textbox(
                                    label="결재 의견",
                                    placeholder="승인/반려 의견",
                                    lines=2
                                )
                                with gr.Row():
                                    cs_quick_approve_btn = gr.Button("✅ 승인", variant="primary")
                                    cs_quick_reject_btn = gr.Button("❌ 반려", variant="stop")
                                cs_quick_result = gr.Markdown()

                            with gr.Column(scale=2):
                                cs_quick_status = gr.Markdown(value="결재번호를 입력하면 결재 상태가 표시됩니다.")

                        gr.Markdown("---")
                        gr.Markdown("#### 📧 기한 초과 알림 발송")

                        with gr.Row():
                            with gr.Column(scale=1):
                                gr.Markdown("기한이 초과되었거나 임박한 결재에 대해 담당자에게 알림 이메일을 발송합니다.")
                                cs_send_all_notifications_btn = gr.Button("📧 전체 알림 발송", variant="primary")
                                cs_notification_result = gr.Markdown()

                            with gr.Column(scale=2):
                                gr.Markdown("##### 개별 알림 발송")
                                with gr.Row():
                                    cs_notification_approval_id = gr.Textbox(
                                        label="결재번호",
                                        placeholder="알림 발송할 결재번호"
                                    )
                                    cs_notification_type = gr.Dropdown(
                                        choices=["overdue", "urgent", "warning"],
                                        label="알림 유형",
                                        value="overdue"
                                    )
                                cs_send_single_notification_btn = gr.Button("📧 개별 알림 발송", variant="secondary")
                                cs_single_notification_result = gr.Markdown()

                        gr.Markdown("---")
                        gr.Markdown("##### 📜 알림 발송 이력")
                        with gr.Row():
                            cs_history_approval_id = gr.Textbox(
                                label="결재번호",
                                placeholder="이력 조회할 결재번호"
                            )
                            cs_history_btn = gr.Button("이력 조회", variant="secondary")
                        cs_notification_history_result = gr.Markdown()

                # ===== 이벤트 연결 =====
                # Step 1: 불만 접수
                cs_init_btn.click(
                    fn=init_cs_workflow,
                    outputs=[cs_init_result]
                )

                cs_list_btn.click(
                    fn=get_complaints_list,
                    outputs=[cs_list_result]
                )

                def create_and_generate_email(customer, product, lot_id, cell_id, defect_type, description, severity):
                    # 불만 접수 생성
                    result, complaint_id = create_cs_complaint(customer, product, lot_id, cell_id, defect_type, description, severity)

                    if complaint_id:
                        # 이메일 자동 생성
                        email_content, recipient_email, recipient_name = generate_complaint_email(
                            customer, product, lot_id, cell_id, defect_type, description, severity, complaint_id
                        )
                        return (result, complaint_id, complaint_id, complaint_id, complaint_id,
                                email_content, recipient_email, recipient_name)
                    else:
                        return (result, "", "", "", "", "", "", "")

                cs_create_btn.click(
                    fn=create_and_generate_email,
                    inputs=[cs_customer, cs_product, cs_lot_id, cs_cell_id,
                            cs_defect_type, cs_description, cs_severity],
                    outputs=[cs_create_result, cs_complaint_id_display,
                             cs_first_complaint_id, cs_second_complaint_id, cs_report_complaint_id,
                             cs_email_content, cs_recipient_email, cs_recipient_name]
                )

                # 이메일 언어 번역
                def translate_and_update(email_content, target_language):
                    if not email_content.strip():
                        return "이메일 내용이 없습니다."
                    translated = translate_email(email_content, target_language)
                    return translated

                cs_translate_btn.click(
                    fn=translate_and_update,
                    inputs=[cs_email_content, cs_email_language],
                    outputs=[cs_email_content]
                )

                # Step 1: 고객 이메일 결재 요청
                def create_email_approval_handler(complaint_id, email_content):
                    if not complaint_id.strip():
                        return "", "결재 요청 실패: 불만 접수 ID가 없습니다.", "결재 요청을 먼저 생성하세요."
                    if not email_content.strip():
                        return "", "결재 요청 실패: 이메일 내용이 없습니다.", "결재 요청을 먼저 생성하세요."

                    approval_id, approval = create_approval_request(
                        "COMPLAINT_EMAIL",
                        complaint_id,
                        f"고객 확인 이메일 발송 승인 요청 - {complaint_id}",
                        email_content
                    )
                    return approval_id, f"결재 요청이 생성되었습니다. 결재번호: {approval_id}", format_approval_status(approval)

                cs_email_approval_btn.click(
                    fn=create_email_approval_handler,
                    inputs=[cs_complaint_id_display, cs_email_content],
                    outputs=[cs_email_approval_id, cs_email_approval_result, cs_email_approval_status]
                )

                # Step 1: 고객 이메일 결재 승인
                def approve_email_handler(approval_id, comment):
                    if not approval_id.strip():
                        return "결재번호가 없습니다.", format_approval_status(None)
                    success, message = process_approval(approval_id, "APPROVE", comment)
                    approval = get_approval_status(approval_id)
                    return message, format_approval_status(approval)

                cs_email_approve_btn.click(
                    fn=approve_email_handler,
                    inputs=[cs_email_approval_id, cs_email_approval_comment],
                    outputs=[cs_email_approval_result, cs_email_approval_status]
                )

                # Step 1: 고객 이메일 결재 반려
                def reject_email_handler(approval_id, comment):
                    if not approval_id.strip():
                        return "결재번호가 없습니다.", format_approval_status(None)
                    if not comment.strip():
                        return "반려 사유를 입력하세요.", format_approval_status(get_approval_status(approval_id))
                    success, message = process_approval(approval_id, "REJECT", comment)
                    approval = get_approval_status(approval_id)
                    return message, format_approval_status(approval)

                cs_email_reject_btn.click(
                    fn=reject_email_handler,
                    inputs=[cs_email_approval_id, cs_email_approval_comment],
                    outputs=[cs_email_approval_result, cs_email_approval_status]
                )

                # 이메일 전송 (결재 완료 후)
                def send_email_handler(recipient_email, recipient_name, email_content, complaint_id, approval_id):
                    if not email_content.strip():
                        return "전송할 이메일 내용이 없습니다."
                    if not complaint_id.strip():
                        return "불만 접수 ID가 없습니다."
                    if not approval_id.strip():
                        return "결재를 먼저 요청하세요."

                    # 결재 상태 확인
                    approval = get_approval_status(approval_id)
                    if not approval:
                        return "결재 정보를 찾을 수 없습니다."
                    if approval["status"] != "APPROVED":
                        return f"결재가 완료되지 않았습니다. (현재 상태: {APPROVAL_STATUS.get(approval['status'], approval['status'])})"

                    return send_complaint_email(recipient_email, recipient_name, email_content, complaint_id)

                cs_send_email_btn.click(
                    fn=send_email_handler,
                    inputs=[cs_recipient_email, cs_recipient_name, cs_email_content, cs_complaint_id_display, cs_email_approval_id],
                    outputs=[cs_email_result]
                )

                # 미팅 요청 이메일 생성
                def generate_meeting_email_handler(customer, product, defect_type, description, severity, complaint_id):
                    if not complaint_id.strip():
                        return "불만 접수 ID가 없습니다. 먼저 불만을 접수하세요.", "", "", ""
                    email_content, attendees, responsible_dept, meeting_time = generate_meeting_request_email(
                        customer, product, defect_type, description, severity, complaint_id
                    )
                    return email_content, attendees, responsible_dept, meeting_time

                cs_generate_meeting_btn.click(
                    fn=generate_meeting_email_handler,
                    inputs=[cs_customer, cs_product, cs_defect_type, cs_description, cs_severity, cs_complaint_id_display],
                    outputs=[cs_meeting_email_content, cs_meeting_attendees, cs_meeting_responsible_dept, cs_meeting_time]
                )

                # 미팅 요청 이메일 전송
                def send_meeting_handler(attendees, email_content, complaint_id, meeting_time):
                    if not email_content.strip():
                        return "미팅 요청 이메일 내용이 없습니다. 먼저 '미팅 요청 메일 생성' 버튼을 클릭하세요."
                    if not complaint_id.strip():
                        return "불만 접수 ID가 없습니다."
                    return send_meeting_request_email(attendees, email_content, complaint_id, meeting_time)

                cs_send_meeting_btn.click(
                    fn=send_meeting_handler,
                    inputs=[cs_meeting_attendees, cs_meeting_email_content, cs_complaint_id_display, cs_meeting_time],
                    outputs=[cs_meeting_result]
                )

                # Step 2: 1차 분석
                def first_analysis_and_update(complaint_id):
                    result, analysis_id = perform_first_analysis(complaint_id)
                    return result, analysis_id, analysis_id

                cs_first_analysis_btn.click(
                    fn=first_analysis_and_update,
                    inputs=[cs_first_complaint_id],
                    outputs=[cs_first_result, cs_first_analysis_id_display, cs_second_first_analysis_id]
                )

                # Step 2: 미팅 결과 정리
                def generate_summary_handler(complaint_id, first_result, meeting_notes):
                    if not complaint_id.strip():
                        return "불만 접수 ID가 없습니다."
                    if not meeting_notes.strip():
                        return "미팅 노트를 입력하세요."
                    return generate_meeting_summary(complaint_id, first_result, meeting_notes)

                cs_generate_summary_btn.click(
                    fn=generate_summary_handler,
                    inputs=[cs_first_complaint_id, cs_first_result, cs_meeting_notes],
                    outputs=[cs_meeting_summary]
                )

                # Step 2: 업무 할당 이메일 생성
                def generate_task_email_handler(complaint_id, dept, task_desc, deadline):
                    if not complaint_id.strip():
                        return "불만 접수 ID가 없습니다."
                    tasks_data = f"담당부서: {dept}\n업무내용: {task_desc}\n납기일: {deadline}"
                    return generate_task_assignment_email(complaint_id, tasks_data)

                cs_generate_task_email_btn.click(
                    fn=generate_task_email_handler,
                    inputs=[cs_first_complaint_id, cs_task_dept, cs_task_description, cs_task_deadline],
                    outputs=[cs_task_email_content]
                )

                # Step 2: 업무 할당 이메일 전송
                def send_task_email_handler(dept, email_content, complaint_id, task_desc, deadline):
                    if not email_content.strip():
                        return "업무 할당 이메일 내용이 없습니다."
                    dept_info = INTERNAL_DEPARTMENTS.get(dept, {})
                    recipient_email = dept_info.get("email", "unknown@sdc.com")
                    return send_task_assignment_email(dept, recipient_email, email_content, complaint_id, task_desc, deadline)

                cs_send_task_email_btn.click(
                    fn=send_task_email_handler,
                    inputs=[cs_task_dept, cs_task_email_content, cs_first_complaint_id, cs_task_description, cs_task_deadline],
                    outputs=[cs_task_email_result]
                )

                # Step 2: 산출물 등록
                outputs_list = []
                def register_output_handler(complaint_id, output_type, output_desc, current_summary):
                    if not complaint_id.strip():
                        return "불만 접수 ID가 없습니다.", current_summary
                    result = register_first_output(complaint_id, output_type, output_desc)
                    new_entry = f"• {output_type}: {output_desc}"
                    updated_summary = (current_summary + "\n" + new_entry) if current_summary.strip() else new_entry
                    return result, updated_summary

                cs_register_output_btn.click(
                    fn=register_output_handler,
                    inputs=[cs_first_complaint_id, cs_output_type, cs_output_description, cs_outputs_summary],
                    outputs=[cs_output_result, cs_outputs_summary]
                )

                # Step 2: 1차 산출물 보고서 생성
                def generate_output_report_handler(complaint_id, meeting_summary, outputs_summary):
                    if not complaint_id.strip():
                        return "불만 접수 ID가 없습니다.", ""
                    tasks_summary = "업무 할당 이력은 이메일 로그를 참조하세요."
                    return generate_first_output_report(complaint_id, meeting_summary, tasks_summary, outputs_summary)

                cs_generate_output_report_btn.click(
                    fn=generate_output_report_handler,
                    inputs=[cs_first_complaint_id, cs_meeting_summary, cs_outputs_summary],
                    outputs=[cs_output_report_result, cs_output_report_path]
                )

                # Step 2: 1차 분석 결재 요청
                def create_first_approval_handler(complaint_id, first_result, meeting_summary, outputs_summary, report_path):
                    if not complaint_id.strip():
                        return "", "결재 요청 실패: 불만 접수 ID가 없습니다.", "결재를 요청하려면 먼저 1차 분석을 완료하세요."
                    if not first_result.strip() or "분석 결과" not in first_result:
                        return "", "결재 요청 실패: 1차 분석을 먼저 수행하세요.", "결재를 요청하려면 먼저 1차 분석을 완료하세요."

                    content = f"""1차 분석 결과 결재 요청

[분석 결과]
{first_result}

[미팅 결과]
{meeting_summary if meeting_summary else '(미팅 기록 없음)'}

[산출물]
{outputs_summary if outputs_summary else '(등록된 산출물 없음)'}

[보고서 경로]
{report_path if report_path else '(보고서 미생성)'}
"""
                    approval_id, approval = create_approval_request(
                        "FIRST_ANALYSIS",
                        complaint_id,
                        f"1차 분석 결과 승인 요청 - {complaint_id}",
                        content
                    )
                    return approval_id, f"결재 요청이 생성되었습니다. 결재번호: {approval_id}", format_approval_status(approval)

                cs_first_approval_btn.click(
                    fn=create_first_approval_handler,
                    inputs=[cs_first_complaint_id, cs_first_result, cs_meeting_summary, cs_outputs_summary, cs_output_report_path],
                    outputs=[cs_first_approval_id, cs_first_approval_result, cs_first_approval_status]
                )

                # Step 2: 1차 분석 결재 승인
                def approve_first_handler(approval_id, comment):
                    if not approval_id.strip():
                        return "결재번호가 없습니다.", format_approval_status(None)
                    success, message = process_approval(approval_id, "APPROVE", comment)
                    approval = get_approval_status(approval_id)
                    return message, format_approval_status(approval)

                cs_first_approve_btn.click(
                    fn=approve_first_handler,
                    inputs=[cs_first_approval_id, cs_first_approval_comment],
                    outputs=[cs_first_approval_result, cs_first_approval_status]
                )

                # Step 2: 1차 분석 결재 반려
                def reject_first_handler(approval_id, comment):
                    if not approval_id.strip():
                        return "결재번호가 없습니다.", format_approval_status(None)
                    if not comment.strip():
                        return "반려 사유를 입력하세요.", format_approval_status(get_approval_status(approval_id))
                    success, message = process_approval(approval_id, "REJECT", comment)
                    approval = get_approval_status(approval_id)
                    return message, format_approval_status(approval)

                cs_first_reject_btn.click(
                    fn=reject_first_handler,
                    inputs=[cs_first_approval_id, cs_first_approval_comment],
                    outputs=[cs_first_approval_result, cs_first_approval_status]
                )

                # Step 3: 2차 분석
                def second_analysis_and_update(complaint_id, first_analysis_id):
                    result, analysis_id = perform_second_analysis(complaint_id, first_analysis_id)
                    return result, analysis_id

                cs_second_analysis_btn.click(
                    fn=second_analysis_and_update,
                    inputs=[cs_second_complaint_id, cs_second_first_analysis_id],
                    outputs=[cs_second_result, cs_second_analysis_id_display]
                )

                # Step 3: 2차 산출물 등록
                def register_second_output_handler(complaint_id, output_type, output_desc, current_summary):
                    if not complaint_id.strip():
                        return "불만 접수 ID가 없습니다.", current_summary
                    result = register_second_output(complaint_id, output_type, output_desc)
                    new_entry = f"• {output_type}: {output_desc}"
                    updated_summary = (current_summary + "\n" + new_entry) if current_summary.strip() else new_entry
                    return result, updated_summary

                cs_second_register_output_btn.click(
                    fn=register_second_output_handler,
                    inputs=[cs_second_complaint_id, cs_second_output_type, cs_second_output_description, cs_second_outputs_summary],
                    outputs=[cs_second_output_result, cs_second_outputs_summary]
                )

                # Step 3: 2차 산출물 보고서 생성
                def generate_second_output_report_handler(complaint_id, second_result, outputs_summary):
                    if not complaint_id.strip():
                        return "불만 접수 ID가 없습니다.", ""
                    return generate_second_output_report(complaint_id, second_result, outputs_summary)

                cs_generate_second_output_report_btn.click(
                    fn=generate_second_output_report_handler,
                    inputs=[cs_second_complaint_id, cs_second_result, cs_second_outputs_summary],
                    outputs=[cs_second_output_report_result, cs_second_output_report_path]
                )

                # Step 4: 보고서 생성
                cs_report_btn.click(
                    fn=generate_final_report,
                    inputs=[cs_report_complaint_id],
                    outputs=[cs_report_result, cs_report_path]
                )

                # Step 4: 보고서 다운로드
                def download_report(report_path):
                    import os
                    if not report_path or not report_path.strip():
                        return None
                    if os.path.exists(report_path):
                        return report_path
                    return None

                cs_view_report_btn.click(
                    fn=download_report,
                    inputs=[cs_report_path],
                    outputs=[cs_report_file]
                )

                # Step 4: 최종 보고서 결재 요청
                def create_final_approval_handler(complaint_id, report_path, report_result):
                    if not complaint_id.strip():
                        return "", "결재 요청 실패: 불만 접수 ID가 없습니다.", "보고서 생성 후 결재를 요청하세요."
                    if not report_path.strip():
                        return "", "결재 요청 실패: 먼저 보고서를 생성하세요.", "보고서 생성 후 결재를 요청하세요."

                    content = f"""최종 보고서 결재 요청

[불만 접수 ID]
{complaint_id}

[보고서 경로]
{report_path}

[보고서 정보]
{report_result}
"""
                    approval_id, approval = create_approval_request(
                        "FINAL_REPORT",
                        complaint_id,
                        f"최종 보고서 승인 요청 - {complaint_id}",
                        content
                    )
                    return approval_id, f"결재 요청이 생성되었습니다. 결재번호: {approval_id}", format_approval_status(approval)

                cs_final_approval_btn.click(
                    fn=create_final_approval_handler,
                    inputs=[cs_report_complaint_id, cs_report_path, cs_report_result],
                    outputs=[cs_final_approval_id, cs_final_approval_result, cs_final_approval_status]
                )

                # Step 4: 최종 보고서 결재 승인
                def approve_final_handler(approval_id, comment):
                    if not approval_id.strip():
                        return "결재번호가 없습니다.", format_approval_status(None)
                    success, message = process_approval(approval_id, "APPROVE", comment)
                    approval = get_approval_status(approval_id)
                    return message, format_approval_status(approval)

                cs_final_approve_btn.click(
                    fn=approve_final_handler,
                    inputs=[cs_final_approval_id, cs_final_approval_comment],
                    outputs=[cs_final_approval_result, cs_final_approval_status]
                )

                # Step 4: 최종 보고서 결재 반려
                def reject_final_handler(approval_id, comment):
                    if not approval_id.strip():
                        return "결재번호가 없습니다.", format_approval_status(None)
                    if not comment.strip():
                        return "반려 사유를 입력하세요.", format_approval_status(get_approval_status(approval_id))
                    success, message = process_approval(approval_id, "REJECT", comment)
                    approval = get_approval_status(approval_id)
                    return message, format_approval_status(approval)

                cs_final_reject_btn.click(
                    fn=reject_final_handler,
                    inputs=[cs_final_approval_id, cs_final_approval_comment],
                    outputs=[cs_final_approval_result, cs_final_approval_status]
                )

                # Step 4: 고객 회신 메일 생성
                def generate_reply_email_handler(customer, complaint_id, defect_type, report_path):
                    if not complaint_id.strip():
                        return "", "", "", "불만 접수 ID가 없습니다."
                    if not report_path.strip():
                        return "", "", "", "먼저 보고서를 생성하세요."

                    email_content, recipient_email, recipient_name = generate_customer_reply_email(
                        customer, complaint_id, defect_type, report_path
                    )
                    return email_content, recipient_email, recipient_name, ""

                cs_reply_generate_btn.click(
                    fn=generate_reply_email_handler,
                    inputs=[cs_report_customer, cs_report_complaint_id, cs_report_defect_type, cs_report_path],
                    outputs=[cs_reply_email_content, cs_reply_recipient_email, cs_reply_recipient_name, cs_reply_result]
                )

                # Step 4: 고객 회신 메일 언어 번역
                def translate_reply_email_handler(email_content, target_language):
                    if not email_content.strip():
                        return "이메일 내용이 없습니다."
                    translated = translate_email(email_content, target_language)
                    return translated

                cs_reply_translate_btn.click(
                    fn=translate_reply_email_handler,
                    inputs=[cs_reply_email_content, cs_reply_language],
                    outputs=[cs_reply_email_content]
                )

                # Step 4: 고객 회신 메일 결재 요청
                def create_reply_approval_handler(complaint_id, email_content, recipient_email, recipient_name):
                    if not complaint_id.strip():
                        return "", "결재 요청 실패: 불만 접수 ID가 없습니다.", "회신 메일 생성 후 결재를 요청하세요."
                    if not email_content.strip():
                        return "", "결재 요청 실패: 회신 메일을 먼저 생성하세요.", "회신 메일 생성 후 결재를 요청하세요."

                    content = f"""고객 회신 메일 결재 요청

[수신자]
{recipient_name} <{recipient_email}>

[메일 내용]
{email_content}
"""
                    approval_id, approval = create_approval_request(
                        "CUSTOMER_REPLY",
                        complaint_id,
                        f"고객 회신 메일 승인 요청 - {complaint_id}",
                        content
                    )
                    return approval_id, f"결재 요청이 생성되었습니다. 결재번호: {approval_id}", format_approval_status(approval)

                cs_reply_approval_btn.click(
                    fn=create_reply_approval_handler,
                    inputs=[cs_report_complaint_id, cs_reply_email_content, cs_reply_recipient_email, cs_reply_recipient_name],
                    outputs=[cs_reply_approval_id, cs_reply_approval_result, cs_reply_approval_status]
                )

                # Step 4: 고객 회신 메일 결재 승인
                def approve_reply_handler(approval_id, comment):
                    if not approval_id.strip():
                        return "결재번호가 없습니다.", format_approval_status(None)
                    success, message = process_approval(approval_id, "APPROVE", comment)
                    approval = get_approval_status(approval_id)
                    return message, format_approval_status(approval)

                cs_reply_approve_btn.click(
                    fn=approve_reply_handler,
                    inputs=[cs_reply_approval_id, cs_reply_approval_comment],
                    outputs=[cs_reply_approval_result, cs_reply_approval_status]
                )

                # Step 4: 고객 회신 메일 결재 반려
                def reject_reply_handler(approval_id, comment):
                    if not approval_id.strip():
                        return "결재번호가 없습니다.", format_approval_status(None)
                    if not comment.strip():
                        return "반려 사유를 입력하세요.", format_approval_status(get_approval_status(approval_id))
                    success, message = process_approval(approval_id, "REJECT", comment)
                    approval = get_approval_status(approval_id)
                    return message, format_approval_status(approval)

                cs_reply_reject_btn.click(
                    fn=reject_reply_handler,
                    inputs=[cs_reply_approval_id, cs_reply_approval_comment],
                    outputs=[cs_reply_approval_result, cs_reply_approval_status]
                )

                # Step 4: 고객 회신 메일 전송 (결재 완료 후)
                def send_reply_email_handler(recipient_email, recipient_name, email_content, complaint_id, report_path, approval_id):
                    if not email_content.strip():
                        return "전송할 이메일 내용이 없습니다."
                    if not complaint_id.strip():
                        return "불만 접수 ID가 없습니다."
                    if not approval_id.strip():
                        return "결재를 먼저 요청하세요."

                    # 결재 상태 확인
                    approval = get_approval_status(approval_id)
                    if not approval:
                        return "결재 정보를 찾을 수 없습니다."
                    if approval["status"] != "APPROVED":
                        return f"결재가 완료되지 않았습니다. (현재 상태: {APPROVAL_STATUS.get(approval['status'], approval['status'])})"

                    return send_customer_reply_email(recipient_email, recipient_name, email_content, complaint_id, report_path)

                cs_reply_send_btn.click(
                    fn=send_reply_email_handler,
                    inputs=[cs_reply_recipient_email, cs_reply_recipient_name, cs_reply_email_content, cs_report_complaint_id, cs_report_path, cs_reply_approval_id],
                    outputs=[cs_reply_result]
                )

                # ===== 결재 대시보드 이벤트 =====
                # 대시보드 새로고침
                def refresh_dashboard():
                    summary = get_approval_dashboard_summary()
                    deadline_summary_md = get_deadline_summary()
                    deadline_alerts_md = get_deadline_alerts()
                    pending = get_pending_approvals_list()
                    recent = get_recent_approvals_list()
                    return summary, deadline_summary_md, deadline_alerts_md, pending, recent

                cs_dashboard_refresh_btn.click(
                    fn=refresh_dashboard,
                    outputs=[cs_dashboard_summary, cs_deadline_summary, cs_deadline_alerts, cs_pending_list, cs_recent_list]
                )

                # 결재 검색
                def search_approval_handler(search_type, search_value):
                    if not search_value.strip():
                        return "검색어를 입력하세요."
                    return search_approvals(search_type, search_value)

                cs_search_btn.click(
                    fn=search_approval_handler,
                    inputs=[cs_search_type, cs_search_value],
                    outputs=[cs_search_result]
                )

                # 결재 상세 조회 (기한 정보 포함)
                def detail_approval_handler(approval_id):
                    if not approval_id.strip():
                        return "결재번호를 입력하세요."
                    approval = get_approval_status(approval_id)
                    if not approval:
                        return "결재 정보를 찾을 수 없습니다."
                    return format_approval_with_deadline(approval) + f"\n\n---\n#### 결재 내용\n```\n{approval['content'][:1000]}{'...' if len(approval['content']) > 1000 else ''}\n```"

                cs_detail_btn.click(
                    fn=detail_approval_handler,
                    inputs=[cs_detail_approval_id],
                    outputs=[cs_detail_result]
                )

                # 빠른 결재 - 결재번호 입력 시 상태 조회 (기한 정보 포함)
                def quick_status_handler(approval_id):
                    if not approval_id.strip():
                        return "결재번호를 입력하면 결재 상태가 표시됩니다."
                    approval = get_approval_status(approval_id)
                    if not approval:
                        return "결재 정보를 찾을 수 없습니다."
                    return format_approval_with_deadline(approval)

                cs_quick_approval_id.change(
                    fn=quick_status_handler,
                    inputs=[cs_quick_approval_id],
                    outputs=[cs_quick_status]
                )

                # 빠른 결재 승인
                def quick_approve_handler(approval_id, comment):
                    if not approval_id.strip():
                        return "결재번호가 없습니다.", format_approval_with_deadline(None) if None else "결재 정보가 없습니다."
                    success, message = process_approval(approval_id, "APPROVE", comment)
                    approval = get_approval_status(approval_id)
                    return message, format_approval_with_deadline(approval) if approval else "결재 정보가 없습니다."

                cs_quick_approve_btn.click(
                    fn=quick_approve_handler,
                    inputs=[cs_quick_approval_id, cs_quick_comment],
                    outputs=[cs_quick_result, cs_quick_status]
                )

                # 빠른 결재 반려
                def quick_reject_handler(approval_id, comment):
                    if not approval_id.strip():
                        return "결재번호가 없습니다.", "결재 정보가 없습니다."
                    if not comment.strip():
                        approval = get_approval_status(approval_id)
                        return "반려 사유를 입력하세요.", format_approval_with_deadline(approval) if approval else "결재 정보가 없습니다."
                    success, message = process_approval(approval_id, "REJECT", comment)
                    approval = get_approval_status(approval_id)
                    return message, format_approval_with_deadline(approval) if approval else "결재 정보가 없습니다."

                cs_quick_reject_btn.click(
                    fn=quick_reject_handler,
                    inputs=[cs_quick_approval_id, cs_quick_comment],
                    outputs=[cs_quick_result, cs_quick_status]
                )

                # 결재 현황 대시보드: 전체 알림 발송
                def send_all_notifications_handler():
                    sent_count, results = send_all_overdue_notifications()
                    return format_notification_results(sent_count, results)

                cs_send_all_notifications_btn.click(
                    fn=send_all_notifications_handler,
                    outputs=[cs_notification_result]
                )

                # 결재 현황 대시보드: 개별 알림 발송
                def send_single_notification_handler(approval_id, notification_type):
                    if not approval_id.strip():
                        return "결재번호를 입력하세요."
                    approval = get_approval_status(approval_id)
                    if not approval:
                        return f"결재번호 {approval_id}를 찾을 수 없습니다."
                    if approval["status"] != "PENDING":
                        return f"이미 처리된 결재입니다. (상태: {APPROVAL_STATUS.get(approval['status'], approval['status'])})"

                    success = send_deadline_notification_email(approval, notification_type)
                    if success:
                        return f"### ✅ 알림 발송 완료\n\n- **결재번호**: {approval_id}\n- **알림 유형**: {notification_type}\n- **발송 대상**: {approval['current_approver'].get('name', 'N/A')} ({approval['current_approver'].get('email', 'N/A')})"
                    else:
                        return f"### ❌ 알림 발송 실패\n\n결재번호 {approval_id}에 대한 알림 발송에 실패했습니다."

                cs_send_single_notification_btn.click(
                    fn=send_single_notification_handler,
                    inputs=[cs_notification_approval_id, cs_notification_type],
                    outputs=[cs_single_notification_result]
                )

                # 결재 현황 대시보드: 알림 발송 이력 조회
                def get_notification_history_handler(approval_id):
                    if not approval_id.strip():
                        return "결재번호를 입력하세요."
                    history = get_notification_history_for_approval(approval_id)
                    if not history:
                        return f"결재번호 {approval_id}에 대한 알림 발송 이력이 없습니다."

                    result = f"### 📜 알림 발송 이력 - {approval_id}\n\n"
                    result += "| 발송일시 | 알림유형 | 수신자 | 이메일 |\n"
                    result += "|----------|----------|--------|--------|\n"
                    for h in history:
                        sent_time = h.get("sent_at", "N/A")
                        notif_type = h.get("notification_type", "N/A")
                        recipient = h.get("recipient_name", "N/A")
                        email = h.get("recipient_email", "N/A")
                        result += f"| {sent_time} | {notif_type} | {recipient} | {email} |\n"

                    return result

                cs_history_btn.click(
                    fn=get_notification_history_handler,
                    inputs=[cs_history_approval_id],
                    outputs=[cs_notification_history_result]
                )



            # ===== 탭 2: 빅데이터 분석 =====
            with gr.TabItem("2. 빅데이터 분석", id="bigdata"):
                gr.Markdown("### 데이터 레이크 기반 불량 분석 파이프라인")

                with gr.Row():
                    with gr.Column(scale=1):
                        gr.Markdown("#### 불량 케이스 정보 입력")

                        with gr.Group():
                            bd_case_id = gr.Textbox(label="케이스 ID", placeholder="CASE20250103001")
                            bd_cell_id = gr.Textbox(label="셀 ID", placeholder="CELL001")
                            bd_defect_type = gr.Dropdown(
                                label="결함 유형",
                                choices=["dead_pixel", "bright_spot", "line_defect", "mura", "scratch", "particle", "unknown"],
                                value="unknown"
                            )
                            bd_defect_date = gr.Textbox(label="발생일", placeholder="2025-01-03", value="2025-01-03")

                        with gr.Row():
                            bd_customer = gr.Textbox(label="고객사", placeholder="Customer_A", scale=1)
                            bd_severity = gr.Dropdown(
                                label="심각도",
                                choices=["LOW", "MEDIUM", "HIGH", "CRITICAL"],
                                value="MEDIUM",
                                scale=1
                            )

                        bd_description = gr.Textbox(label="불량 설명", placeholder="불량 상세 내용 입력", lines=2)

                        with gr.Row():
                            bd_init_btn = gr.Button("파이프라인 초기화", variant="secondary")
                            bd_run_btn = gr.Button("분석 실행", variant="primary")

                        bd_init_status = gr.Textbox(label="초기화 상태", interactive=False)
                        bd_init_btn.click(init_bigdata_pipeline, outputs=[bd_init_status])

                    with gr.Column(scale=1):
                        gr.Markdown("#### 파이프라인 실행 결과")
                        bd_result_html = gr.HTML(label="")

                        with gr.Accordion("상세 데이터 요약", open=False):
                            bd_data_summary = gr.Textbox(label="추출 데이터", lines=10, interactive=False)
                            bd_summary_btn = gr.Button("데이터 요약 조회", variant="secondary")
                            bd_summary_btn.click(get_extracted_data_summary, outputs=[bd_data_summary])

                        with gr.Accordion("JSON 결과", open=False):
                            bd_raw_json = gr.Textbox(label="Raw JSON", lines=12, interactive=False)

                bd_run_btn.click(
                    run_defect_analysis_pipeline,
                    inputs=[bd_case_id, bd_cell_id, bd_defect_type, bd_defect_date, bd_customer, bd_severity, bd_description],
                    outputs=[bd_result_html, bd_raw_json]
                )

                gr.Markdown("---")
                gr.Markdown("""
                #### 데이터 파이프라인 아키텍처

                ```
                ┌─────────────┐    ┌─────────────┐    ┌─────────────┐    ┌─────────────┐
                │   QMS/YMS   │───▶│  Greenplum  │───▶│   Parquet   │───▶│    Spark    │
                │  (원천시스템) │    │ (데이터레이크)│    │   (변환)    │    │ (데이터마트) │
                └─────────────┘    └─────────────┘    └─────────────┘    └─────────────┘
                                          │                                    │
                                          ▼                                    ▼
                                   ┌─────────────┐                      ┌─────────────┐
                                   │   Oracle    │                      │   분석결과   │
                                   │ (실시간 FDC)│──────────────────────▶│ (데이터마트) │
                                   └─────────────┘                      └─────────────┘
                ```

                **처리 단계:**
                1. **데이터 추출**: 불량 셀의 제품 이력, 개발 이력, 변경점 데이터 추출
                2. **Parquet 변환**: 컬럼 지향 포맷으로 변환 (Snappy 압축)
                3. **S3 업로드**: 데이터 레이크에 저장
                4. **데이터마트 구성**: Spark로 분석용 데이터 통합
                """)

            # ===== 탭 3: 결함이미지분석 =====
            with gr.TabItem("3. 결함이미지분석", id="analysis"):
                gr.Markdown("### 이미지 기반 결함 분석 (VLM 채팅)")

                # 이미지 이름 상태 저장용
                current_image_name = gr.State(value=None)

                with gr.Row(equal_height=True, elem_id="analysis-main-row"):
                    # 왼쪽 사이드바: 채팅 히스토리
                    with gr.Column(scale=1, min_width=180, elem_id="history-sidebar"):
                        gr.Markdown("#### 채팅 히스토리")
                        # 스크롤 가능한 라디오 리스트
                        with gr.Column(elem_id="history-list-container"):
                            history_dropdown = gr.Radio(
                                choices=get_chat_history_list(),
                                label="저장된 대화 목록",
                                interactive=True,
                                elem_id="history-radio-list",
                            )
                        with gr.Row():
                            load_history_btn = gr.Button("불러오기", size="sm", variant="secondary")
                            delete_history_btn = gr.Button("삭제", size="sm", variant="stop")
                        history_status = gr.Markdown("", elem_classes=["info-text"])

                    # 중앙: 이미지 영역
                    with gr.Column(scale=2):
                        # 샘플 이미지 (맨 위)
                        if SAMPLE_IMAGES:
                            gr.Markdown(f"#### 분석할 결함 이미지 선택 ({len(SAMPLE_IMAGES)}개)")
                            gr.Markdown("*아래 이미지를 클릭하면 결함 위치가 자동 표시되고, VLM 채팅 분석 대상이 됩니다.*")
                            sample_gallery = gr.Gallery(
                                value=[(str(img), img.name) for img in SAMPLE_IMAGES],
                                columns=10,
                                rows=5,
                                height=280,
                                object_fit="cover",
                                show_label=False,
                                allow_preview=False,
                            )

                        gr.Markdown("---")

                        # 결함 좌표 시각화 (클릭시 팝업 확대)
                        gr.Markdown("#### 결함 좌표 시각화 (이미지 클릭하여 확대)")
                        visualized_image = gr.Image(
                            type="pil",
                            label="결함 위치가 표시된 이미지 (클릭하여 확대)",
                            height=350,
                            interactive=False,
                            elem_id="defect-visualization",
                        )
                        visualization_info = gr.Markdown(value="위에서 샘플 이미지를 클릭하면 결함 위치가 자동으로 표시됩니다.")
                        selected_image_info = gr.Markdown(value="", elem_classes=["info-box"])

                        # JavaScript 라이트박스 (이미지 클릭시 팝업)
                        gr.HTML("""
                        <div id="lightbox-overlay" style="display:none; position:fixed; top:0; left:0; width:100vw; height:100vh; background:rgba(0,0,0,0.9); z-index:9999; justify-content:center; align-items:center; cursor:zoom-out;">
                            <img id="lightbox-img" src="" style="max-width:90vw; max-height:90vh; object-fit:contain; border-radius:8px; box-shadow:0 0 30px rgba(255,255,255,0.3);">
                            <div style="position:absolute; top:20px; right:30px; color:white; font-size:40px; cursor:pointer; background:rgba(0,0,0,0.5); width:50px; height:50px; border-radius:50%; display:flex; justify-content:center; align-items:center;" onclick="document.getElementById('lightbox-overlay').style.display='none';">&times;</div>
                            <div style="position:absolute; bottom:20px; color:white; font-size:16px;">클릭하여 닫기 / ESC 키</div>
                        </div>
                        <script>
                        (function() {
                            // 라이트박스 닫기
                            var overlay = document.getElementById('lightbox-overlay');
                            overlay.addEventListener('click', function() {
                                this.style.display = 'none';
                            });
                            // ESC 키로 닫기
                            document.addEventListener('keydown', function(e) {
                                if (e.key === 'Escape') {
                                    overlay.style.display = 'none';
                                }
                            });
                            // 이미지 클릭 이벤트 설정
                            function setupImageClick() {
                                var container = document.getElementById('defect-visualization');
                                if (container) {
                                    var img = container.querySelector('img');
                                    if (img && !img.dataset.lightboxSetup) {
                                        img.dataset.lightboxSetup = 'true';
                                        img.style.cursor = 'zoom-in';
                                        img.addEventListener('click', function(e) {
                                            e.stopPropagation();
                                            var lightboxImg = document.getElementById('lightbox-img');
                                            lightboxImg.src = this.src;
                                            overlay.style.display = 'flex';
                                        });
                                    }
                                }
                            }
                            // 초기 설정 및 주기적 체크 (동적 이미지 변경 대응)
                            setInterval(setupImageClick, 500);
                            setupImageClick();
                        })();
                        </script>
                        """)

                        # 숨겨진 이미지 입력 (채팅용)
                        chat_image_input = gr.Image(type="pil", visible=False)

                        with gr.Accordion("자동 분석 결과", open=False):
                            gr.Markdown("#### 분석 결과")
                            result_html = gr.HTML(label="", show_label=False)
                            analyze_btn = gr.Button("자동 분석 실행", variant="secondary", size="sm")
                            with gr.Accordion("Raw 응답", open=False):
                                raw_output = gr.Textbox(label="모델 응답", lines=5)

                    # 오른쪽: VLM 채팅 영역
                    with gr.Column(scale=1):
                        with gr.Row():
                            gr.Markdown("#### VLM 채팅")
                            demo_mode_checkbox = gr.Checkbox(
                                label="데모 모드 (빠른 응답)",
                                value=True,
                                interactive=True,
                            )
                        demo_mode_status = gr.Markdown("**✅ 데모 모드 활성화** (빠른 응답, 메타데이터 기반)")
                        gr.Markdown("이미지를 업로드하거나 샘플을 선택한 후 자연어로 질문하세요.", elem_classes=["info-text"])

                        vlm_chatbot = gr.Chatbot(
                            label="VLM 대화",
                            height=320,
                            show_label=False,
                        )

                        with gr.Row():
                            vlm_chat_input = gr.Textbox(
                                label="",
                                placeholder="예: 어떤 결함이 있나요? / x,y 좌표 알려줘 / 유사 이미지 찾아줘",
                                show_label=False,
                                scale=5,
                            )
                            vlm_send_btn = gr.Button("전송", variant="primary", scale=1)

                        with gr.Row():
                            vlm_clear_btn = gr.Button("대화 초기화", variant="secondary", size="sm")
                            vlm_save_btn = gr.Button("채팅 저장", variant="primary", size="sm")
                        vlm_save_status = gr.Markdown("")

                        gr.Markdown("""
                        ---
                        **질문 예시:**
                        - "이 이미지에 어떤 결함이 있나요?"
                        - "**x,y 좌표 알려줘**" (결함 좌표)
                        - "**좌표 시각화해줘**" (이미지에 표시)
                        - "**유사 이미지 찾아줘**" (비슷한 결함 검색)
                        - "결함의 원인은 무엇인가요?"
                        - "심각도는 어느 정도인가요?"
                        - "어떤 조치가 필요한가요?"
                        - "결함 크기는 얼마인가요?"
                        """)

                # 이벤트 연결
                analyze_btn.click(analyze_image, inputs=[chat_image_input], outputs=[result_html, raw_output])

                # 데모 모드 전환 이벤트
                demo_mode_checkbox.change(
                    toggle_demo_mode,
                    inputs=[demo_mode_checkbox],
                    outputs=[demo_mode_status]
                )

                # 채팅 이벤트 (이미지 이름 포함)
                vlm_send_btn.click(
                    vlm_chat,
                    inputs=[vlm_chat_input, vlm_chatbot, chat_image_input, current_image_name],
                    outputs=[vlm_chatbot, vlm_chat_input]
                )
                vlm_chat_input.submit(
                    vlm_chat,
                    inputs=[vlm_chat_input, vlm_chatbot, chat_image_input, current_image_name],
                    outputs=[vlm_chatbot, vlm_chat_input]
                )
                vlm_clear_btn.click(
                    clear_chat,
                    outputs=[vlm_chatbot, chat_image_input]
                )

                # VLM 채팅 영역 내 저장 버튼 이벤트
                vlm_save_btn.click(
                    save_chat_history,
                    inputs=[vlm_chatbot, current_image_name],
                    outputs=[vlm_save_status, history_dropdown]
                )

                # 샘플 이미지 선택 이벤트 (메타데이터 연동)
                if SAMPLE_IMAGES:
                    sample_gallery.select(
                        select_sample_image,
                        outputs=[chat_image_input, current_image_name]
                    ).then(
                        get_image_info,
                        inputs=[current_image_name],
                        outputs=[selected_image_info]
                    ).then(
                        # 자동 시각화
                        visualize_defect_coordinates,
                        inputs=[chat_image_input, current_image_name],
                        outputs=[visualized_image, visualization_info]
                    )

                # 채팅 히스토리 이벤트
                load_history_btn.click(
                    load_chat_history,
                    inputs=[history_dropdown],
                    outputs=[vlm_chatbot, current_image_name, history_status]
                )
                delete_history_btn.click(
                    delete_chat_history,
                    inputs=[history_dropdown],
                    outputs=[history_status, history_dropdown]
                )

            # ===== 탭 4: 품질 대시보드 =====
            with gr.TabItem("4. 품질 대시보드", id="dashboard"):
                gr.Markdown("### 품질 분석 대시보드")
                gr.Markdown("목업 데이터를 기반으로 품질 현황을 시각화합니다.")

                with gr.Row():
                    with gr.Column(scale=3):
                        dashboard_img = gr.Image(label="종합 대시보드", height=550)
                    with gr.Column(scale=1):
                        gr.Markdown("#### 대시보드 생성")
                        dashboard_gen_btn = gr.Button("종합 대시보드 생성", variant="primary", size="lg")
                        dashboard_status = gr.Textbox(label="분석 요약", interactive=False, lines=2)

                        gr.Markdown("---")
                        gr.Markdown("#### 개별 차트")
                        defect_chart_btn = gr.Button("불량 유형 분포", variant="secondary")
                        equipment_chart_btn = gr.Button("설비별 불량률", variant="secondary")
                        customer_chart_btn = gr.Button("고객사별 품질", variant="secondary")

                dashboard_gen_btn.click(generate_quality_dashboard, outputs=[dashboard_img, dashboard_status])
                defect_chart_btn.click(generate_defect_chart, outputs=[dashboard_img])
                equipment_chart_btn.click(generate_equipment_chart, outputs=[dashboard_img])
                customer_chart_btn.click(generate_customer_chart, outputs=[dashboard_img])

                gr.Markdown("---")

                with gr.Row():
                    with gr.Column():
                        gr.Markdown("""
                        #### 대시보드 구성

                        | 차트 | 설명 |
                        |------|------|
                        | 불량 유형별 분포 | 8가지 불량 유형의 비율 (파이 차트) |
                        | 심각도별 불량 | CRITICAL/MAJOR/MINOR/COSMETIC 분포 |
                        | 설비별 불량률 | 불량률 상위 10개 설비 및 평균선 |
                        | 고객사별 품질 | 수율 및 통과율 비교 |
                        | 라인별 생산 | 양품/불량 수량 및 수율 |
                        | KPI 요약 | 핵심 품질 지표 |
                        """)
                    with gr.Column():
                        gr.Markdown("""
                        #### 사용 방법

                        1. **목업 데이터 탭**에서 먼저 데이터 생성
                        2. **종합 대시보드 생성** 버튼 클릭
                        3. 개별 차트 버튼으로 상세 분석

                        #### 데이터 요구사항
                        - Quality 데이터: 불량 유형, 심각도 분석
                        - Manufacturing 데이터: 고객사별 수율
                        - MES 데이터: 라인별 생산 실적
                        """)

            # ===== 탭 5: GraphRAG 관리 =====
            with gr.TabItem("5. GraphRAG 관리", id="graphrag"):
                gr.Markdown("### 지식 그래프 데이터 관리")

                with gr.Row():
                    with gr.Column(scale=1):
                        gr.Markdown("#### 초기화")
                        with gr.Row():
                            use_neo4j = gr.Checkbox(label="Neo4j 사용", value=False)
                            neo4j_pwd = gr.Textbox(label="Neo4j 비밀번호", value="password", type="password")
                        init_btn = gr.Button("지식 베이스 초기화", variant="primary")
                        init_status = gr.Textbox(label="상태", interactive=False)

                        init_btn.click(init_knowledge_base, inputs=[use_neo4j, neo4j_pwd], outputs=[init_status])

                    with gr.Column(scale=1):
                        gr.Markdown("#### 데이터 조회")
                        with gr.Row():
                            view_nodes_btn = gr.Button("노드 조회")
                            view_rels_btn = gr.Button("관계 조회")
                        nodes_output = gr.Textbox(label="결과", lines=12, interactive=False)

                        view_nodes_btn.click(get_all_nodes, outputs=[nodes_output])
                        view_rels_btn.click(get_all_relations, outputs=[nodes_output])

                gr.Markdown("---")

                # 엑셀 업로드 섹션
                with gr.Row():
                    with gr.Column(scale=1):
                        gr.Markdown("#### 엑셀 템플릿 다운로드")
                        gr.Markdown("데이터 입력 형식을 확인하려면 템플릿을 다운로드하세요.")
                        template_btn = gr.Button("템플릿 생성", variant="secondary")
                        template_file = gr.File(label="템플릿 파일", interactive=False)
                        template_status = gr.Textbox(label="상태", interactive=False, lines=2)

                        template_btn.click(create_excel_template, outputs=[template_file, template_status])

                    with gr.Column(scale=1):
                        gr.Markdown("#### 엑셀 데이터 업로드")
                        gr.Markdown("템플릿 형식에 맞게 작성한 엑셀 파일을 업로드하세요.")
                        excel_file = gr.File(label="엑셀 파일 (.xlsx)", file_types=[".xlsx", ".xls"])
                        upload_btn = gr.Button("데이터 업로드", variant="primary")
                        upload_status = gr.Textbox(label="업로드 결과", interactive=False, lines=6)

                        upload_btn.click(upload_excel_data, inputs=[excel_file], outputs=[upload_status])

                gr.Markdown("---")
                gr.Markdown("#### 또는 개별 노드 직접 입력")

                with gr.Row():
                    # 결함 노드 추가
                    with gr.Column():
                        gr.Markdown("#### 결함 노드 추가")
                        def_id = gr.Textbox(label="결함 ID", placeholder="DEF007")
                        def_type = gr.Dropdown(
                            label="결함 유형",
                            choices=["dead_pixel", "bright_spot", "line_defect", "mura", "scratch", "particle", "custom"],
                            value="custom"
                        )
                        def_name = gr.Textbox(label="한글명", placeholder="색편차")
                        def_desc = gr.Textbox(label="설명", placeholder="색상이 기준값에서 벗어난 결함")
                        def_severity = gr.Textbox(label="심각도 (쉼표구분)", placeholder="low, medium, high")
                        def_visual = gr.Textbox(label="시각적 특징", placeholder="특정 영역의 색상 변화")
                        add_def_btn = gr.Button("결함 추가", variant="secondary")
                        add_def_status = gr.Textbox(label="결과", interactive=False)

                        add_def_btn.click(add_defect_node,
                                          inputs=[def_id, def_type, def_name, def_desc, def_severity, def_visual],
                                          outputs=[add_def_status])

                    # 원인 노드 추가
                    with gr.Column():
                        gr.Markdown("#### 원인 노드 추가")
                        cause_id = gr.Textbox(label="원인 ID", placeholder="RC009")
                        cause_type = gr.Textbox(label="원인 유형", placeholder="temperature_variation")
                        cause_name = gr.Textbox(label="한글명", placeholder="온도 변화")
                        cause_desc = gr.Textbox(label="설명", placeholder="공정 중 온도 제어 불량")
                        cause_cat = gr.Dropdown(
                            label="분류",
                            choices=["equipment", "process", "material", "environment", "human"],
                            value="process"
                        )
                        add_cause_btn = gr.Button("원인 추가", variant="secondary")
                        add_cause_status = gr.Textbox(label="결과", interactive=False)

                        add_cause_btn.click(add_cause_node,
                                            inputs=[cause_id, cause_type, cause_name, cause_desc, cause_cat],
                                            outputs=[add_cause_status])

                    # 조치 노드 추가
                    with gr.Column():
                        gr.Markdown("#### 조치 노드 추가")
                        act_id = gr.Textbox(label="조치 ID", placeholder="ACT006")
                        act_type = gr.Textbox(label="조치 유형", placeholder="temperature_control")
                        act_name = gr.Textbox(label="한글명", placeholder="온도 제어 강화")
                        act_desc = gr.Textbox(label="설명", placeholder="공정 온도 모니터링 강화")
                        act_priority = gr.Dropdown(
                            label="우선순위",
                            choices=["immediate", "high", "medium", "low"],
                            value="medium"
                        )
                        add_act_btn = gr.Button("조치 추가", variant="secondary")
                        add_act_status = gr.Textbox(label="결과", interactive=False)

                        add_act_btn.click(add_action_node,
                                          inputs=[act_id, act_type, act_name, act_desc, act_priority],
                                          outputs=[add_act_status])

                gr.Markdown("---")
                gr.Markdown("#### 관계 추가")

                with gr.Row():
                    with gr.Column():
                        gr.Markdown("##### CAUSED_BY (결함 → 원인)")
                        cb_def_id = gr.Textbox(label="결함 ID", placeholder="DEF007")
                        cb_cause_id = gr.Textbox(label="원인 ID", placeholder="RC009")
                        cb_prob = gr.Slider(label="확률", minimum=0, maximum=1, value=0.5, step=0.05)
                        cb_evidence = gr.Textbox(label="근거", placeholder="온도 변화로 인한 박막 불균일")
                        add_cb_btn = gr.Button("관계 추가", variant="secondary")
                        add_cb_status = gr.Textbox(label="결과", interactive=False)

                        add_cb_btn.click(add_caused_by_relation,
                                         inputs=[cb_def_id, cb_cause_id, cb_prob, cb_evidence],
                                         outputs=[add_cb_status])

                    with gr.Column():
                        gr.Markdown("##### REQUIRES (원인 → 조치)")
                        rq_cause_id = gr.Textbox(label="원인 ID", placeholder="RC009")
                        rq_act_id = gr.Textbox(label="조치 ID", placeholder="ACT006")
                        rq_eff = gr.Slider(label="효과성", minimum=0, maximum=1, value=0.5, step=0.05)
                        add_rq_btn = gr.Button("관계 추가", variant="secondary")
                        add_rq_status = gr.Textbox(label="결과", interactive=False)

                        add_rq_btn.click(add_requires_relation,
                                         inputs=[rq_cause_id, rq_act_id, rq_eff],
                                         outputs=[add_rq_status])

                    with gr.Column():
                        gr.Markdown("##### 결함 분석 쿼리")
                        query_type = gr.Dropdown(
                            label="결함 유형",
                            choices=["dead_pixel", "bright_spot", "line_defect", "mura", "scratch", "particle"],
                            value="dead_pixel"
                        )
                        query_btn = gr.Button("분석 쿼리", variant="primary")
                        query_result = gr.Textbox(label="분석 결과", lines=10, interactive=False)

                        query_btn.click(query_defect_analysis, inputs=[query_type], outputs=[query_result])

            # ===== 탭 6: 그래프 시각화 =====
            with gr.TabItem("6. 그래프 시각화", id="visualization"):
                gr.Markdown("### 지식 그래프 시각화")

                with gr.Row():
                    with gr.Column(scale=1):
                        gr.Markdown("#### 전체 그래프")
                        full_graph_btn = gr.Button("전체 그래프 생성", variant="primary")
                        full_graph_status = gr.Textbox(label="통계", interactive=False)
                        full_graph_img = gr.Image(label="지식 그래프", height=500)

                        full_graph_btn.click(create_graph_visualization,
                                             outputs=[full_graph_img, full_graph_status])

                    with gr.Column(scale=1):
                        gr.Markdown("#### 결함별 서브그래프")
                        subgraph_type = gr.Dropdown(
                            label="결함 유형 선택",
                            choices=["dead_pixel", "bright_spot", "line_defect", "mura", "scratch", "particle"],
                            value="dead_pixel"
                        )
                        subgraph_btn = gr.Button("서브그래프 생성", variant="primary")
                        subgraph_status = gr.Textbox(label="정보", interactive=False)
                        subgraph_img = gr.Image(label="결함 분석 그래프", height=500)

                        subgraph_btn.click(create_subgraph_visualization,
                                           inputs=[subgraph_type],
                                           outputs=[subgraph_img, subgraph_status])

            # ===== 탭 7: 목업 데이터 생성 =====
            with gr.TabItem("7. 목업 데이터", id="mockdata"):
                gr.Markdown("### 시연용 빅데이터 생성")
                gr.Markdown("개발단계, 제조현장, MES 실적 데이터를 포함한 종합 목업 데이터 생성")

                with gr.Row():
                    with gr.Column(scale=1):
                        gr.Markdown("#### 데이터 생성 설정")

                        with gr.Group():
                            mock_num_lots = gr.Slider(
                                label="로트 수",
                                minimum=10, maximum=500, value=100, step=10,
                                info="생성할 로트 수 (10~500)"
                            )
                            mock_num_cells = gr.Slider(
                                label="셀/로트",
                                minimum=10, maximum=100, value=50, step=10,
                                info="로트당 셀 수 (10~100)"
                            )
                            mock_num_days = gr.Slider(
                                label="기간 (일)",
                                minimum=7, maximum=90, value=30, step=7,
                                info="데이터 생성 기간 (7~90일)"
                            )

                        with gr.Row():
                            mock_init_btn = gr.Button("생성기 초기화", variant="secondary")
                            mock_gen_btn = gr.Button("데이터 생성", variant="primary")

                        mock_init_status = gr.Textbox(label="초기화 상태", interactive=False)
                        mock_init_btn.click(init_mock_generator, outputs=[mock_init_status])

                        gr.Markdown("---")
                        gr.Markdown("#### 데이터 미리보기")
                        mock_preview_cat = gr.Dropdown(
                            label="카테고리 선택",
                            choices=["development", "equipment", "material", "inspection", "quality", "manufacturing", "mes", "traceability"],
                            value="manufacturing"
                        )
                        mock_preview_btn = gr.Button("미리보기", variant="secondary")
                        mock_preview_output = gr.Textbox(label="미리보기", lines=15, interactive=False)
                        mock_preview_btn.click(get_mock_data_preview, inputs=[mock_preview_cat], outputs=[mock_preview_output])

                    with gr.Column(scale=1):
                        gr.Markdown("#### 생성 결과")
                        mock_result_html = gr.HTML(label="")

                        with gr.Accordion("JSON 상세 정보", open=False):
                            mock_raw_json = gr.Textbox(label="Raw JSON", lines=12, interactive=False)

                mock_gen_btn.click(
                    generate_mock_data,
                    inputs=[mock_num_lots, mock_num_cells, mock_num_days],
                    outputs=[mock_result_html, mock_raw_json]
                )

                gr.Markdown("---")
                gr.Markdown("""
                #### 생성되는 데이터 종류

                | 카테고리 | 설명 | 주요 필드 |
                |----------|------|-----------|
                | 🔬 Development | 개발단계 데이터 (EVT/DVT/PVT) | project_id, phase, test_result, measurements |
                | ⚙️ Equipment | 설비 마스터 | equipment_id, type, line, status, parameters |
                | 📦 Material | 자재 데이터 | material_id, lot_no, vendor, specifications |
                | 🔍 Inspection | 검사 데이터 | inspection_id, result, defect_codes, measurements |
                | ✅ Quality | 양/불량 데이터 | defect_type, severity, root_cause, action |
                | 🏭 Manufacturing | 제조현장 통합 | lot_id, cell_id, process_step, yield_rate |
                | 📊 MES | MES 실적 | work_order_id, plan_qty, actual_qty, yield_rate |
                | 🔗 Traceability | 이력추적 | trace_id, event_type, parameters, linked_defects |

                #### 데이터 관계

                ```
                ┌──────────────┐     ┌──────────────┐     ┌──────────────┐
                │  Development │────▶│ Manufacturing│────▶│     MES      │
                │  (개발단계)   │     │  (제조현장)   │     │   (실적)     │
                └──────────────┘     └──────────────┘     └──────────────┘
                       │                    │                    │
                       ▼                    ▼                    ▼
                ┌──────────────┐     ┌──────────────┐     ┌──────────────┐
                │  Equipment   │     │  Inspection  │     │   Quality    │
                │   (설비)     │     │   (검사)     │     │  (품질결과)  │
                └──────────────┘     └──────────────┘     └──────────────┘
                       │                    │                    │
                       └────────────────────┴────────────────────┘
                                           │
                                           ▼
                                   ┌──────────────┐
                                   │ Traceability │
                                   │  (이력추적)   │
                                   └──────────────┘
                ```
                """)

            # ===== 탭 8: 시스템 정보 =====
            with gr.TabItem("8. 시스템 정보", id="info"):
                with gr.Row():
                    with gr.Column(scale=1):
                        gr.HTML("""
<div style="padding: 20px; font-family: 'Malgun Gothic', sans-serif;">
    <h2 style="color: #1a73e8; border-bottom: 2px solid #1a73e8; padding-bottom: 10px;">시스템 개요</h2>
    <p style="font-size: 15px; color: #333;"><strong>SDC Customer Quality System</strong>은 디스플레이 제조 품질 관리를 위한 AI 기반 통합 솔루션입니다.</p>

    <h3 style="color: #333; margin-top: 25px;">시스템 아키텍처</h3>
    <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); border-radius: 12px; padding: 25px; margin: 15px 0; color: white;">
        <div style="text-align: center; font-size: 18px; font-weight: bold; margin-bottom: 15px;">디스플레이 결함 분석 시스템 v2.0</div>
        <div style="background: rgba(255,255,255,0.15); border-radius: 8px; padding: 15px; margin: 10px 0;">
            <div style="display: flex; justify-content: center; align-items: center; gap: 10px; flex-wrap: wrap;">
                <span style="background: #fff; color: #667eea; padding: 8px 15px; border-radius: 20px; font-weight: bold;">이미지 입력</span>
                <span style="font-size: 20px;">→</span>
                <span style="background: #fff; color: #667eea; padding: 8px 15px; border-radius: 20px; font-weight: bold;">VLM 추론<br><small style="font-size:10px;">Cosmos Reason 7B</small></span>
                <span style="font-size: 20px;">→</span>
                <span style="background: #fff; color: #667eea; padding: 8px 15px; border-radius: 20px; font-weight: bold;">GraphRAG<br><small style="font-size:10px;">Neo4j 지식그래프</small></span>
                <span style="font-size: 20px;">→</span>
                <span style="background: #fff; color: #667eea; padding: 8px 15px; border-radius: 20px; font-weight: bold;">분석 결과</span>
            </div>
        </div>
        <div style="background: rgba(255,255,255,0.15); border-radius: 8px; padding: 15px; margin: 10px 0;">
            <div style="display: flex; justify-content: center; align-items: center; gap: 10px; flex-wrap: wrap;">
                <span style="background: #fff; color: #764ba2; padding: 8px 15px; border-radius: 20px; font-weight: bold;">CS 불만접수</span>
                <span style="font-size: 20px;">→</span>
                <span style="background: #fff; color: #764ba2; padding: 8px 15px; border-radius: 20px; font-weight: bold;">1차 분석<br><small style="font-size:10px;">품질부서</small></span>
                <span style="font-size: 20px;">→</span>
                <span style="background: #fff; color: #764ba2; padding: 8px 15px; border-radius: 20px; font-weight: bold;">2차 분석<br><small style="font-size:10px;">귀책부서</small></span>
                <span style="font-size: 20px;">→</span>
                <span style="background: #fff; color: #764ba2; padding: 8px 15px; border-radius: 20px; font-weight: bold;">보고서 생성</span>
            </div>
        </div>
    </div>

    <h3 style="color: #333; margin-top: 25px;">핵심 기능</h3>
    <table style="width: 100%; border-collapse: collapse; margin: 15px 0; font-size: 14px;">
        <thead>
            <tr style="background: #1a73e8; color: white;">
                <th style="padding: 12px; text-align: left; border: 1px solid #ddd;">기능</th>
                <th style="padding: 12px; text-align: left; border: 1px solid #ddd;">설명</th>
            </tr>
        </thead>
        <tbody>
            <tr style="background: #f9f9f9;">
                <td style="padding: 10px; border: 1px solid #ddd; font-weight: bold;">CS 워크플로우</td>
                <td style="padding: 10px; border: 1px solid #ddd;">고객 불만 접수부터 최종 보고서까지 전체 프로세스 관리</td>
            </tr>
            <tr>
                <td style="padding: 10px; border: 1px solid #ddd; font-weight: bold;">결재 시스템</td>
                <td style="padding: 10px; border: 1px solid #ddd;">다단계 결재 라인, 기한 관리, 알림 발송</td>
            </tr>
            <tr style="background: #f9f9f9;">
                <td style="padding: 10px; border: 1px solid #ddd; font-weight: bold;">VLM 분석</td>
                <td style="padding: 10px; border: 1px solid #ddd;">Cosmos Reason 7B 기반 결함 이미지 분석</td>
            </tr>
            <tr>
                <td style="padding: 10px; border: 1px solid #ddd; font-weight: bold;">GraphRAG</td>
                <td style="padding: 10px; border: 1px solid #ddd;">Neo4j 지식그래프 기반 유사 사례 검색</td>
            </tr>
            <tr style="background: #f9f9f9;">
                <td style="padding: 10px; border: 1px solid #ddd; font-weight: bold;">빅데이터 분석</td>
                <td style="padding: 10px; border: 1px solid #ddd;">Spark 기반 대용량 품질 데이터 분석</td>
            </tr>
        </tbody>
    </table>

    <h3 style="color: #333; margin-top: 25px;">기술 스택</h3>
    <table style="width: 100%; border-collapse: collapse; margin: 15px 0; font-size: 14px;">
        <thead>
            <tr style="background: #34a853; color: white;">
                <th style="padding: 12px; text-align: left; border: 1px solid #ddd;">구성요소</th>
                <th style="padding: 12px; text-align: left; border: 1px solid #ddd;">기술</th>
                <th style="padding: 12px; text-align: left; border: 1px solid #ddd;">버전</th>
            </tr>
        </thead>
        <tbody>
            <tr style="background: #f9f9f9;"><td style="padding: 10px; border: 1px solid #ddd; font-weight: bold;">VLM</td><td style="padding: 10px; border: 1px solid #ddd;">Cosmos Reason 7B (NVIDIA)</td><td style="padding: 10px; border: 1px solid #ddd;">-</td></tr>
            <tr><td style="padding: 10px; border: 1px solid #ddd; font-weight: bold;">Fine-tuning</td><td style="padding: 10px; border: 1px solid #ddd;">LoRA</td><td style="padding: 10px; border: 1px solid #ddd;">r=64, alpha=128</td></tr>
            <tr style="background: #f9f9f9;"><td style="padding: 10px; border: 1px solid #ddd; font-weight: bold;">Graph DB</td><td style="padding: 10px; border: 1px solid #ddd;">Neo4j</td><td style="padding: 10px; border: 1px solid #ddd;">5.x</td></tr>
            <tr><td style="padding: 10px; border: 1px solid #ddd; font-weight: bold;">빅데이터</td><td style="padding: 10px; border: 1px solid #ddd;">Apache Spark</td><td style="padding: 10px; border: 1px solid #ddd;">3.x</td></tr>
            <tr style="background: #f9f9f9;"><td style="padding: 10px; border: 1px solid #ddd; font-weight: bold;">API</td><td style="padding: 10px; border: 1px solid #ddd;">FastAPI</td><td style="padding: 10px; border: 1px solid #ddd;">-</td></tr>
            <tr><td style="padding: 10px; border: 1px solid #ddd; font-weight: bold;">UI</td><td style="padding: 10px; border: 1px solid #ddd;">Gradio</td><td style="padding: 10px; border: 1px solid #ddd;">4.x</td></tr>
            <tr style="background: #f9f9f9;"><td style="padding: 10px; border: 1px solid #ddd; font-weight: bold;">Container</td><td style="padding: 10px; border: 1px solid #ddd;">Docker</td><td style="padding: 10px; border: 1px solid #ddd;">-</td></tr>
            <tr><td style="padding: 10px; border: 1px solid #ddd; font-weight: bold;">문서생성</td><td style="padding: 10px; border: 1px solid #ddd;">python-docx</td><td style="padding: 10px; border: 1px solid #ddd;">-</td></tr>
        </tbody>
    </table>
</div>
                        """)

                    with gr.Column(scale=1):
                        gr.HTML("""
<div style="padding: 20px; font-family: 'Malgun Gothic', sans-serif;">
    <h2 style="color: #1a73e8; border-bottom: 2px solid #1a73e8; padding-bottom: 10px;">사용 가이드</h2>

    <h3 style="color: #333; margin-top: 20px;">1. CS 워크플로우</h3>
    <div style="margin: 10px 0;">
        <div style="background: #e8f0fe; border-left: 4px solid #1a73e8; padding: 12px; margin: 8px 0; border-radius: 0 8px 8px 0;">
            <strong style="color: #1a73e8;">Step 1. 불만 접수</strong>
            <ul style="margin: 8px 0 0 0; padding-left: 20px; color: #333;">
                <li>고객사, 제품, 결함 유형 입력</li>
                <li>고객 확인 이메일 생성 및 발송</li>
                <li>사내 미팅 요청 이메일 발송</li>
            </ul>
        </div>
        <div style="background: #e6f4ea; border-left: 4px solid #34a853; padding: 12px; margin: 8px 0; border-radius: 0 8px 8px 0;">
            <strong style="color: #34a853;">Step 2. 1차 분석</strong>
            <ul style="margin: 8px 0 0 0; padding-left: 20px; color: #333;">
                <li>품질부서에서 기본 분석 수행</li>
                <li>미팅 결과 정리 (LLM 활용)</li>
                <li>업무 할당 및 납기 등록</li>
                <li>1차 산출물 보고서 생성</li>
            </ul>
        </div>
        <div style="background: #fef7e0; border-left: 4px solid #f9ab00; padding: 12px; margin: 8px 0; border-radius: 0 8px 8px 0;">
            <strong style="color: #f9ab00;">Step 3. 2차 분석</strong>
            <ul style="margin: 8px 0 0 0; padding-left: 20px; color: #333;">
                <li>귀책부서에서 상세 분석 수행</li>
                <li>재발 방지 대책 수립</li>
                <li>2차 산출물 보고서 생성</li>
            </ul>
        </div>
        <div style="background: #fce8e6; border-left: 4px solid #ea4335; padding: 12px; margin: 8px 0; border-radius: 0 8px 8px 0;">
            <strong style="color: #ea4335;">Step 4. 보고서 생성</strong>
            <ul style="margin: 8px 0 0 0; padding-left: 20px; color: #333;">
                <li>최종 보고서 생성 (Word)</li>
                <li>고객 회신 메일 작성 및 발송</li>
            </ul>
        </div>
    </div>

    <h3 style="color: #333; margin-top: 25px;">2. 결재 시스템</h3>
    <ul style="padding-left: 20px; color: #333; line-height: 1.8;">
        <li><strong>결재 유형</strong>: 이메일 발송, 분석 결과, 최종 보고서</li>
        <li><strong>결재 라인</strong>: 팀장 → 부장 → 상무 (유형별 상이)</li>
        <li><strong>기한 관리</strong>: 유형별 기한 설정, 초과 시 알림</li>
        <li><strong>알림 발송</strong>: 기한 초과/임박 시 담당자 알림</li>
    </ul>

    <h3 style="color: #333; margin-top: 25px;">3. 접속 정보</h3>
    <table style="width: 100%; border-collapse: collapse; margin: 15px 0; font-size: 14px;">
        <thead>
            <tr style="background: #5f6368; color: white;">
                <th style="padding: 10px; text-align: left; border: 1px solid #ddd;">서비스</th>
                <th style="padding: 10px; text-align: left; border: 1px solid #ddd;">URL</th>
                <th style="padding: 10px; text-align: left; border: 1px solid #ddd;">비고</th>
            </tr>
        </thead>
        <tbody>
            <tr style="background: #f9f9f9;"><td style="padding: 10px; border: 1px solid #ddd; font-weight: bold;">본 시스템</td><td style="padding: 10px; border: 1px solid #ddd;">http://localhost:7860</td><td style="padding: 10px; border: 1px solid #ddd;">Gradio UI</td></tr>
            <tr><td style="padding: 10px; border: 1px solid #ddd; font-weight: bold;">Neo4j Browser</td><td style="padding: 10px; border: 1px solid #ddd;">http://localhost:7474</td><td style="padding: 10px; border: 1px solid #ddd;">그래프 DB</td></tr>
            <tr style="background: #f9f9f9;"><td style="padding: 10px; border: 1px solid #ddd; font-weight: bold;">Neo4j Bolt</td><td style="padding: 10px; border: 1px solid #ddd;">bolt://localhost:7687</td><td style="padding: 10px; border: 1px solid #ddd;">연결 포트</td></tr>
        </tbody>
    </table>
    <p style="color: #333;"><strong>Neo4j 기본 계정</strong>: <code style="background: #f1f3f4; padding: 2px 6px; border-radius: 4px;">neo4j</code> / <code style="background: #f1f3f4; padding: 2px 6px; border-radius: 4px;">password</code></p>

    <h3 style="color: #333; margin-top: 25px;">4. 디렉토리 구조</h3>
    <div style="background: #f8f9fa; border: 1px solid #ddd; border-radius: 8px; padding: 15px; font-family: monospace; font-size: 13px;">
        <div style="color: #333;">/tmp/</div>
        <div style="color: #333; padding-left: 20px;">├── cs_emails/ <span style="color: #666;"># 발송된 이메일 로그</span></div>
        <div style="color: #333; padding-left: 20px;">├── cs_approval_notifications/ <span style="color: #666;"># 알림 발송 로그</span></div>
        <div style="color: #333; padding-left: 20px;">└── cs_reports/ <span style="color: #666;"># 생성된 보고서 파일</span></div>
        <div style="color: #333; padding-left: 50px;">├── first_output/ <span style="color: #666;"># 1차 산출물 보고서</span></div>
        <div style="color: #333; padding-left: 50px;">├── second_output/ <span style="color: #666;"># 2차 산출물 보고서</span></div>
        <div style="color: #333; padding-left: 50px;">└── final/ <span style="color: #666;"># 최종 보고서</span></div>
    </div>

    <h3 style="color: #333; margin-top: 25px;">5. 문의 및 지원</h3>
    <ul style="padding-left: 20px; color: #333; line-height: 1.8;">
        <li><strong>개발팀</strong>: 품질관리시스템개발팀</li>
        <li><strong>이메일</strong>: qms-support@sdc.com</li>
    </ul>
</div>
                        """)

        # 푸터
        gr.HTML("""
        <div style="text-align: center; padding: 10px; margin-top: 15px; color: #666666; font-size: 12px; background: #ffffff; border-top: 1px solid #e0e0e0;">
            SDC Customer Quality - Cosmos Reason VLM + GraphRAG Demo
        </div>
        """)

    return demo


if __name__ == "__main__":
    import os
    import subprocess
    import time

    PORT = 7860

    # 기존 포트 사용 프로세스 종료 (현재 프로세스 제외)
    print(f"\n포트 {PORT} 정리 중...")
    current_pid = os.getpid()
    # fuser로 포트 사용 프로세스만 종료
    subprocess.run(f"fuser -k {PORT}/tcp 2>/dev/null || true", shell=True)
    time.sleep(2)
    print("완료")

    # 자동으로 인메모리 KB 초기화
    init_knowledge_base(use_neo4j=False)

    print(f"\n서버 시작: http://localhost:{PORT}")

    demo = create_demo()
    demo.launch(
        server_name="0.0.0.0",
        server_port=PORT,
        share=False,
        theme=gr.themes.Default(),
        css=LIGHT_CSS,
    )
