"""
보고서 생성 모듈
최종 분석 보고서 Word 문서 생성
"""

import json
import os
from datetime import datetime
from typing import Dict, Any, List, Optional
import logging

logger = logging.getLogger(__name__)

# python-docx 설치 확인
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed. Word document generation will be limited.")


class ReportGenerator:
    """보고서 생성기"""

    def __init__(self, output_dir: str = "/tmp/cs_reports"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

        # 과거 사례 문서 디렉토리
        self.past_cases_dir = os.path.join(output_dir, "past_cases")
        os.makedirs(self.past_cases_dir, exist_ok=True)

    def generate_final_report(self,
                              complaint_data: Dict[str, Any],
                              first_analysis: Dict[str, Any],
                              second_analysis: Dict[str, Any],
                              include_images: bool = True) -> str:
        """
        최종 분석 보고서 생성

        Args:
            complaint_data: CS 불만 접수 데이터
            first_analysis: 1차 분석 결과
            second_analysis: 2차 분석 결과
            include_images: 이미지 포함 여부

        Returns:
            생성된 보고서 파일 경로
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        complaint_id = complaint_data.get("complaint_id", "UNKNOWN")

        if DOCX_AVAILABLE:
            return self._generate_word_report(
                complaint_data, first_analysis, second_analysis,
                complaint_id, timestamp, include_images
            )
        else:
            return self._generate_text_report(
                complaint_data, first_analysis, second_analysis,
                complaint_id, timestamp
            )

    def _generate_word_report(self,
                               complaint_data: Dict,
                               first_analysis: Dict,
                               second_analysis: Dict,
                               complaint_id: str,
                               timestamp: str,
                               include_images: bool) -> str:
        """Word 문서 형식 보고서 생성"""
        doc = Document()

        # 문서 제목
        title = doc.add_heading('고객 품질 불량 분석 보고서', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 문서 정보
        doc.add_paragraph(f"보고서 번호: RPT-{complaint_id}-{timestamp}")
        doc.add_paragraph(f"작성일: {datetime.now().strftime('%Y년 %m월 %d일')}")
        doc.add_paragraph("")

        # 1. 불만 접수 정보
        doc.add_heading('1. 불만 접수 정보', level=1)

        info_table = doc.add_table(rows=8, cols=2)
        info_table.style = 'Table Grid'

        info_items = [
            ("접수번호", complaint_data.get("complaint_id", "")),
            ("접수일시", str(complaint_data.get("receipt_date", ""))),
            ("고객사", complaint_data.get("customer", "")),
            ("제품모델", complaint_data.get("product_model", "")),
            ("LOT ID", complaint_data.get("lot_id", "")),
            ("CELL ID", complaint_data.get("cell_id", "")),
            ("결함유형", complaint_data.get("defect_type", "")),
            ("심각도", complaint_data.get("severity", ""))
        ]

        for i, (label, value) in enumerate(info_items):
            info_table.rows[i].cells[0].text = label
            info_table.rows[i].cells[1].text = str(value)

        doc.add_paragraph("")

        # 2. 결함 현상
        doc.add_heading('2. 결함 현상', level=1)
        doc.add_paragraph(complaint_data.get("defect_description", "결함 설명 없음"))

        # 3. 1차 기본 분석 결과
        doc.add_heading('3. 1차 기본 분석 결과 (품질부서)', level=1)

        doc.add_heading('3.1 빅데이터 분석 결과', level=2)
        bigdata = first_analysis.get("bigdata_result", {})

        if bigdata.get("defect_statistics"):
            stats = bigdata["defect_statistics"]
            doc.add_paragraph(f"• 총 발생 건수: {stats.get('total_count', 'N/A')}건")
            doc.add_paragraph(f"• 심각도 분포: {stats.get('severity_distribution', {})}")

        if bigdata.get("equipment_analysis", {}).get("high_risk_equipment"):
            doc.add_paragraph("• 고위험 설비:")
            for eq in bigdata["equipment_analysis"]["high_risk_equipment"]:
                doc.add_paragraph(f"  - {eq.get('equipment_id')}: 불량률 {eq.get('defect_rate', 0):.2f}%")

        doc.add_heading('3.2 귀책 부서 판정', level=2)
        doc.add_paragraph(f"• 귀책 부서: {first_analysis.get('responsible_dept', '미정')}")
        doc.add_paragraph(f"• 신뢰도: {first_analysis.get('confidence_score', 0)*100:.1f}%")

        doc.add_heading('3.3 유사 사례', level=2)
        for case in first_analysis.get("similar_cases", []):
            doc.add_paragraph(f"• {case}")

        doc.add_paragraph("")

        # 4. 2차 상세 분석 결과
        doc.add_heading('4. 2차 상세 분석 결과 (귀책부서)', level=1)

        doc.add_heading('4.1 결함 이미지 분석', level=2)
        image_result = second_analysis.get("image_result", {})
        if image_result.get("detected_defects"):
            for defect in image_result["detected_defects"]:
                doc.add_paragraph(f"• 결함 유형: {defect.get('type', 'N/A')}")
                doc.add_paragraph(f"• 크기: {defect.get('size_mm', 'N/A')}mm")
                doc.add_paragraph(f"• 신뢰도: {defect.get('confidence', 0)*100:.1f}%")

        if image_result.get("defect_characteristics"):
            chars = image_result["defect_characteristics"]
            doc.add_paragraph(f"• 형상: {chars.get('shape', 'N/A')}")
            doc.add_paragraph(f"• 패턴: {chars.get('pattern', 'N/A')}")
            doc.add_paragraph(f"• 추정 원인: {chars.get('estimated_cause', 'N/A')}")

        doc.add_heading('4.2 GraphRAG 지식그래프 분석', level=2)
        graphrag = second_analysis.get("graphrag_result", {})

        cause_analysis = graphrag.get("cause_analysis", {})
        if cause_analysis:
            doc.add_paragraph(f"• 주요 원인: {cause_analysis.get('primary_cause', 'N/A')}")
            if cause_analysis.get("secondary_causes"):
                doc.add_paragraph("• 부가 원인:")
                for cause in cause_analysis["secondary_causes"]:
                    doc.add_paragraph(f"  - {cause}")

        doc.add_heading('4.3 과거 사례 분석', level=2)
        past_case = second_analysis.get("past_case_result", {})

        if past_case.get("cases_analyzed"):
            for case in past_case["cases_analyzed"]:
                doc.add_paragraph(f"• 사례 ID: {case.get('case_id', 'N/A')}")
                doc.add_paragraph(f"  - 발생일: {case.get('occurrence_date', 'N/A')}")
                doc.add_paragraph(f"  - 고객사: {case.get('customer', 'N/A')}")
                doc.add_paragraph(f"  - 원인: {case.get('root_cause', 'N/A')}")
                doc.add_paragraph(f"  - 대책: {case.get('countermeasure', 'N/A')}")
                doc.add_paragraph(f"  - 유사도: {case.get('similarity_score', 0)*100:.1f}%")
                doc.add_paragraph("")

        # 5. 근본 원인 분석
        doc.add_heading('5. 근본 원인 분석', level=1)
        doc.add_paragraph(second_analysis.get("root_cause", "원인 분석 중"))

        # 6. 대책 및 재발방지
        doc.add_heading('6. 대책 수립', level=1)

        doc.add_heading('6.1 즉각 대책', level=2)
        for i, measure in enumerate(second_analysis.get("countermeasures", []), 1):
            doc.add_paragraph(f"{i}. {measure}")

        doc.add_heading('6.2 재발 방지 대책', level=2)
        for i, measure in enumerate(second_analysis.get("prevention_measures", []), 1):
            doc.add_paragraph(f"{i}. {measure}")

        # 7. 결론
        doc.add_heading('7. 결론', level=1)
        conclusion = f"""
본 불량 건에 대한 분석 결과, 주요 원인은 '{second_analysis.get("root_cause", "미확인")}'으로 판단됩니다.
귀책 부서는 '{second_analysis.get("responsible_dept", "미정")}'이며,
상기 대책을 통해 유사 불량의 재발을 방지하고자 합니다.

분석 신뢰도: {second_analysis.get("confidence_score", 0)*100:.1f}%
"""
        doc.add_paragraph(conclusion)

        # 8. 승인
        doc.add_heading('8. 승인', level=1)

        approval_table = doc.add_table(rows=2, cols=4)
        approval_table.style = 'Table Grid'

        headers = ["작성", "검토", "승인", "결재"]
        for i, header in enumerate(headers):
            approval_table.rows[0].cells[i].text = header
            approval_table.rows[1].cells[i].text = ""

        # 저장
        filename = f"품질불량분석보고서_{complaint_id}_{timestamp}.docx"
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)

        logger.info(f"Word 보고서 생성 완료: {filepath}")
        return filepath

    def _generate_text_report(self,
                               complaint_data: Dict,
                               first_analysis: Dict,
                               second_analysis: Dict,
                               complaint_id: str,
                               timestamp: str) -> str:
        """텍스트 형식 보고서 생성 (python-docx 미설치 시)"""

        report_content = f"""
================================================================================
                      고객 품질 불량 분석 보고서
================================================================================

보고서 번호: RPT-{complaint_id}-{timestamp}
작성일: {datetime.now().strftime('%Y년 %m월 %d일')}

================================================================================
1. 불만 접수 정보
================================================================================
접수번호: {complaint_data.get("complaint_id", "")}
접수일시: {complaint_data.get("receipt_date", "")}
고객사: {complaint_data.get("customer", "")}
제품모델: {complaint_data.get("product_model", "")}
LOT ID: {complaint_data.get("lot_id", "")}
CELL ID: {complaint_data.get("cell_id", "")}
결함유형: {complaint_data.get("defect_type", "")}
심각도: {complaint_data.get("severity", "")}

================================================================================
2. 결함 현상
================================================================================
{complaint_data.get("defect_description", "결함 설명 없음")}

================================================================================
3. 1차 기본 분석 결과 (품질부서)
================================================================================
■ 귀책 부서: {first_analysis.get('responsible_dept', '미정')}
■ 신뢰도: {first_analysis.get('confidence_score', 0)*100:.1f}%
■ 유사 사례: {', '.join(first_analysis.get('similar_cases', []))}

================================================================================
4. 2차 상세 분석 결과 (귀책부서)
================================================================================
■ 이미지 분석 결과:
{json.dumps(second_analysis.get('image_result', {}), ensure_ascii=False, indent=2)}

■ GraphRAG 분석 결과:
{json.dumps(second_analysis.get('graphrag_result', {}), ensure_ascii=False, indent=2)}

================================================================================
5. 근본 원인 분석
================================================================================
{second_analysis.get("root_cause", "원인 분석 중")}

================================================================================
6. 대책 수립
================================================================================
■ 즉각 대책:
{chr(10).join(f"  {i+1}. {m}" for i, m in enumerate(second_analysis.get("countermeasures", [])))}

■ 재발 방지 대책:
{chr(10).join(f"  {i+1}. {m}" for i, m in enumerate(second_analysis.get("prevention_measures", [])))}

================================================================================
7. 결론
================================================================================
본 불량 건에 대한 분석 결과, 주요 원인은 '{second_analysis.get("root_cause", "미확인")}'으로
판단됩니다. 귀책 부서는 '{second_analysis.get("responsible_dept", "미정")}'이며,
상기 대책을 통해 유사 불량의 재발을 방지하고자 합니다.

분석 신뢰도: {second_analysis.get("confidence_score", 0)*100:.1f}%

================================================================================
8. 승인
================================================================================
작성:                  검토:                  승인:                  결재:

================================================================================
"""

        filename = f"품질불량분석보고서_{complaint_id}_{timestamp}.txt"
        filepath = os.path.join(self.output_dir, filename)

        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(report_content)

        logger.info(f"텍스트 보고서 생성 완료: {filepath}")
        return filepath

    def generate_past_case_document(self, case_data: Dict[str, Any]) -> str:
        """과거 사례 문서 생성"""
        case_id = case_data.get("case_id", f"CASE_{datetime.now().strftime('%Y%m%d%H%M%S')}")

        content = {
            "case_id": case_id,
            "created_date": datetime.now().isoformat(),
            "complaint_info": case_data.get("complaint_info", {}),
            "analysis_result": case_data.get("analysis_result", {}),
            "root_cause": case_data.get("root_cause", ""),
            "countermeasures": case_data.get("countermeasures", []),
            "prevention_measures": case_data.get("prevention_measures", []),
            "result": case_data.get("result", ""),
            "keywords": case_data.get("keywords", []),
            "tags": case_data.get("tags", [])
        }

        filename = f"case_{case_id}.json"
        filepath = os.path.join(self.past_cases_dir, filename)

        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(content, f, ensure_ascii=False, indent=2, default=str)

        logger.info(f"과거 사례 문서 생성: {filepath}")
        return filepath

    def load_past_cases(self) -> List[Dict[str, Any]]:
        """과거 사례 문서 로드"""
        cases = []

        for filename in os.listdir(self.past_cases_dir):
            if filename.endswith('.json'):
                filepath = os.path.join(self.past_cases_dir, filename)
                with open(filepath, 'r', encoding='utf-8') as f:
                    cases.append(json.load(f))

        return cases

    def generate_sample_past_cases(self, count: int = 5) -> List[str]:
        """샘플 과거 사례 문서 생성"""
        import random

        sample_cases = [
            {
                "case_id": "CASE_DEAD_PIXEL_2024001",
                "complaint_info": {
                    "customer": "APPLE",
                    "product_model": "OLED_67_FHD",
                    "defect_type": "DEAD_PIXEL",
                    "occurrence_date": "2024-03-15"
                },
                "analysis_result": {
                    "responsible_dept": "TFT공정",
                    "confidence": 0.92
                },
                "root_cause": "TFT Array 공정 중 ESD 발생으로 인한 TFT 손상",
                "countermeasures": [
                    "ESD 방지 장비 점검 및 교체",
                    "작업자 ESD 교육 재실시",
                    "Ionizer 성능 확인"
                ],
                "prevention_measures": [
                    "ESD 모니터링 시스템 구축",
                    "정전기 발생 위험 공정 개선"
                ],
                "result": "불량률 70% 감소",
                "keywords": ["DEAD_PIXEL", "ESD", "TFT", "Array"],
                "tags": ["TFT공정", "APPLE", "OLED"]
            },
            {
                "case_id": "CASE_MURA_2024002",
                "complaint_info": {
                    "customer": "SAMSUNG_MOBILE",
                    "product_model": "LTPO_68_QHD",
                    "defect_type": "MURA",
                    "occurrence_date": "2024-05-22"
                },
                "analysis_result": {
                    "responsible_dept": "CF공정",
                    "confidence": 0.85
                },
                "root_cause": "CF 스페이서 높이 불균일로 인한 셀갭 변동",
                "countermeasures": [
                    "스페이서 도포량 조정",
                    "CF 공정 온도 프로파일 최적화",
                    "검사 기준 강화"
                ],
                "prevention_measures": [
                    "셀갭 실시간 모니터링 도입",
                    "스페이서 자재 입고 검사 강화"
                ],
                "result": "불량률 60% 감소",
                "keywords": ["MURA", "스페이서", "셀갭", "CF"],
                "tags": ["CF공정", "SAMSUNG", "LTPO"]
            },
            {
                "case_id": "CASE_LINE_2024003",
                "complaint_info": {
                    "customer": "GOOGLE",
                    "product_model": "OLED_61_FHD",
                    "defect_type": "LINE_DEFECT",
                    "occurrence_date": "2024-07-10"
                },
                "analysis_result": {
                    "responsible_dept": "TFT공정",
                    "confidence": 0.88
                },
                "root_cause": "게이트 드라이버 IC 접합 불량",
                "countermeasures": [
                    "COG 본딩 조건 최적화",
                    "ACF 자재 변경",
                    "본딩 장비 점검"
                ],
                "prevention_measures": [
                    "본딩 품질 인라인 검사 강화",
                    "IC 자재 관리 프로세스 개선"
                ],
                "result": "불량률 80% 감소",
                "keywords": ["LINE_DEFECT", "Gate Driver", "COG", "본딩"],
                "tags": ["TFT공정", "GOOGLE", "OLED"]
            },
            {
                "case_id": "CASE_SCRATCH_2024004",
                "complaint_info": {
                    "customer": "XIAOMI",
                    "product_model": "LCD_109_2K",
                    "defect_type": "SCRATCH",
                    "occurrence_date": "2024-08-25"
                },
                "analysis_result": {
                    "responsible_dept": "CELL공정",
                    "confidence": 0.91
                },
                "root_cause": "Cell 공정 이송 중 Roller 마모로 인한 스크래치 발생",
                "countermeasures": [
                    "이송 Roller 교체",
                    "이송 속도 조정",
                    "표면 보호 필름 적용"
                ],
                "prevention_measures": [
                    "Roller 마모도 정기 점검 체계 구축",
                    "이송 장비 PM 주기 단축"
                ],
                "result": "불량률 90% 감소",
                "keywords": ["SCRATCH", "Roller", "이송", "마모"],
                "tags": ["CELL공정", "XIAOMI", "LCD"]
            },
            {
                "case_id": "CASE_TOUCH_2024005",
                "complaint_info": {
                    "customer": "META",
                    "product_model": "OLED_76_FOLD",
                    "defect_type": "TOUCH_FAIL",
                    "occurrence_date": "2024-09-30"
                },
                "analysis_result": {
                    "responsible_dept": "모듈공정",
                    "confidence": 0.87
                },
                "root_cause": "TSP 적층 공정 중 접착력 부족으로 인한 터치 불량",
                "countermeasures": [
                    "OCA 접착제 변경",
                    "적층 압력 조건 최적화",
                    "Clean room 습도 관리 강화"
                ],
                "prevention_measures": [
                    "TSP 적층 품질 검사 항목 추가",
                    "자재 변경점 관리 프로세스 수립"
                ],
                "result": "불량률 75% 감소",
                "keywords": ["TOUCH_FAIL", "TSP", "OCA", "적층"],
                "tags": ["모듈공정", "META", "FOLD"]
            }
        ]

        generated_files = []
        for case in sample_cases[:count]:
            filepath = self.generate_past_case_document(case)
            generated_files.append(filepath)

        return generated_files

    def get_report_summary(self, report_path: str) -> Dict[str, Any]:
        """보고서 요약 정보 반환"""
        summary = {
            "path": report_path,
            "filename": os.path.basename(report_path),
            "created": datetime.fromtimestamp(os.path.getctime(report_path)).isoformat(),
            "size_kb": os.path.getsize(report_path) / 1024
        }

        if report_path.endswith('.docx'):
            summary["format"] = "Word Document"
        elif report_path.endswith('.txt'):
            summary["format"] = "Text File"
        elif report_path.endswith('.json'):
            summary["format"] = "JSON"

        return summary
